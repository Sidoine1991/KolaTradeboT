from docx import Document
from docx.shared import Pt

report = {
    "title": "Audit statique — SMC_Universal.mq5",
    "meta": "Date: 2026-07-29 21:14:08 +01:00\nAuteur: Copilot (assistant)",
    "exec_summary": (
        "L'EA contient de nombreuses protections (SafeOrderSend, gates GOM, limites terminales) mais utilise "
        "aussi des appels externes WebRequest, I/O fichiers et des stubs permissifs. Risques majeurs: dépendance "
        "au serveur AI via WebRequest, erreurs réseau non toujours correctement retriées, et potentiel de conditions "
        "de course autour de l'état global. Recommandation: durcir WebRequest, ajouter retries, vérifier retcodes, "
        "sécuriser la persistence et limiter les opérations bloquantes dans OnTick."
    ),
    "findings": [
        {
            "severity": "High",
            "title": "Dépendances réseau (WebRequest) et sécurité",
            "details": (
                "Multiples appels WebRequest vers g_state.config.aiServerURL et webhooks (WhatsApp, PsychoBot, etc.). "
                "MT5 exige que chaque URL soit autorisée dans Options > Expert Advisors. Les timeouts varient; certaines "
                "fonctions assument succès (code 200) sans retry ni validation JSON robuste. Risque: délai/indisponibilité du "
                "serveur qui bloque la logique d'exécution et provoque verdicts GOM indisponibles."
            ),
            "files": ["mt5/modules/HTTPTransport.mqh", "mt5/SMC_GOM_Pipeline.mqh"]
        },
        {
            "severity": "High",
            "title": "I/O fichier et journalisation",
            "details": (
                "Utilisation de FileOpen/FileWrite dans plusieurs modules (TradeJournal, ConcordanceStore). "
                "Pas toujours de vérification d'échec d'ouverture, possible contention quand plusieurs instances ou threads "
                "accèdent aux mêmes fichiers."
            ),
            "files": ["mt5/modules/SMC_TradeJournal.mqh", "mt5/modules/SMC_ConcordanceStore.mqh"]
        },
        {
            "severity": "Medium",
            "title": "Contrôles d'ordre et retcodes",
            "details": (
                "SafeOrderSend centralise les contrôles (terminal full, GOM verdict, direction rules). C'est positif. "
                "Cependant il faut s'assurer que tous les chemins d'appel vérifient res.retcode et gèrent erreurs transitoires (ex: TRADE_RETCODE_REQUOTE, TRADE_RETCODE_REJECT)."
            ),
            "files": ["mt5/modules/SMC_Stubs.mqh", "mt5/modules/SMC_GOM_Pipeline.mqh"]
        },
        {
            "severity": "Medium",
            "title": "Stubs permissifs et état global",
            "details": (
                "Le fichier SMC_Stubs.mqh expose des implémentations temporaires qui retournent 'true' ou valeurs neutres. "
                "Si ces stubs sont utilisés en production, cela affaiblit les contrôles. L'état global (g_state, g_smcGomVerdict) "
                "introduit du couplage et risque de conditions de course."
            ),
            "files": ["mt5/modules/SMC_Stubs.mqh", "mt5/modules/TMState.mqh"]
        },
        {
            "severity": "Low",
            "title": "Magic numbers et configuration",
            "details": (
                "De nombreux inputs et hardcodes (timeouts, magic numbers, default thresholds) existent. Les rendre configurables via inputs aide la sécurité opérationnelle."
            ),
            "files": ["mt5/modules/TMState.mqh", "mt5/SMC_Universal.mq5"]
        }
    ],
    "recommendations": [
        {
            "prio": 1,
            "text": "Durcir les appels HTTP: forcer HTTPS, vérifier le code HTTP et le contenu JSON; ajouter retry exponentiel (3 essais) et circuit-breaker; journaliser échecs détaillés. Implémenter dans HTTPTransport.mqh."
        },
        {
            "prio": 1,
            "text": "Harmoniser et vérifier les retcodes après OrderSend. Ajouter retries pour erreurs transitoires et backoff, et centraliser la gestion des erreurs retournées par la plateforme."
        },
        {
            "prio": 2,
            "text": "Protéger l'I/O: vérifier handle != INVALID_HANDLE, utiliser FILE_COMMON avec noms uniques, faire FileClose, et écriture atomique via fichier temporaire + FileDelete/FileRename."
        },
        {
            "prio": 2,
            "text": "Remplacer stubs permissifs par implémentations de production ou verrouiller leur usage derrière un flag 'use_stubs' afin d'éviter comportements non sûrs en prod."
        },
        {
            "prio": 3,
            "text": "Documenter et exposer tous les magic numbers comme inputs user-modifiables; ajouter des sanity checks lors du OnInit."
        }
    ],
    "quick_fixes": [
        {
            "text": "Ajout d'un wrapper HTTP_PostWithRetry() dans HTTPTransport.mqh (3 retries, timeout configurable, validate JSON)."
        },
        {
            "text": "Vérification systématique du résultat FileOpen() et fallback si échec (log + tenter dossier Common)."
        }
    ],
    "next_steps": (
        "1) Revue ligne-par-ligne des sections d'entrée/sortie et du state machine (je peux continuer et produire un patch suggéré).\n"
        "2) Tests de compilation sur MetaTrader 5 et exécution en mode démo pour valider les chemins critiques.\n"
        "3) Optionnel: générer patchs non-destructifs et commits locaux (besoin de confirmation)."
    ),
    "appendix": {
        "scanned_files": [
            "mt5/SMC_Universal.mq5",
            "mt5/modules/HTTPTransport.mqh",
            "mt5/modules/SMC_Stubs.mqh",
            "mt5/modules/SMC_TradeJournal.mqh",
            "mt5/modules/SMC_ConcordanceStore.mqh"
        ],
        "notes": "Rapport basé sur analyse statique partielle du code dans le dépôt. Tests runtime non effectués."
    }
}

# Create Word document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

doc.add_heading(report['title'], level=1)
for line in report['meta'].split('\n'):
    p = doc.add_paragraph(line)
    p.style = 'Intense Quote'

# Executive summary
doc.add_heading('Résumé exécutif', level=2)
doc.add_paragraph(report['exec_summary'])

# Findings
doc.add_heading('Constats détaillés', level=2)
for f in report['findings']:
    doc.add_heading(f"[{f['severity']}] {f['title']}", level=3)
    doc.add_paragraph(f['details'])
    if 'files' in f and f['files']:
        doc.add_paragraph('Fichiers: ' + ', '.join(f['files']))

# Recommendations
doc.add_heading('Recommandations', level=2)
for r in report['recommendations']:
    doc.add_paragraph(f"Priorité {r['prio']}: {r['text']}")

# Quick fixes
doc.add_heading('Corrections rapides proposées', level=2)
for q in report['quick_fixes']:
    doc.add_paragraph('- ' + q['text'])

# Next steps
doc.add_heading('Étapes suivantes', level=2)
doc.add_paragraph(report['next_steps'])

# Appendix
doc.add_heading('Annexe', level=2)
doc.add_paragraph('Fichiers scannés:')
for f in report['appendix']['scanned_files']:
    doc.add_paragraph('- ' + f)

doc.add_paragraph(report['appendix']['notes'])

out_path = r"D:\Dev\TradBOT\SMC_Universal_audit.docx"
doc.save(out_path)
print('Saved', out_path)
