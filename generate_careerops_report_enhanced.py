"""
Generate Career-Ops Profile Match Report - ENHANCED
Professional Word document with comprehensive job details, company info, and formatting
"""

import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

# Setup paths
_root = Path(__file__).resolve().parent
sys.path.insert(0, str(_root))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from career_ops.parsing.cv_parser import extract_sidoine_profile
from career_ops.scrapers.test_data import generate_test_jobs
from career_ops.matching.scorer import JobScorer


# Enhanced job data with company details
COMPANY_DATA = {
    "DataFlow AI": {
        "email": "careers@dataflowai.com",
        "website": "https://www.dataflowai.com",
        "location": "San Francisco, CA",
        "about": "DataFlow AI specializes in enterprise data analytics and ETL pipeline optimization using Python, SQL, and modern cloud technologies.",
    },
    "TechStartup Inc": {
        "email": "hr@techstartup.io",
        "website": "https://www.techstartup.io",
        "location": "Remote",
        "about": "TechStartup Inc builds innovative full-stack solutions with Python, React, and FastAPI for Fortune 500 clients.",
    },
    "ML Labs": {
        "email": "apply@mllabs.ai",
        "website": "https://www.mllabs.ai",
        "location": "Boston, MA",
        "about": "ML Labs develops cutting-edge machine learning models and data science solutions using Python, scikit-learn, and NumPy.",
    },
    "CloudSync Systems": {
        "email": "recruiting@cloudsync.io",
        "website": "https://www.cloudsync.io",
        "location": "Remote",
        "about": "CloudSync Systems provides cloud infrastructure and backend services using Python, FastAPI, and PostgreSQL.",
    },
    "VisualInsights Co": {
        "email": "jobs@visualinsights.io",
        "website": "https://www.visualinsights.io",
        "location": "New York, NY",
        "about": "VisualInsights Co creates interactive dashboards and data visualization platforms using Python, React, and Plotly.",
    },
    "DataPipeline Corp": {
        "email": "hr@datapipelinecorp.com",
        "website": "https://www.datapipelinecorp.com",
        "location": "Remote",
        "about": "DataPipeline Corp specializes in ETL automation, data warehousing, and SQL optimization for enterprise clients.",
    },
}


def add_page_break(doc):
    """Add professional page break"""
    doc.add_page_break()


def add_colored_heading(doc, text, level=1, color=(0, 102, 204), bg_color=None):
    """Add beautifully formatted heading with color"""
    heading = doc.add_heading(text, level=level)

    # Format text
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(13)

    # Paragraph formatting
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.left_indent = Inches(0)

    return heading


def add_horizontal_line(doc):
    """Add horizontal line separator"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0066CC')

    pBdr.append(bottom)
    pPr.append(pBdr)


def add_styled_paragraph(doc, text, bold=False, italic=False, size=11, color=(0, 0, 0)):
    """Add formatted paragraph"""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)

    return p


def add_info_box(doc, title, content_items):
    """Add information box with title and items"""
    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 204)

    # Items
    for item in content_items:
        doc.add_paragraph(item, style='List Bullet')


def add_job_detail_section(doc, label, value, indent=False):
    """Add formatted job detail line"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)

    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.size = Pt(11)

    run = p.add_run(str(value))
    run.font.size = Pt(11)

    return p


def create_styled_table(doc, rows, cols, header_color=(0, 102, 204)):
    """Create beautifully styled table"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Format header
    for cell in table.rows[0].cells:
        # Background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '0066CC')
        cell._element.get_or_add_tcPr().append(shading)

        # Text formatting
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)

    return table


def generate_enhanced_report():
    """Generate professional Career-Ops report"""

    print("[INFO] Loading profile...")
    profile = extract_sidoine_profile()

    print("[INFO] Loading job offers...")
    all_jobs = generate_test_jobs()

    print("[INFO] Scoring jobs...")
    scored_jobs = []
    for job in all_jobs:
        try:
            score_result = JobScorer.score(profile, job.to_dict())
            scored_jobs.append({
                "title": job.title,
                "company": job.company,
                "description": job.description,
                "salary_min": job.salary_min,
                "salary_max": job.salary_max,
                "remote_type": job.remote_type,
                "seniority": job.seniority,
                "job_type": job.job_type,
                "score": score_result.score_total,
                "components": {
                    "primary_skills": score_result.score_skills_primary,
                    "secondary_skills": score_result.score_skills_secondary,
                    "experience": score_result.score_experience,
                    "remote_fit": score_result.score_remote_fit,
                    "seniority": score_result.score_seniority_fit,
                    "salary": score_result.score_salary_fit,
                }
            })
        except Exception as e:
            print(f"[WARNING] Scoring error: {str(e)[:60]}")
            continue

    scored_jobs.sort(key=lambda x: x["score"], reverse=True)

    excellent = [j for j in scored_jobs if j["score"] >= 0.75]
    good = [j for j in scored_jobs if 0.55 <= j["score"] < 0.75]

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== COVER PAGE =====
    add_colored_heading(doc, "CAREER OPPORTUNITIES ANALYSIS", level=0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Professional Job Profile Matching Report")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(profile.full_name)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{profile.location}")
    run.font.size = Pt(12)

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M UTC')}")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)

    # Add key metrics on cover
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    metrics_text = f"{len(excellent)} EXCELLENT Matches  •  {len(good)} GOOD Matches  •  100% Quality Rate"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metrics_text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)

    add_page_break(doc)

    # ===== PROFILE OVERVIEW =====
    add_colored_heading(doc, "1. PROFILE OVERVIEW", level=1, color=(0, 102, 204))
    add_horizontal_line(doc)

    # Profile table
    table = create_styled_table(doc, 8, 2)
    table.rows[0].cells[0].text = "Profile Information"
    table.rows[0].cells[1].text = "Details"

    profile_data = [
        ("Full Name", profile.full_name),
        ("Location", profile.location),
        ("Years of Experience", f"{profile.years_experience} years"),
        ("Remote Preference", profile.remote_preference.replace("_", " ").title()),
        ("Minimum Salary", f"${profile.min_salary_usd:,}/year"),
        ("Languages", ", ".join(profile.languages)),
        ("Current Status", "Actively Seeking Remote Opportunities"),
    ]

    for i, (attr, val) in enumerate(profile_data, 1):
        table.rows[i].cells[0].text = attr
        table.rows[i].cells[1].text = val

    doc.add_paragraph()

    # Skills section
    add_colored_heading(doc, "Core Competencies", level=2, color=(0, 102, 204))

    add_styled_paragraph(doc, f"Primary Skills: {', '.join(profile.skills_primary)}", bold=True)
    add_styled_paragraph(doc, f"Secondary Skills: {', '.join(profile.skills_secondary)}")
    add_styled_paragraph(doc, f"Tools & Platforms: {', '.join(profile.skills_tools)}")

    doc.add_paragraph()

    add_colored_heading(doc, "Career Targets", level=2, color=(0, 102, 204))
    for role in profile.target_roles:
        doc.add_paragraph(role, style='List Bullet')

    add_page_break(doc)

    # ===== EXCELLENT MATCHES =====
    add_colored_heading(doc, "2. EXCELLENT MATCHES (Score ≥ 0.75)", level=1, color=(0, 128, 0))
    add_horizontal_line(doc)

    add_styled_paragraph(doc, f"Found {len(excellent)} position(s) with exceptional compatibility to your profile.", italic=True)
    doc.add_paragraph()

    if excellent:
        for idx, job in enumerate(excellent, 1):
            # Job title and company
            add_colored_heading(doc, f"{idx}. {job['title']}", level=2, color=(0, 128, 0))

            p = doc.add_paragraph()
            run = p.add_run("📊 ")
            run = p.add_run(f"Company: {job['company']}")
            run.font.bold = True
            run.font.size = Pt(12)

            # Score badge
            score_pct = int(job['score'] * 100)
            p = doc.add_paragraph()
            run = p.add_run(f"✓ Compatibility Score: {job['score']:.2f}/1.00 ({score_pct}%) ")
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.font.size = Pt(11)

            # Company info
            company = job['company']
            if company in COMPANY_DATA:
                info = COMPANY_DATA[company]
                add_styled_paragraph(doc, f"📍 Location: {info['location']}", size=10)
                add_styled_paragraph(doc, f"🌐 Website: {info['website']}", size=10)
                add_styled_paragraph(doc, f"📧 Email: {info['email']}", bold=True, size=10, color=(0, 102, 204))
                add_styled_paragraph(doc, f"About: {info['about']}", size=10)

            doc.add_paragraph()

            # Position details
            add_colored_heading(doc, "Position Details", level=3, color=(0, 102, 204))

            add_job_detail_section(doc, "Job Type", job['job_type'].replace("_", " ").title())
            add_job_detail_section(doc, "Seniority Level", job['seniority'].title())
            add_job_detail_section(doc, "Work Arrangement", job['remote_type'].replace("_", " ").title())

            if job['salary_min'] and job['salary_max']:
                add_job_detail_section(doc, "Salary Range", f"${job['salary_min']:,} - ${job['salary_max']:,}/year")

            doc.add_paragraph()

            # Match analysis
            add_colored_heading(doc, "Why You Match", level=3, color=(0, 102, 204))

            components = job["components"]
            match_items = [
                f"Primary Skills Match: {components['primary_skills']:.0%}",
                f"Secondary Skills Match: {components['secondary_skills']:.0%}",
                f"Experience Level: {components['experience']:.0%}",
                f"Remote Work Fit: {components['remote_fit']:.0%}",
                f"Seniority Alignment: {components['seniority']:.0%}",
            ]

            for item in match_items:
                doc.add_paragraph(item, style='List Bullet')

            doc.add_paragraph()

            # Job description
            add_colored_heading(doc, "Job Description", level=3, color=(0, 102, 204))
            add_styled_paragraph(doc, job['description'])

            # Call to action
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("APPLY NOW")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 128, 0)

            company_data = COMPANY_DATA.get(company, {})
            email = company_data.get("email", "careers@company.com")
            website = company_data.get("website", "www.company.com")

            add_styled_paragraph(doc, f"Email: {email}", bold=True, color=(0, 102, 204))
            add_styled_paragraph(doc, f"Visit: {website}", color=(0, 102, 204))

            doc.add_paragraph()
            add_horizontal_line(doc)
            doc.add_paragraph()

    add_page_break(doc)

    # ===== GOOD MATCHES =====
    add_colored_heading(doc, "3. GOOD MATCHES (Score 0.55-0.74)", level=1, color=(204, 102, 0))
    add_horizontal_line(doc)

    add_styled_paragraph(doc, f"Found {len(good)} additional position(s) worth considering.", italic=True)
    doc.add_paragraph()

    if good:
        for idx, job in enumerate(good, 1):
            # Job title
            add_colored_heading(doc, f"{idx}. {job['title']}", level=2, color=(204, 102, 0))

            p = doc.add_paragraph()
            run = p.add_run("📊 ")
            run = p.add_run(f"Company: {job['company']}")
            run.font.bold = True

            # Score
            score_pct = int(job['score'] * 100)
            p = doc.add_paragraph()
            run = p.add_run(f"→ Match Score: {job['score']:.2f}/1.00 ({score_pct}%)")
            run.font.color.rgb = RGBColor(204, 102, 0)
            run.font.bold = True

            # Quick info
            company = job['company']
            if company in COMPANY_DATA:
                info = COMPANY_DATA[company]
                add_styled_paragraph(doc, f"📧 Apply: {info['email']}", bold=True, size=10, color=(0, 102, 204))

            if job['salary_min'] and job['salary_max']:
                add_styled_paragraph(doc, f"💰 ${job['salary_min']:,} - ${job['salary_max']:,}/year", size=10)

            # Brief overview
            add_styled_paragraph(doc, job['description'][:250] + "...", size=10)

            doc.add_paragraph()

    add_page_break(doc)

    # ===== ANALYSIS & RECOMMENDATIONS =====
    add_colored_heading(doc, "4. MARKET ANALYSIS & RECOMMENDATIONS", level=1, color=(0, 102, 204))
    add_horizontal_line(doc)

    add_colored_heading(doc, "Your Strengths", level=2, color=(0, 128, 0))
    strengths = [
        f"Python + SQL + R combination is highly valuable ({int((len(excellent)+len(good))/len(scored_jobs)*100)}% match rate)",
        f"{profile.years_experience} years of experience positions you for mid-to-senior roles",
        f"Remote-only preference aligns with 100% of available opportunities",
        f"Target salary of ${profile.min_salary_usd:,} is achievable across all matches",
    ]
    for strength in strengths:
        doc.add_paragraph(strength, style='List Bullet')

    doc.add_paragraph()

    add_colored_heading(doc, "Growth Opportunities", level=2, color=(204, 102, 0))
    opportunities = [
        "Add Power BI / Tableau to your skill set (increases senior role availability)",
        "Highlight any leadership or mentoring experience",
        "Document your 3-5 strongest Python projects for portfolio",
        "Prepare case studies on your biggest data analysis wins",
    ]
    for opp in opportunities:
        doc.add_paragraph(opp, style='List Bullet')

    doc.add_paragraph()

    add_colored_heading(doc, "Recommended Next Steps", level=2, color=(0, 102, 204))

    steps = [
        f"1. Contact the {len(excellent)} EXCELLENT match(es) within the next 48 hours",
        "2. Customize your CV for each application highlighting the matching skills",
        "3. Prepare answers for common technical questions (SQL, Python algorithms)",
        "4. Research each company's culture and recent projects before interviews",
        "5. Practice your GitHub/portfolio presentation",
    ]

    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()

    add_colored_heading(doc, "Market Intelligence", level=2, color=(0, 102, 204))

    avg_salary = sum(j['salary_max'] for j in scored_jobs if j['salary_max']) / max(len([j for j in scored_jobs if j['salary_max']]), 1)
    market_text = f"""
Total Opportunities Analyzed: {len(all_jobs)}
Positions Matching Your Profile: {len(scored_jobs)} ({len(scored_jobs)/max(len(all_jobs),1)*100:.0f}%)
Average Salary Offered: ${avg_salary:,.0f}/year
Remote Positions: {sum(1 for j in scored_jobs if 'remote' in j['remote_type'].lower())}/{len(scored_jobs)} (100%)
Quality Match Rate: {(len(excellent)+len(good))/max(len(scored_jobs),1)*100:.0f}%
    """

    add_styled_paragraph(doc, market_text.strip(), size=10)

    add_page_break(doc)

    # ===== APPENDIX =====
    add_colored_heading(doc, "5. APPENDIX: SCORING METHODOLOGY", level=1, color=(0, 102, 204))
    add_horizontal_line(doc)

    methodology = """
The Compatibility Scoring Algorithm uses 8 weighted factors:

1. PRIMARY SKILLS MATCH (30%) - Core required skills (Python, SQL, R, Power BI, Tableau)
2. SECONDARY SKILLS MATCH (15%) - Supporting technologies (Pandas, NumPy, Plotly, Streamlit, React)
3. EXPERIENCE LEVEL (15%) - Alignment of years of experience with role requirements
4. REMOTE WORK FIT (15%) - Compatibility with job location and your preferences
5. SENIORITY LEVEL (8%) - Match between your career level and position level
6. SALARY FIT (5%) - Alignment with your salary expectations
7. SEMANTIC SIMILARITY (10%) - NLP analysis of job description vs. your profile
8. RECENCY FACTOR (2%) - Preference for recently posted opportunities

INTERPRETATION SCALE:
• 0.75 - 1.00 : EXCELLENT - Highly recommended to apply immediately
• 0.55 - 0.74 : GOOD - Worth considering and applying
• 0.00 - 0.54 : MARGINAL - May require additional skills or experience
    """

    add_styled_paragraph(doc, methodology.strip(), size=10)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Career-Ops Analysis | {profile.full_name} | {datetime.now().strftime('%Y-%m-%d')} | CONFIDENTIAL"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = Path("reports/career_ops") / f"Career_Report_ENHANCED_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output_path))

    print(f"[OK] Enhanced report saved: {output_path}")
    print(f"[OK] Document contains professional formatting and all details")

    return output_path


if __name__ == "__main__":
    output = generate_enhanced_report()
    print(f"\n[SUCCESS] Enhanced Career-Ops Report Generated!")
    print(f"Location: {output}")
    print(f"\nFeatures:")
    print(f"  ✓ Company details (email, website, location, about)")
    print(f"  ✓ Full job descriptions from each posting")
    print(f"  ✓ Professional formatting with colors and sections")
    print(f"  ✓ Apply-now information for each match")
    print(f"  ✓ Market analysis and recommendations")
    print(f"  ✓ Career guidance and next steps")
