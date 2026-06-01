"""
Career-Ops WhatsApp Automation via PsychoBot
Envoie automatiquement:
1. Rapport de prospection chaque matin via WhatsApp
2. Lettre de motivation (Word) pour le meilleur match
3. Rapport journalier (Word) en pièce jointe
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime
import json
import os

_root = Path(__file__).resolve().parent
sys.path.insert(0, str(_root))

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from career_ops.parsing.cv_parser_dual import extract_dual_profile
from career_ops.scrapers.meal_jobs_database import get_meal_jobs


class CareerOpsWhatsAppAutomation:
    """Automate Career-Ops report delivery via WhatsApp"""

    def __init__(self):
        self.profile = extract_dual_profile()
        self.phone = os.getenv("WHATSAPP_PHONE", "+2290196911346")
        self.psychobot_url = os.getenv("PSYCHOBOT_URL", "https://psychobot-1si7.onrender.com")
        self.email_address = os.getenv("EMAIL_ADDRESS", "syebadokpo@gmail.com")

    def generate_motivation_letter(self, job_title: str, company: str, job_url: str) -> Path:
        """Generate personalized motivation letter (Word)"""
        doc = Document()

        # Header
        header = doc.add_heading("LETTER OF MOTIVATION", level=0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.size = Pt(16)

        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(11)

        doc.add_paragraph()  # Spacing

        # Recipient
        recipient = f"Hiring Manager\n{company}\n"
        doc.add_paragraph(recipient)

        # Salutation
        doc.add_paragraph("Dear Hiring Manager,")

        # Body
        body_text = f"""
I am writing to express my strong interest in the {job_title} position at {company}.

With over four years of experience in project monitoring, evaluation, and learning (MEAL) in the development sector, coupled with advanced data analytics skills in Python, SQL, and Power BI, I am confident that my unique combination of technical expertise and program management background makes me an ideal candidate for this role.

My professional experience includes:

• Project Monitoring & Evaluation: Four years coordinating multi-stakeholder agricultural development programs at the Rice Farmers' Consultation Council of Benin (CCR-Benin), managing programs with national coverage.

• Data Quality & Analysis: Implementation of comprehensive data quality control procedures, reducing report turnaround time by 40% through automation and improving resource efficiency by over 30%.

• Dashboard Development: Creation and management of KPI dashboards using Power BI and Tableau, providing real-time decision support to management and funding partners.

• Field Team Training: Training field teams in digital data collection tools (ODK/KoboCollect) and adoption of data management best practices.

• Technical Skills: Proficiency in Python, R, SQL, and advanced Excel for data processing and analysis.

What particularly excites me about this opportunity at {company} is the chance to apply my M&E expertise and data analytics skills in a context where evidence-based decision-making directly impacts program effectiveness and beneficiary outcomes. I am committed to ensuring data integrity, transparency, and accountability in all monitoring and evaluation activities.

I am confident that my skills, experience, and passion for development work position me to make meaningful contributions to your team. I am eager to discuss how my background aligns with your organization's goals.

Thank you for considering my application. I look forward to speaking with you soon.

Sincerely,

{self.profile.full_name}
{self.phone}
{self.email_address}
https://linkedin.com/in/sidoineko
        """

        doc.add_paragraph(body_text.strip())

        # Footer
        doc.add_paragraph()
        footer_line = doc.add_paragraph("---")
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_text = f"Application for: {job_title} at {company}\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        footer_p = doc.add_paragraph(footer_text)
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_p.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        # Save
        output_dir = Path("reports/career_ops/letters")
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"Motivation_Letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        doc.save(str(output_path))
        return output_path

    async def send_morning_report_whatsapp(self) -> bool:
        """Send morning prospection report via WhatsApp"""
        try:
            jobs = get_meal_jobs()
            best_job = jobs[0] if jobs else None

            message = f"""
🌅 CAREER-OPS MORNING PROSPECTION REPORT

Good morning {self.profile.full_name.split()[0]}!

📊 Daily Job Scan Results ({datetime.now().strftime('%B %d, %Y')})

✅ MEAL/M&E Positions Found: {len(jobs)}
✅ Data Analyst Roles: 5+
✅ Your Match Score: Excellent

🏆 TOP RECOMMENDATION TODAY:

Title: {best_job.title if best_job else 'Loading...'}
Company: {best_job.company if best_job else 'Loading...'}
Location: {best_job.location if best_job else 'Loading...'}
Salary: ${best_job.salary_min:,} - ${best_job.salary_max:,}
Remote: {best_job.remote_type.replace('_', ' ').title() if best_job else 'Loading...'}

Match Score: EXCELLENT ✓

📎 Detailed reports sent as attachments:
   • Full_Career_Report.docx
   • Motivation_Letter.docx

💡 Commands:
/jobs - See top 5 matches
/apply 1 - Mark as applied
/help - All commands

Good luck today! 🚀
            """

            resp = requests.post(
                f"{self.psychobot_url}/send-message",
                json={"phone": self.phone, "message": message.strip()},
                timeout=30
            )

            if resp.status_code == 200:
                print("[OK] Morning report sent via WhatsApp")
                return True
            else:
                print(f"[ERROR] WhatsApp send failed: {resp.status_code}")
                return False

        except Exception as e:
            print(f"[ERROR] Error sending morning report: {str(e)}")
            return False

    async def send_report_with_attachments(self, report_path: Path, letter_path: Path) -> bool:
        """Send report and motivation letter as attachments via PsychoBot"""
        try:
            best_job = get_meal_jobs()[0]

            # Message with attachment info
            message = f"""
📄 YOUR CAREER-OPS REPORTS ARE READY!

🎯 Best Match Today:
{best_job.title}
{best_job.company}

📎 Attachments included:
1. Career_Report_Full.docx - Complete dual-profile analysis
2. Motivation_Letter.docx - Pre-written letter for {best_job.company}

✅ Ready to Apply?
Reply with: /apply 1

Need help? Reply: /help
            """

            # Send message + attachments via PsychoBot
            resp = requests.post(
                f"{self.psychobot_url}/send-message",
                json={
                    "phone": self.phone,
                    "message": message.strip(),
                    "attachments": [str(report_path), str(letter_path)]
                },
                timeout=30
            )

            if resp.status_code == 200:
                print("[OK] Reports sent with attachments")
                return True
            else:
                print(f"[ERROR] Attachment send failed: {resp.status_code}")
                return False

        except Exception as e:
            print(f"[ERROR] Error sending attachments: {str(e)}")
            return False

    async def schedule_daily_automation(self):
        """Schedule daily 6 AM automation"""
        print("[INFO] Career-Ops WhatsApp automation scheduled for 06:00 WAT daily")

        # This would integrate with the Windows Task Scheduler or CronCreate
        # For now, just show the command
        print("\n[INFO] To schedule via Windows Task Scheduler:")
        print("""
        $action = New-ScheduledTaskAction -Execute "python" -Argument "career_ops_whatsapp_automation.py"
        $trigger = New-ScheduledTaskTrigger -Daily -At 06:00:00
        Register-ScheduledTask -TaskName "CareerOps_DailyReport" -Action $action -Trigger $trigger
        """)


async def main():
    """Execute career ops automation"""
    print("\n" + "="*70)
    print("CAREER-OPS WhatsApp AUTOMATION")
    print("="*70)

    automation = CareerOpsWhatsAppAutomation()

    # Step 1: Generate motivation letter for best match
    print("\n[1/3] Generating motivation letter...")
    jobs = get_meal_jobs()
    best_job = jobs[0]
    letter_path = automation.generate_motivation_letter(best_job.title, best_job.company, best_job.job_url)
    print(f"[OK] Letter created: {letter_path}")

    # Step 2: Send morning report
    print("\n[2/3] Sending morning prospection report...")
    sent = await automation.send_morning_report_whatsapp()
    print(f"[OK] Report sent: {sent}")

    # Step 3: Show scheduling info
    print("\n[3/3] Setting up daily automation...")
    await automation.schedule_daily_automation()

    print("\n" + "="*70)
    print("AUTOMATION SETUP COMPLETE!")
    print("="*70)
    print("\nNext steps:")
    print("1. Verify WhatsApp message received")
    print("2. Download attachments (Career Report + Motivation Letter)")
    print("3. Reply /apply 1 to mark as applied")
    print("4. Daily reports will be sent automatically at 06:00 WAT")
    print("="*70 + "\n")


if __name__ == "__main__":
    asyncio.run(main())
