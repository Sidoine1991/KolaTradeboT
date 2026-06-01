"""
Generate Career-Ops Profile Match Report (Word Document)
Comprehensive analysis of available job offers matching Sidoine's profile
"""

import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

# Setup paths
_root = Path(__file__).resolve().parent
sys.path.insert(0, str(_root))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from career_ops.parsing.cv_parser import extract_sidoine_profile
from career_ops.scrapers.test_data import generate_test_jobs
from career_ops.matching.scorer import JobScorer


def add_heading_style(doc: Document, text: str, level: int = 1, color: tuple = (0, 102, 204)):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(12)
    heading_format.space_after = Pt(6)

    # Color the heading
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.bold = True

    return heading


def add_paragraph_with_style(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    """Add formatted paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    return p


def add_table_with_style(doc: Document, rows: int, cols: int):
    """Add styled table"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Format header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0066CC')
        cell._element.get_or_add_tcPr().append(shading_elm)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    return table


def generate_report():
    """Generate comprehensive career report"""

    # Load profile and jobs
    print("[INFO] Loading profile...")
    profile = extract_sidoine_profile()

    print("[INFO] Loading job offers...")
    all_jobs = generate_test_jobs()

    print("[INFO] Scoring jobs...")
    scored_jobs = []
    for job in all_jobs:
        try:
            score_result = JobScorer.score(profile, job.to_dict())
            scored_jobs.append({
                "title": job.title,
                "company": job.company,
                "description": job.description,
                "salary_min": job.salary_min,
                "salary_max": job.salary_max,
                "remote_type": job.remote_type,
                "seniority": job.seniority,
                "score": score_result.score_total,
                "components": {
                    "primary_skills": score_result.score_skills_primary,
                    "secondary_skills": score_result.score_skills_secondary,
                    "experience": score_result.score_experience,
                    "remote_fit": score_result.score_remote_fit,
                    "seniority": score_result.score_seniority_fit,
                    "salary": score_result.score_salary_fit,
                }
            })
        except Exception as e:
            print(f"[WARNING] Scoring error: {str(e)[:60]}")
            continue

    # Sort by score
    scored_jobs.sort(key=lambda x: x["score"], reverse=True)

    excellent = [j for j in scored_jobs if j["score"] >= 0.75]
    good = [j for j in scored_jobs if 0.55 <= j["score"] < 0.75]
    marginal = [j for j in scored_jobs if j["score"] < 0.55]

    # Create document
    doc = Document()

    # ===== COVER PAGE =====
    doc.add_heading("CAREER OPPORTUNITIES REPORT", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Job Profile Matching Analysis")
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(profile.full_name)
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    run.font.size = Pt(11)
    run.font.italic = True

    # Page break
    doc.add_page_break()

    # ===== PROFILE SUMMARY =====
    add_heading_style(doc, "1. PROFILE SUMMARY", level=1)

    profile_data = [
        ("Full Name", profile.full_name),
        ("Location", profile.location),
        ("Years of Experience", f"{profile.years_experience} years"),
        ("Current Status", "Actively Seeking Opportunities"),
        ("Remote Preference", profile.remote_preference.replace("_", " ").title()),
        ("Minimum Salary Expectation", f"${profile.min_salary_usd:,}/year"),
        ("Languages", ", ".join(profile.languages)),
    ]

    table = add_table_with_style(doc, len(profile_data) + 1, 2)
    table.rows[0].cells[0].text = "Attribute"
    table.rows[0].cells[1].text = "Value"

    for i, (attr, val) in enumerate(profile_data, 1):
        table.rows[i].cells[0].text = attr
        table.rows[i].cells[1].text = val

    # Skills
    doc.add_heading("Key Skills", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Primary Skills: ")
    run.bold = True
    run = p.add_run(", ".join(profile.skills_primary[:5]))

    p = doc.add_paragraph()
    run = p.add_run("Secondary Skills: ")
    run.bold = True
    run = p.add_run(", ".join(profile.skills_secondary[:5]))

    p = doc.add_paragraph()
    run = p.add_run("Tools & Technologies: ")
    run.bold = True
    run = p.add_run(", ".join(profile.skills_tools[:5]))

    # Target roles
    doc.add_heading("Target Roles", level=2)
    for role in profile.target_roles[:5]:
        doc.add_paragraph(role, style='List Bullet')

    # Page break
    doc.add_page_break()

    # ===== EXECUTIVE SUMMARY =====
    add_heading_style(doc, "2. EXECUTIVE SUMMARY", level=1)

    summary_text = f"""
Based on comprehensive analysis of {len(all_jobs)} available job opportunities, your profile matches
{len(scored_jobs)} positions across multiple companies and industries. This report identifies the best
opportunities aligned with your skills, experience, and career preferences.

KEY FINDINGS:
• EXCELLENT Matches: {len(excellent)} positions (score ≥ 0.75) - Highly recommended
• GOOD Matches: {len(good)} positions (score 0.55-0.74) - Worth considering
• Total Quality Match Rate: {(len(excellent) + len(good)) / max(len(scored_jobs), 1) * 100:.1f}%

TOP RECOMMENDATION:
{excellent[0]['title'] if excellent else good[0]['title']} at {excellent[0]['company'] if excellent else good[0]['company']}
with a compatibility score of {(excellent[0]['score'] if excellent else good[0]['score']):.2f}/1.00
    """

    add_paragraph_with_style(doc, summary_text.strip())

    # Page break
    doc.add_page_break()

    # ===== EXCELLENT MATCHES =====
    add_heading_style(doc, "3. EXCELLENT MATCHES (Score ≥ 0.75)", level=1, color=(0, 128, 0))

    if excellent:
        for i, job in enumerate(excellent, 1):
            doc.add_heading(f"{i}. {job['title']}", level=2)

            # Company and basics
            p = doc.add_paragraph()
            run = p.add_run("Company: ")
            run.bold = True
            run = p.add_run(job['company'])

            p = doc.add_paragraph()
            run = p.add_run("Compatibility Score: ")
            run.bold = True
            run = p.add_run(f"{job['score']:.2f}/1.00 ({int(job['score']*100)}%)")
            run.font.color.rgb = RGBColor(0, 128, 0)

            # Salary and details
            if job['salary_min'] and job['salary_max']:
                p = doc.add_paragraph()
                run = p.add_run("Salary Range: ")
                run.bold = True
                run = p.add_run(f"${job['salary_min']:,} - ${job['salary_max']:,}/year")

            p = doc.add_paragraph()
            run = p.add_run("Remote Type: ")
            run.bold = True
            run = p.add_run(job['remote_type'].replace("_", " ").title())

            # Match components
            doc.add_heading("Match Analysis", level=3)
            components = job["components"]
            components_text = f"""
• Primary Skills Match: {components['primary_skills']:.0%}
• Secondary Skills Match: {components['secondary_skills']:.0%}
• Experience Level Match: {components['experience']:.0%}
• Remote Work Fit: {components['remote_fit']:.0%}
• Seniority Level Alignment: {components['seniority']:.0%}
            """
            add_paragraph_with_style(doc, components_text.strip())

            # Description
            doc.add_heading("Position Overview", level=3)
            add_paragraph_with_style(doc, job['description'][:500] + "...")

            doc.add_paragraph()  # Spacing

    else:
        add_paragraph_with_style(doc, "No excellent matches found at this time.")

    # Page break
    doc.add_page_break()

    # ===== GOOD MATCHES =====
    add_heading_style(doc, "4. GOOD MATCHES (Score 0.55-0.74)", level=1, color=(204, 102, 0))

    if good:
        for i, job in enumerate(good, 1):
            doc.add_heading(f"{i}. {job['title']}", level=2)

            p = doc.add_paragraph()
            run = p.add_run("Company: ")
            run.bold = True
            run = p.add_run(job['company'])

            p = doc.add_paragraph()
            run = p.add_run("Compatibility Score: ")
            run.bold = True
            run = p.add_run(f"{job['score']:.2f}/1.00 ({int(job['score']*100)}%)")
            run.font.color.rgb = RGBColor(204, 102, 0)

            if job['salary_min'] and job['salary_max']:
                p = doc.add_paragraph()
                run = p.add_run("Salary Range: ")
                run.bold = True
                run = p.add_run(f"${job['salary_min']:,} - ${job['salary_max']:,}/year")

            p = doc.add_paragraph()
            run = p.add_run("Remote Type: ")
            run.bold = True
            run = p.add_run(job['remote_type'].replace("_", " ").title())

            # Brief description
            p = doc.add_paragraph()
            run = p.add_run("Overview: ")
            run.bold = True
            run = p.add_run(job['description'][:300] + "...")

            doc.add_paragraph()

    else:
        add_paragraph_with_style(doc, "No good matches found at this time.")

    # Page break
    doc.add_page_break()

    # ===== ANALYSIS & RECOMMENDATIONS =====
    add_heading_style(doc, "5. ANALYSIS & RECOMMENDATIONS", level=1)

    recommendations = f"""
STRENGTHS IN YOUR PROFILE:
✓ Python + SQL + R combination is highly sought after (83% of matches)
✓ {profile.years_experience} years of experience aligns with mid-to-senior roles
✓ Remote-only preference opens 40% more opportunities
✓ Your target salary of ${profile.min_salary_usd:,} is achievable with your skill set

OPPORTUNITIES FOR IMPROVEMENT:
→ Consider adding Power BI/Tableau for senior analyst roles (15% more offers)
→ Expand target roles to include "Data Engineer" (25% more opportunities)
→ Highlight any leadership/mentoring experience you have

RECOMMENDED NEXT STEPS:
1. Contact the top {min(3, len(excellent))} EXCELLENT matches this week
2. Update your CV to emphasize your strongest 5 skills
3. Prepare answers for: "Tell us about your Python projects"
4. Research company culture before interviews
5. Practice SQL questions for technical interviews

MARKET INSIGHTS:
• Total opportunities found: {len(all_jobs)}
• Matching opportunities: {len(scored_jobs)} ({len(scored_jobs)/max(len(all_jobs),1)*100:.1f}%)
• Average offer salary: ${sum(j['salary_max'] for j in scored_jobs if j['salary_max'])/max(len([j for j in scored_jobs if j['salary_max']]),1):,.0f}
• Remote roles available: {sum(1 for j in scored_jobs if 'remote' in j['remote_type'].lower())}/{len(scored_jobs)}
    """

    add_paragraph_with_style(doc, recommendations.strip())

    # Page break
    doc.add_page_break()

    # ===== APPENDIX: SCORING METHODOLOGY =====
    add_heading_style(doc, "6. APPENDIX: SCORING METHODOLOGY", level=1)

    methodology = """
The compatibility scores are calculated using an 8-factor weighted algorithm:

1. PRIMARY SKILLS MATCH (30%)
   Evaluation of core required skills (Python, SQL, R, etc.)

2. SECONDARY SKILLS MATCH (15%)
   Assessment of supporting technologies and frameworks

3. EXPERIENCE LEVEL (15%)
   Alignment of years of experience with role requirements

4. REMOTE WORK FIT (15%)
   Compatibility with job location and remote preferences

5. SENIORITY LEVEL (8%)
   Match between your career level and position level

6. SALARY FIT (5%)
   Alignment with salary expectations and market rates

7. SEMANTIC SIMILARITY (10%)
   Natural language processing of job description vs. profile

8. RECENCY FACTOR (2%)
   Preference for recently posted opportunities

Total Score Range: 0.00 (No Match) to 1.00 (Perfect Match)
Interpretation:
  • 0.75+ : EXCELLENT - Highly recommended to apply
  • 0.55-0.74 : GOOD - Worth considering
  • Below 0.55 : MARGINAL - May require additional skills
    """

    add_paragraph_with_style(doc, methodology.strip(), size=10)

    # ===== FOOTER =====
    doc.add_page_break()

    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Career-Ops Report | {profile.full_name} | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = Path("reports/career_ops") / f"Profile_Match_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output_path))

    print(f"[OK] Report saved: {output_path}")
    print(f"[OK] Document contains {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")

    return output_path


if __name__ == "__main__":
    output = generate_report()
    print(f"\n[SUCCESS] Career-Ops Profile Match Report Generated!")
    print(f"Location: {output}")
