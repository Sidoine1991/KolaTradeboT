"""
Génère rapport Word professionnel des résultats backtest GoldSMC
Usage: python Python/generate_backtest_word_report.py "chemin/backtest.xlsx"
"""

import sys
import io
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import the analysis function
from analyze_goldsmc_backtest import extract_metrics_from_excel


def set_cell_background(cell, color):
    """Définit couleur de fond cellule."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_styled_heading(doc, text, level=1, color=None):
    """Ajoute titre stylé."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_metric_table(doc, metrics, targets):
    """Ajoute tableau comparatif métriques vs objectifs."""

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # En-têtes
    headers = table.rows[0].cells
    headers[0].text = 'Métrique'
    headers[1].text = 'Résultat'
    headers[2].text = 'Objectif'
    headers[3].text = 'Statut'

    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Données
    metric_configs = [
        ('Profit Factor', metrics.get('profit_factor', 0), targets['profit_factor'], 'greater'),
        ('Win Rate (%)', metrics.get('win_rate', 0), targets['win_rate'], 'greater'),
        ('Max Drawdown (%)', metrics.get('max_dd_pct', 0), targets['max_dd'], 'less'),
        ('Recovery Factor', metrics.get('recovery_factor', 0), targets['recovery_factor'], 'greater'),
        ('Net Profit ($)', metrics.get('net_profit', 0), 0, 'greater'),
        ('Total Trades', metrics.get('total_trades', 0), 20, 'greater'),
    ]

    for metric_name, value, target, comparison in metric_configs:
        row = table.add_row().cells
        row[0].text = metric_name

        # Formatage valeur
        if 'Rate' in metric_name or 'Drawdown' in metric_name:
            row[1].text = f"{value:.2f}%"
        elif 'Profit' in metric_name and '$' in metric_name:
            row[2].text = f"${value:,.2f}"
        elif 'Factor' in metric_name:
            row[1].text = f"{value:.2f}"
        else:
            row[1].text = str(int(value))

        # Target
        if 'Rate' in metric_name or 'Drawdown' in metric_name:
            row[2].text = f"{'≥' if comparison == 'greater' else '≤'} {target}%"
        elif 'Factor' in metric_name:
            row[2].text = f"≥ {target}"
        elif 'Trades' in metric_name:
            row[2].text = f"≥ {target}"
        else:
            row[2].text = "N/A"

        # Statut
        if comparison == 'greater':
            passed = value >= target
        else:
            passed = value <= target

        if 'Trades' in metric_name:
            passed = value >= target
        elif 'Profit' in metric_name and '$' in metric_name:
            passed = value > 0

        row[3].text = '✅ OK' if passed else '❌ ÉCHEC'

        # Couleur statut
        if passed:
            set_cell_background(row[3], 'C6EFCE')
        else:
            set_cell_background(row[3], 'FFC7CE')


def generate_word_report(excel_path: str, output_path: str = None):
    """Génère rapport Word complet."""

    print(f"📊 Génération rapport Word...")
    print(f"   Source: {excel_path}")

    # Extraire métriques
    metrics = extract_metrics_from_excel(excel_path)

    if not metrics:
        print("❌ Impossible d'extraire les métriques")
        return None

    # Détecter régime probable basé sur période
    period_start = metrics.get('period_start', '')
    period_end = metrics.get('period_end', '')

    if '2022' in period_start or '2023' in period_start or '2024' in period_start:
        regime = 'BEAR'
        targets = {
            'profit_factor': 2.0,
            'win_rate': 50.0,
            'max_dd': 20.0,
            'recovery_factor': 2.5
        }
    elif '2017' in period_start or '2019' in period_start or '2020' in period_start:
        regime = 'BULL'
        targets = {
            'profit_factor': 5.0,
            'win_rate': 55.0,
            'max_dd': 20.0,
            'recovery_factor': 5.0
        }
    else:
        regime = 'TRANSITION'
        targets = {
            'profit_factor': 1.8,
            'win_rate': 52.0,
            'max_dd': 15.0,
            'recovery_factor': 2.0
        }

    # Créer document
    doc = Document()

    # Style document
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # En-tête
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = header_para.add_run('📊 RAPPORT BACKTEST GOLDSMC V5\n')
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(68, 114, 196)

    subtitle_run = header_para.add_run(f'Régime: {regime} Market Analysis')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(112, 112, 112)

    doc.add_paragraph()

    # Informations générales
    add_styled_heading(doc, '1. INFORMATIONS GÉNÉRALES', level=1, color=(68, 114, 196))

    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Light Grid Accent 1'

    info_data = [
        ('Symbole', metrics.get('symbol', 'XAUUSD')),
        ('Expert Advisor', metrics.get('expert', 'GoldSMC_EA_v5')),
        ('Période', f"{period_start} → {period_end}"),
        ('Dépôt initial', f"${metrics.get('initial_deposit', 500):,.2f}"),
        ('Magic Number', str(metrics.get('magic_number', 'N/A'))),
        ('Total Trades', str(int(metrics.get('total_trades', 0)))),
        ('Rapport généré', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row[0], 'E7E6E6')
        row[1].text = value

    doc.add_paragraph()

    # Résultats vs Objectifs
    add_styled_heading(doc, '2. RÉSULTATS VS OBJECTIFS', level=1, color=(68, 114, 196))

    doc.add_paragraph(f"Objectifs {regime} Market:", style='List Bullet')
    doc.add_paragraph(f"• Profit Factor ≥ {targets['profit_factor']}", style='List Bullet 2')
    doc.add_paragraph(f"• Win Rate ≥ {targets['win_rate']}%", style='List Bullet 2')
    doc.add_paragraph(f"• Max Drawdown < {targets['max_dd']}%", style='List Bullet 2')
    doc.add_paragraph(f"• Recovery Factor ≥ {targets['recovery_factor']}", style='List Bullet 2')

    doc.add_paragraph()
    add_metric_table(doc, metrics, targets)

    doc.add_paragraph()

    # Analyse détaillée
    add_styled_heading(doc, '3. ANALYSE DÉTAILLÉE', level=1, color=(68, 114, 196))

    # Performance
    add_styled_heading(doc, '3.1 Performance Globale', level=2)

    perf_para = doc.add_paragraph()
    perf_para.add_run(f"Net Profit: ").font.bold = True
    profit_run = perf_para.add_run(f"${metrics.get('net_profit', 0):,.2f}")
    profit_run.font.color.rgb = RGBColor(0, 176, 80) if metrics.get('net_profit', 0) > 0 else RGBColor(255, 0, 0)
    profit_run.font.bold = True

    perf_para = doc.add_paragraph()
    perf_para.add_run(f"Gross Profit: ").font.bold = True
    perf_para.add_run(f"${metrics.get('gross_profit', 0):,.2f}")

    perf_para = doc.add_paragraph()
    perf_para.add_run(f"Gross Loss: ").font.bold = True
    perf_para.add_run(f"${metrics.get('gross_loss', 0):,.2f}")

    perf_para = doc.add_paragraph()
    perf_para.add_run(f"Profit Factor: ").font.bold = True
    pf_run = perf_para.add_run(f"{metrics.get('profit_factor', 0):.2f}")
    pf_run.font.bold = True
    pf_run.font.color.rgb = RGBColor(0, 176, 80) if metrics.get('profit_factor', 0) >= targets['profit_factor'] else RGBColor(255, 0, 0)

    # Win Rate
    add_styled_heading(doc, '3.2 Taux de Réussite', level=2)

    win_total = metrics.get('winning_trades', 0)
    loss_total = metrics.get('losing_trades', 0)
    total = metrics.get('total_trades', 0)

    doc.add_paragraph(f"Trades gagnants: {int(win_total)} / {int(total)}", style='List Bullet')
    doc.add_paragraph(f"Trades perdants: {int(loss_total)} / {int(total)}", style='List Bullet')

    wr_para = doc.add_paragraph()
    wr_para.add_run(f"Win Rate: ").font.bold = True
    wr_run = wr_para.add_run(f"{metrics.get('win_rate', 0):.2f}%")
    wr_run.font.bold = True
    wr_run.font.color.rgb = RGBColor(0, 176, 80) if metrics.get('win_rate', 0) >= targets['win_rate'] else RGBColor(255, 0, 0)

    # Drawdown
    add_styled_heading(doc, '3.3 Gestion du Risque', level=2)

    dd_para = doc.add_paragraph()
    dd_para.add_run(f"Max Drawdown: ").font.bold = True
    dd_run = dd_para.add_run(f"{metrics.get('max_dd_pct', 0):.2f}%")
    dd_run.font.bold = True
    dd_run.font.color.rgb = RGBColor(0, 176, 80) if metrics.get('max_dd_pct', 0) <= targets['max_dd'] else RGBColor(255, 0, 0)

    dd_para = doc.add_paragraph()
    dd_para.add_run(f"Max Drawdown ($): ").font.bold = True
    dd_para.add_run(f"${metrics.get('max_dd', 0):,.2f}")

    dd_para = doc.add_paragraph()
    dd_para.add_run(f"Recovery Factor: ").font.bold = True
    rf_run = dd_para.add_run(f"{metrics.get('recovery_factor', 0):.2f}")
    rf_run.font.bold = True
    rf_run.font.color.rgb = RGBColor(0, 176, 80) if metrics.get('recovery_factor', 0) >= targets['recovery_factor'] else RGBColor(255, 0, 0)

    doc.add_paragraph()

    # Conclusion
    add_styled_heading(doc, '4. CONCLUSION', level=1, color=(68, 114, 196))

    # Calcul statut global
    pf_ok = metrics.get('profit_factor', 0) >= targets['profit_factor']
    wr_ok = metrics.get('win_rate', 0) >= targets['win_rate']
    dd_ok = metrics.get('max_dd_pct', 0) <= targets['max_dd']
    rf_ok = metrics.get('recovery_factor', 0) >= targets['recovery_factor']

    passed_count = sum([pf_ok, wr_ok, dd_ok, rf_ok])

    if passed_count >= 3:
        conclusion = "✅ BACKTEST VALIDÉ"
        conclusion_color = RGBColor(0, 176, 80)
        recommendation = f"Les paramètres {regime} sont validés. Procéder à la phase suivante (Walk-Forward Analysis ou tests démo)."
    elif passed_count >= 2:
        conclusion = "⚠️ BACKTEST PARTIEL"
        conclusion_color = RGBColor(255, 192, 0)
        recommendation = f"Certains objectifs {regime} non atteints. Ajuster paramètres et re-tester avant WFA."
    else:
        conclusion = "❌ BACKTEST ÉCHOUÉ"
        conclusion_color = RGBColor(255, 0, 0)
        recommendation = f"Objectifs {regime} non atteints. Revoir stratégie et paramètres avant de poursuivre."

    conclusion_para = doc.add_paragraph()
    conclusion_run = conclusion_para.add_run(conclusion)
    conclusion_run.font.size = Pt(16)
    conclusion_run.font.bold = True
    conclusion_run.font.color.rgb = conclusion_color

    doc.add_paragraph()
    doc.add_paragraph(recommendation)

    doc.add_paragraph()

    # Prochaines étapes
    add_styled_heading(doc, '5. PROCHAINES ÉTAPES', level=1, color=(68, 114, 196))

    if passed_count >= 3:
        doc.add_paragraph("1. Backtest autres régimes (BULL/BEAR/TRANSITION)", style='List Number')
        doc.add_paragraph("2. Si tous régimes validés → Lancer Walk-Forward Analysis (26 itérations)", style='List Number')
        doc.add_paragraph("3. Validation OOS sur 6 mois par itération", style='List Number')
        doc.add_paragraph("4. Tests démo avec adaptation automatique", style='List Number')
        doc.add_paragraph("5. Production progressive", style='List Number')
    else:
        doc.add_paragraph("1. Identifier métriques défaillantes", style='List Number')
        doc.add_paragraph("2. Ajuster paramètres .set concernés", style='List Number')
        doc.add_paragraph("3. Re-lancer backtest sur même période", style='List Number')
        doc.add_paragraph("4. Analyser avec: python Python/analyze_goldsmc_backtest.py", style='List Number')
        doc.add_paragraph("5. Validation avant phase suivante", style='List Number')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f'TradBOT — GoldSMC v5 Optimization System\nGénéré le {datetime.now().strftime("%Y-%m-%d à %H:%M:%S")}')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(112, 112, 112)

    # Sauvegarder
    if output_path is None:
        excel_path_obj = Path(excel_path)
        output_dir = Path("D:/Dev/TradBOT/Backtest_report/Word_Reports")
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = output_dir / f"Backtest_{regime}_{timestamp}.docx"

    doc.save(output_path)

    print(f"\n✅ Rapport Word généré: {output_path}")
    print(f"   Régime: {regime}")
    print(f"   Statut: {conclusion}")
    print(f"   Objectifs atteints: {passed_count}/4")

    return str(output_path)


def main():
    if len(sys.argv) < 2:
        print("Usage: python Python/generate_backtest_word_report.py \"chemin/backtest.xlsx\"")
        sys.exit(1)

    excel_path = sys.argv[1]

    if not Path(excel_path).exists():
        print(f"❌ Fichier introuvable: {excel_path}")
        sys.exit(1)

    print("=" * 80)
    print("  GÉNÉRATION RAPPORT WORD BACKTEST")
    print("=" * 80)
    print()

    output_path = generate_word_report(excel_path)

    if output_path:
        print("\n" + "=" * 80)
        print("  TERMINÉ")
        print("=" * 80)
        print(f"\n📂 Rapport: {output_path}")
    else:
        print("\n❌ Échec génération rapport")
        sys.exit(1)


if __name__ == "__main__":
    main()
