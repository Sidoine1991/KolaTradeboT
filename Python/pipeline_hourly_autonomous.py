#!/usr/bin/env python3
"""
Pipeline TradBOT — Hourly Autonomous Execution
Toutes les heures : Scan → Top-5 → TradingAgents → Ordre → Suivi

Fallback: TradingView MCP → Deriv WebSocket → Local JSON
"""
import json
import asyncio
import logging
import requests
import sys
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Any, Optional, Tuple

# Import loss cooldown tracker
try:
    from loss_cooldown_tracker import get_cooldown_tracker, check_symbol_cooldown
    COOLDOWN_TRACKER_AVAILABLE = True
except ImportError:
    COOLDOWN_TRACKER_AVAILABLE = False

# Import spike anticipation
try:
    from spike_anticipation import SpikeAnticipator
    SPIKE_ANTICIPATION_AVAILABLE = True
except ImportError:
    SPIKE_ANTICIPATION_AVAILABLE = False  # Tuple is needed for return type in get_ia_status_v2

# Import docx for Word report generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log_warning = lambda msg: print(f"⚠️  {msg}")

# Fix encoding for Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler("logs/pipeline_hourly.log", encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
log = logging.getLogger(__name__)


def _check_mtf_gate(symbol: str, analysis: dict, action: str) -> tuple:
    """Gate MTF : vérifie que H4+H1+M15 confirment la direction du signal.

    Règles :
    - BUY  valide  : H4==BULL ou (H1==BULL et M15==BULL)
    - SELL valide  : H4==BEAR ou (H1==BEAR et M15==BEAR)
    - Rejet absolu : H4+H1 opposés au signal
    - Cohérence MTF: ≥ 4/6 TF dans la direction du signal requis

    Retourne (ok: bool, raison: str)
    Si aucune donnée TF disponible → laisse passer (pas de données = pas de blocage).
    """
    tfs = {
        "m1":  analysis.get("tf_m1_dir",  "NEUT"),
        "m5":  analysis.get("tf_m5_dir",  "NEUT"),
        "m15": analysis.get("tf_m15_dir", "NEUT"),
        "h1":  analysis.get("tf_h1_dir",  "NEUT"),
        "h4":  analysis.get("tf_h4_dir",  "NEUT"),
        "d1":  analysis.get("tf_d1_dir",  "NEUT"),
    }
    # Si toutes les données TF sont absentes → pas de blocage
    if all(d == "NEUT" for d in tfs.values()):
        return True, ""

    h4  = tfs["h4"]
    h1  = tfs["h1"]
    m15 = tfs["m15"]
    side = "BULL" if action == "BUY" else "BEAR"
    opposite = "BEAR" if action == "BUY" else "BULL"

    # Rejet absolu : H4 ET H1 tous les deux opposés au signal
    if h4 == opposite and h1 == opposite:
        return False, f"MTF rejet absolu — H4={h4} H1={h1} tous deux contre {action}"

    # BUY valide : H4 BULL OU (H1 BULL et M15 BULL)
    # SELL valide : H4 BEAR OU (H1 BEAR et M15 BEAR)
    structure_ok = (h4 == side) or (h1 == side and m15 == side)
    if not structure_ok:
        return False, f"MTF structure insuffisante — H4={h4} H1={h1} M15={m15} pour {action}"

    # Cohérence MTF : nombre de TF dans la direction
    count_side = sum(1 for d in tfs.values() if d == side)
    if count_side < 4:
        return False, f"MTF cohérence {count_side}/6 TF {side} < 4 requis pour {action}"

    return True, ""

AI_SERVER_URL = "http://127.0.0.1:8000"
GOM_SIGNAL_FILE = Path("data/gom_signal.json")

def get_ia_status_v2(symbol: str, gom_data: dict) -> Tuple[float, dict]:
    """
    Calcule IA Status v2 directement depuis les données TF du GOM verdict.

    Le GOM retourne directement tf_m1_dir, tf_m5_dir, etc. — on l'utilise pour
    calculer confiance multi-TF avec M5 prioritaire (sans appel /decision).

    Retourne: (ia_confidence_pct: float, ia_details: dict)
    """
    # Recueillir les directions TF depuis le GOM verdict
    tf_dirs = {
        "M1": gom_data.get("tf_m1_dir", "NEUT").upper(),
        "M5": gom_data.get("tf_m5_dir", "NEUT").upper(),
        "M15": gom_data.get("tf_m15_dir", "NEUT").upper(),
        "M30": gom_data.get("tf_m30_dir", "NEUT").upper(),
        "H1": gom_data.get("tf_h1_dir", "NEUT").upper(),
        "H4": gom_data.get("tf_h4_dir", "NEUT").upper(),
        "D1": gom_data.get("tf_d1_dir", "NEUT").upper(),
        "W1": gom_data.get("tf_w1_dir", "NEUT").upper(),
    }

    # Convertir direction text en votes
    tf_votes = {}
    count_buy = 0
    count_sell = 0
    count_neutral = 0

    for tf, direction in tf_dirs.items():
        if direction == "BULL":
            tf_votes[tf] = 1
            count_buy += 1
        elif direction == "BEAR":
            tf_votes[tf] = -1
            count_sell += 1
        else:
            tf_votes[tf] = 0
            count_neutral += 1

    # Poids de chaque TF (M5 = référence principale)
    tf_weights = {
        "M5": 0.30, "M15": 0.25, "M1": 0.20, "H1": 0.15,
        "M30": 0.05, "H4": 0.03, "D1": 0.02, "W1": 0.00,
    }

    # Calculer score pondéré
    score_buy = 0.0
    score_sell = 0.0
    for tf, vote in tf_votes.items():
        weight = tf_weights.get(tf, 0.0)
        if vote > 0:
            score_buy += weight
        elif vote < 0:
            score_sell += weight

    # Décision finale basée sur M5 + scores pondérés
    m5_vote = tf_votes.get("M5", 0)
    action = "HOLD"
    confidence = 0.5

    if m5_vote > 0:  # M5 BULL
        if score_buy >= score_sell:
            action = "BUY"
            confidence = min(0.95, 0.60 + (count_buy / 8.0) * 0.35)
        else:
            action = "HOLD"
            confidence = 0.45
    elif m5_vote < 0:  # M5 BEAR
        if score_sell > score_buy:
            action = "SELL"
            confidence = min(0.95, 0.60 + (count_sell / 8.0) * 0.35)
        else:
            action = "HOLD"
            confidence = 0.45
    else:  # M5 NEUTRAL
        if score_buy > 0.55:
            action = "BUY"
            confidence = min(0.75, 0.50 + (count_buy / 8.0) * 0.25)
        elif score_sell > 0.55:
            action = "SELL"
            confidence = min(0.75, 0.50 + (count_sell / 8.0) * 0.25)
        else:
            action = "HOLD"
            confidence = 0.40 + (count_neutral / 8.0) * 0.10

    # Calcul confiance finale en %
    confidence_pct = round(confidence * 100.0, 1)

    # Alignment score: % des TF alignés avec action
    if action == "BUY":
        aligned = count_buy
    elif action == "SELL":
        aligned = count_sell
    else:
        aligned = max(count_buy, count_sell, count_neutral)

    total_votes = count_buy + count_sell + count_neutral
    alignment_score = round((aligned / max(total_votes, 1)) * 100.0, 1) if total_votes > 0 else 0.0

    log.debug(f"  ✅ IA v2 pour {symbol}: action={action} {confidence_pct}% (M5={m5_vote}, votes B/S/N: {count_buy}/{count_sell}/{count_neutral})")

    return confidence_pct, {
        "action": action,
        "confidence_pct": confidence_pct,
        "alignment_score": alignment_score,
        "m5_direction": m5_vote,
        "counts": {"buy": count_buy, "sell": count_sell, "neutral": count_neutral},
        "tf_directions": tf_dirs,
        "source": "gom_tfdirs_local_calc"
    }


class PipelineHourly:
    """Pipeline exécuteur horaire avec suivi complet."""

    def __init__(self):
        self.mt5_attached = {}  # {symbol: bool}
        self.top5_results = []
        self.orders_placed = []
        self.errors = []
        self._gom_cache: Dict[str, Any] = {}
        self._recent_orders: Dict[str, float] = {}  # {symbol: timestamp} — dedup within 5 sec

        # Initialize spike anticipator if available
        self.spike_anticipator = None
        if SPIKE_ANTICIPATION_AVAILABLE:
            self.spike_anticipator = SpikeAnticipator(anticipation_pips=5.0)
            log.info("[INIT] Spike anticipation enabled")

    @staticmethod
    def _flatten_directions(v: dict) -> dict:
        """Aplatit le champ imbriqué 'directions' en clés tf_*_dir individuelles.

        Supporte les deux formats:
          - {"directions": {"M1": "BUY", "M5": "BUY", ...}}  (format gom_signal.json)
          - {"tf_m1_dir": "BULL", ...}                        (format /gom-verdict)

        Convertit BUY->BULL, SELL->BEAR, NEUTRAL->NEUT pour uniformité pipeline.
        """
        dirs = v.get("directions")
        if not isinstance(dirs, dict):
            return v

        _MAP = {"BUY": "BULL", "SELL": "BEAR", "NEUTRAL": "NEUT", "NEUT": "NEUT",
                "BULL": "BULL", "BEAR": "BEAR"}
        tf_key_map = {
            "M1": "tf_m1_dir", "M5": "tf_m5_dir", "M15": "tf_m15_dir",
            "M30": "tf_m30_dir", "H1": "tf_h1_dir", "H4": "tf_h4_dir",
            "D1": "tf_d1_dir", "W1": "tf_w1_dir",
        }
        for tf_label, flat_key in tf_key_map.items():
            raw_val = dirs.get(tf_label, "NEUT")
            v[flat_key] = _MAP.get(str(raw_val).upper(), "NEUT")

        # Calculer coherence_pct depuis les directions si absent ou 0
        if not v.get("coherence_pct"):
            verdict_str = str(v.get("verdict", "")).upper()
            is_buy = "BUY" in verdict_str
            target = "BULL" if is_buy else "BEAR"
            aligned = sum(1 for d in dirs.values() if _MAP.get(str(d).upper(), "NEUT") == target)
            total = len(dirs)
            # Confiance = aligned/total * 100 (si >= 4/6 = 67%+, si 5/6 = 83%, 6/6 = 100%)
            v["coherence_pct"] = round((aligned / max(total, 1)) * 100.0, 1)

        return v

    def load_gom_data(self) -> Dict[str, Dict[str, Any]]:
        """Charge les verdicts GOM depuis le serveur live (fallback fichier JSON)."""
        # Priorité 1: endpoint live /gom-verdicts (tous les symboles)
        try:
            r = requests.get(f"{AI_SERVER_URL}/gom-verdicts", timeout=5)
            if r.status_code == 200:
                data = r.json()
                verdicts = data.get("verdicts", data) if isinstance(data, dict) else data
                if isinstance(verdicts, list) and verdicts:
                    log.info(f"  [GOM] {len(verdicts)} verdicts chargés depuis /gom-verdicts (LIVE)")
                    indexed = {v["symbol"]: self._flatten_directions(v) for v in verdicts if "symbol" in v}
                    # Enrichir avec kola/BB/ATR depuis /gom-kola-dashboard (niveaux réels)
                    enriched = 0
                    for sym, v in indexed.items():
                        if float(v.get("kola_buy") or 0) > 0:
                            continue  # déjà enrichi
                        try:
                            import urllib.parse
                            enc = urllib.parse.quote(sym)
                            dr = requests.get(
                                f"{AI_SERVER_URL}/gom-kola-dashboard?symbol={enc}&source=local",
                                timeout=3
                            )
                            if dr.status_code == 200:
                                d = dr.json()
                                for field in ("kola_buy", "kola_sell", "bb_up", "bb_dn", "atr", "atr14"):
                                    val = d.get(field)
                                    if val:
                                        v[field] = float(val)
                                enriched += 1
                        except Exception:
                            pass
                    if enriched:
                        log.info(f"  [GOM] {enriched} symboles enrichis kola/BB/ATR depuis dashboard")
                    return indexed
        except Exception:
            pass

        # Priorité 2: fichier JSON local (peut être stale)
        try:
            if GOM_SIGNAL_FILE.exists():
                raw = json.loads(GOM_SIGNAL_FILE.read_text(encoding="utf-8"))

                # Format A: {"verdicts": [...], "timestamp": ...} (gom_sync export)
                if isinstance(raw, dict) and "verdicts" in raw and isinstance(raw["verdicts"], list):
                    indexed = {v["symbol"]: self._flatten_directions(v)
                               for v in raw["verdicts"] if "symbol" in v}
                    log.info(f"  [GOM] {len(indexed)} verdicts depuis fichier local (format verdicts[])")
                    return indexed

                # Format B: [v1, v2, ...] (ancien format liste)
                if isinstance(raw, list):
                    indexed = {v["symbol"]: self._flatten_directions(v)
                               for v in raw if "symbol" in v}
                    log.warning("  [GOM] Données depuis fichier local (format liste, peut être stale)")
                    return indexed

                # Format C: {"SYMBOL": {...}, ...} (dict par symbole, écrit par _persist_gom_signal_file)
                if isinstance(raw, dict):
                    indexed = {k: self._flatten_directions(v) for k, v in raw.items()
                               if isinstance(v, dict) and "verdict" in v}
                    if indexed:
                        log.info(f"  [GOM] {len(indexed)} verdicts depuis fichier local (format dict)")
                        return indexed
                    return raw
        except Exception as e:
            log.error(f"  [GOM] Erreur lecture fichier: {e}")
        return {}

    def check_mt5_attachment(self, symbol: str) -> bool:
        """Vérifie si le symbole est attaché à SMC_Universal."""
        try:
            # Appel ai_server pour vérifier l'état du symbole
            response = requests.get(
                f"{AI_SERVER_URL}/chart-status",
                params={"symbol": symbol},
                timeout=5
            )
            return response.status_code == 200 and response.json().get("attached", False)
        except Exception as e:
            log.warning(f"⚠️  Impossible vérifier {symbol}: {e}")
            return False

    def send_whatsapp_alert(self, symbol: str, message: str) -> bool:
        """Envoie alerte WhatsApp."""
        try:
            payload = {
                "message": f"⚠️ **PIPELINE ALERT**\n{symbol}\n{message}",
                "to_number": "2290196911346"
            }
            response = requests.post(
                f"{AI_SERVER_URL}/whatsapp/send",
                json=payload,
                timeout=10
            )
            return response.status_code in [200, 201]
        except Exception as e:
            log.warning(f"⚠️  WhatsApp error: {e}")
            return False

    def get_top5_signals(self) -> List[Tuple[str, str, float]]:
        """Phase 1: Scan et sélectionne Top-5."""
        log.info("📊 Phase 1 — Scan des symboles")
        gom_data = self.load_gom_data()
        self._gom_cache = gom_data  # cache pour le fallback Phase 2

        signals = []
        _seen_norm: dict = {}  # norm_key → index in signals (dédup Boom 1000/BOOM 1000)
        for sym, v in gom_data.items():
            verdict = v.get("verdict", "WAIT")
            verdict_num = v.get("verdict_num", 0)
            if verdict_num == 0:
                continue
            score = v.get("verdict_gap", v.get("coherence_pct", 2.0))
            # Normalise pour dédupliquer ex: "Boom 1000 Index" == "BOOM 1000 INDEX"
            sym_norm = sym.upper().replace(" ", "").replace("INDEX", "").replace("DERIV:", "")
            if sym_norm in _seen_norm:
                idx = _seen_norm[sym_norm]
                if score > signals[idx][2]:
                    log.info(f"  [dedup] {sym} remplace {signals[idx][0]} (score {score:.1f} > {signals[idx][2]:.1f})")
                    signals[idx] = (sym, verdict, score)
                continue
            _seen_norm[sym_norm] = len(signals)
            signals.append((sym, verdict, score))
            log.info(f"  ✅ {sym:25s} | {verdict:20s} | Score: {score:.1f}")

        # Tri par score et limite à 5
        top5 = sorted(signals, key=lambda x: x[2], reverse=True)[:5]
        log.info(f"📋 Top-5 sélectionnés: {len(top5)}")
        return top5

    def analyze_with_trading_agents(self, symbol: str, verdict: str) -> Optional[Dict[str, Any]]:
        """Phase 2: Analyse avec TradingAgents via HTTP endpoint. Fallback GOM si indisponible."""
        log.info(f"🤖 Phase 2 — Analyse TradingAgents: {symbol}")

        try:
            response = requests.get(
                f"{AI_SERVER_URL}/ta-analysis",
                params={"symbol": symbol, "date_str": datetime.now().strftime("%Y-%m-%d")},
                timeout=15
            )
            if response.status_code == 200:
                data = response.json()
                if data.get("success"):
                    log.info(f"  ✅ TradingAgents OK: {data.get('opinion')}")
                    # Handle report generation if job_id is available
                    job_id = data.get("job_id")
                    if job_id:
                        self._queue_tradingagents_report(job_id, symbol)
                    # Enrichir avec ia_status_v2 depuis GOM cache si absent (pour gate HOLD)
                    if "ia_status_v2" not in data:
                        g = getattr(self, "_gom_cache", {}).get(symbol, {})
                        if g:
                            _, ia_details = get_ia_status_v2(symbol, g)
                            data["ia_status_v2"] = ia_details
                    return data
                log.warning(f"  ⚠️  TradingAgents not available: {data.get('error')} — fallback GOM")
            else:
                log.warning(f"  ⚠️  HTTP {response.status_code} — fallback GOM")

        except requests.exceptions.Timeout:
            log.warning(f"  ⏱️  TradingAgents timeout (15s) — fallback GOM")
        except Exception as e:
            log.warning(f"  ⚠️  TradingAgents error: {e} — fallback GOM")

        # Fallback: construire l'analyse depuis le verdict GOM live
        return self._analyze_from_gom(symbol, verdict)

    def _analyze_from_gom(self, symbol: str, verdict: str) -> Optional[Dict[str, Any]]:
        """Fallback: construit l'analyse depuis le cache GOM chargé en Phase 1 (pas de réseau)."""
        try:
            g = getattr(self, "_gom_cache", {}).get(symbol)
            if not g:
                log.error(f"  ❌ GOM cache vide pour {symbol}")
                return None

            verdict_num = g.get("verdict_num", 0)

            if verdict_num == 0:
                log.warning(f"  ⚠️  GOM verdict=WAIT pour {symbol} — skip")
                return None

            direction = "BUY" if verdict_num > 0 else "SELL"
            entry = float(g.get("entry", 0.0) or g.get("price", 0.0))
            if entry <= 0:
                log.warning(f"  ⚠️  GOM entry=0 pour {symbol} — skip")
                return None

            atr_raw = float(g.get("atr", 0.0) or g.get("atr14", 0.0))
            kola_buy  = float(g.get("kola_buy")  or 0.0)
            kola_sell = float(g.get("kola_sell") or 0.0)
            bb_up = float(g.get("bb_up") or 0.0)
            bb_dn = float(g.get("bb_dn") or 0.0)
            is_synthetic = "CRASH" in symbol.upper() or "BOOM" in symbol.upper() or "JUMP" in symbol.upper()

            # Multiplicateurs ATR selon stratégie GoldSMC v2
            atr_sl_mult = 1.5 if is_synthetic else 2.0
            atr_tp_rr   = 2.0 if is_synthetic else 1.5
            min_sl_dist = atr_raw * atr_sl_mult if atr_raw > 0 else entry * 0.003

            if direction == "BUY":
                # Entry : utiliser kola_buy (support réel) si en-dessous du prix
                if kola_buy > 0 and kola_buy < entry:
                    entry = kola_buy
                # SL : sous bb_dn ou sous entry d'un ATR×mult
                if bb_dn > 0 and abs(entry - bb_dn) >= min_sl_dist:
                    sl = bb_dn - atr_raw * 0.5 if atr_raw > 0 else bb_dn * 0.998
                else:
                    sl = entry - min_sl_dist
                # TP : kola_sell (résistance) ou bb_up
                if kola_sell > entry:
                    tp = kola_sell
                elif bb_up > entry:
                    tp = bb_up
                else:
                    tp = entry + abs(entry - sl) * atr_tp_rr
            else:
                # Entry : utiliser kola_sell (résistance réelle) si au-dessus du prix
                if kola_sell > 0 and kola_sell > entry:
                    entry = kola_sell
                # SL : au-dessus bb_up ou au-dessus entry d'un ATR×mult
                if bb_up > 0 and abs(bb_up - entry) >= min_sl_dist:
                    sl = bb_up + atr_raw * 0.5 if atr_raw > 0 else bb_up * 1.002
                else:
                    sl = entry + min_sl_dist
                # TP : kola_buy (support) ou bb_dn
                if kola_buy > 0 and kola_buy < entry:
                    tp = kola_buy
                elif bb_dn > 0 and bb_dn < entry:
                    tp = bb_dn
                else:
                    tp = entry - abs(sl - entry) * atr_tp_rr

            # Garantir RR ≥ 1.0 (TP doit couvrir au moins la distance SL)
            sl_dist = abs(entry - sl)
            tp_dist = abs(entry - tp)
            if sl_dist > 0 and tp_dist < sl_dist:
                tp = entry + sl_dist * atr_tp_rr if direction == "BUY" else entry - sl_dist * atr_tp_rr
                log.info(f"  ℹ️  RR corrigé (tp_dist={tp_dist:.2f} < sl_dist={sl_dist:.2f}) → TP={tp:.5f}")

            # ENRICHIR CONFIANCE: IA Status v2 (multi-TF avec M5 prioritaire)
            ia_conf_pct, ia_details = get_ia_status_v2(symbol, g)

            analysis = {
                "success": True,
                "symbol": symbol,
                "opinion": verdict,
                "confidence": ia_conf_pct / 100.0,
                "coherence_pct": ia_conf_pct,  # Remplace GOM coherence par IA v2 confiance
                "ia_status_v2": ia_details,  # Détails complets IA v2
                "entry": round(entry, 5),
                "sl": round(sl, 5),
                "tp": round(tp, 5),
                "lot": 0.01,
                "source": "gom_fallback",
                # Propager les directions TF pour la gate MTF
                "tf_m1_dir":  str(g.get("tf_m1_dir")  or "NEUT"),
                "tf_m5_dir":  str(g.get("tf_m5_dir")  or "NEUT"),
                "tf_m15_dir": str(g.get("tf_m15_dir") or "NEUT"),
                "tf_h1_dir":  str(g.get("tf_h1_dir")  or "NEUT"),
                "tf_h4_dir":  str(g.get("tf_h4_dir")  or "NEUT"),
                "tf_d1_dir":  str(g.get("tf_d1_dir")  or "NEUT"),
                "rsi":        float(g.get("rsi") or g.get("rsi14") or 50),
            }
            log.info(f"  ✅ GOM fallback: {direction} Entry={entry:.2f} SL={sl:.2f} TP={tp:.2f}")
            return analysis

        except Exception as e:
            log.error(f"  ❌ GOM fallback error: {e}")
            return None

    def place_order_on_mt5(self, symbol: str, analysis: Dict[str, Any]) -> bool:
        """Phase 3: Place l'ordre via pending-order endpoint."""
        log.info(f"📈 Phase 3 — Place l'ordre: {symbol}")

        try:
            # GATE WIN-STREAK : pause 1h après 3 gains consécutifs sans perte
            try:
                pause_r = requests.get(f"{AI_SERVER_URL}/trading-pause", timeout=5)
                if pause_r.status_code == 200:
                    pause_data = pause_r.json()
                    if pause_data.get("active"):
                        remaining_min = pause_data.get("remaining_sec", 0) // 60
                        log.warning(f"  🏆 Pause win-streak active — trading suspendu encore {remaining_min}min")
                        self.errors.append((symbol, f"WIN_STREAK_PAUSE_{remaining_min}min"))
                        return False
            except Exception:
                pass  # AI server indisponible → continuer sans gate

            # PRÉ-CHECK: Vérifier si un ordre existe déjà pour ce symbole
            try:
                check = requests.get(
                    f"{AI_SERVER_URL}/pending-order",
                    params={"symbol": symbol},
                    timeout=5
                )
                if check.status_code == 200:
                    existing = check.json()
                    if existing and existing.get("ok") and existing.get("status") in ("ready", "executing"):
                        log.warning(f"  🚫 {symbol}: ORDRE EXISTANT (status={existing.get('status')}) — SKIP placement")
                        self.errors.append((symbol, "ORDER_ALREADY_EXISTS"))
                        return False
            except:
                pass  # Continue si vérification échoue

            # GATE COOLDOWN : Check if symbol is in cooldown after 2 consecutive losses
            if COOLDOWN_TRACKER_AVAILABLE:
                if check_symbol_cooldown(symbol):
                    tracker = get_cooldown_tracker()
                    info = tracker.get_cooldown_info(symbol)
                    remaining_min = info["remaining_minutes"]
                    log.warning(f"  🔴 {symbol}: IN COOLDOWN — {remaining_min}min remaining (2 losses detected)")
                    self.errors.append((symbol, f"COOLDOWN_{remaining_min}min"))
                    return False

            # Build payload matching /pending-order endpoint
            action = "BUY" if analysis.get("opinion", "").upper() in ["BUY", "PERFECT BUY", "GOOD BUY"] else "SELL"

            # GARDE ABSOLUE Boom/Crash — vérifiée en premier avant toute autre gate
            sym_up = symbol.upper()
            if "BOOM" in sym_up and action == "SELL":
                log.warning(f"  🚫 {symbol}: SELL interdit sur Boom (RÈGLE ABSOLUE) — skip")
                self.errors.append((symbol, "BOOM_SELL_FORBIDDEN"))
                return False
            if "CRASH" in sym_up and action == "BUY":
                log.warning(f"  🚫 {symbol}: BUY interdit sur Crash (RÈGLE ABSOLUE) — skip")
                self.errors.append((symbol, "CRASH_BUY_FORBIDDEN"))
                return False

            # GATE IA STATUS : coherence_pct doit être >= seuil (80% standard, 70% Boom/Crash)
            # Note: 0 inclus → données absentes = bloqué (anciennement 0 < _ia laissait passer)
            # Synthétiques (Boom/Crash) : seuil réduit à 70% car volatilité propre + directions fiables
            _ia = float(analysis.get("coherence_pct") or analysis.get("confidence", 0) * 100 or 0)
            _is_synthetic = "BOOM" in sym_up or "CRASH" in sym_up or "JUMP" in sym_up
            _ia_threshold = 70.0 if _is_synthetic else 80.0
            if _ia < _ia_threshold:
                log.warning(f"  🚫 {symbol}: IA status {_ia:.0f}% < {_ia_threshold:.0f}% requis — ordre bloqué")
                self.errors.append((symbol, f"IA_STATUS_{_ia:.0f}pct"))
                return False

            # GATE IA ACTION HOLD : si IA v2 dit HOLD, bloquer même si confiance >= 80%
            _ia_action = (analysis.get("ia_status_v2") or {}).get("action", "")
            if _ia_action == "HOLD":
                log.warning(f"  🚫 {symbol}: IA v2 action=HOLD ({_ia:.0f}%) — aucun signal clair, ordre bloqué")
                self.errors.append((symbol, f"IA_HOLD_{_ia:.0f}pct"))
                return False

            # GATE RSI EXTREME : éviter entrées sur RSI overbought/oversold
            _rsi = float(analysis.get("rsi") or analysis.get("rsi14") or 50)
            if action == "BUY" and _rsi > 75:
                log.warning(f"  🚫 {symbol}: RSI overbought {_rsi:.0f} > 75 — BUY risqué, ordre bloqué")
                self.errors.append((symbol, f"RSI_OVERBOUGHT_{_rsi:.0f}"))
                return False
            if action == "SELL" and _rsi < 25:
                log.warning(f"  🚫 {symbol}: RSI oversold {_rsi:.0f} < 25 — SELL risqué, ordre bloqué")
                self.errors.append((symbol, f"RSI_OVERSOLD_{_rsi:.0f}"))
                return False

            # GATE M15 : M15 ne doit pas être opposé à la direction (tendance intermédiaire)
            _m15 = analysis.get("tf_m15_dir", "NEUT")
            if action == "BUY" and _m15 == "BEAR":
                log.warning(f"  🚫 {symbol}: M15=BEAR opposé à BUY — setup non confirmé, ordre bloqué")
                self.errors.append((symbol, "M15_OPPOSED_BUY"))
                return False
            if action == "SELL" and _m15 == "BULL":
                log.warning(f"  🚫 {symbol}: M15=BULL opposé à SELL — setup non confirmé, ordre bloqué")
                self.errors.append((symbol, "M15_OPPOSED_SELL"))
                return False

            # GATE M1 : M1 ne doit pas être opposé à la direction (micro-tendance)
            _m1 = analysis.get("tf_m1_dir", "NEUT")
            if action == "BUY" and _m1 == "BEAR":
                log.warning(f"  🚫 {symbol}: M1=BEAR opposé à BUY — entrée vouée à l'échec, ordre bloqué")
                self.errors.append((symbol, "M1_OPPOSED_BUY"))
                return False
            if action == "SELL" and _m1 == "BULL":
                log.warning(f"  🚫 {symbol}: M1=BULL opposé à SELL — entrée vouée à l'échec, ordre bloqué")
                self.errors.append((symbol, "M1_OPPOSED_SELL"))
                return False

            # GATE POST-SPIKE : attendre 2 bougies M1 de confirmation après un spike
            _bars_since_spike = int(analysis.get("bars_since_spike") or 99)
            if 0 < _bars_since_spike < 2:
                log.warning(f"  🚫 {symbol}: spike trop récent ({_bars_since_spike} bougie M1) — attente confirmation 2 bougies min")
                self.errors.append((symbol, f"POST_SPIKE_CONFIRM_{_bars_since_spike}bar"))
                return False

            # GATE MTF : H4+H1+M15 doivent confirmer la direction
            # Boom/Crash : gate MTF assouplie (unidirectionnel, M5+M15 suffisent)
            if not _is_synthetic:
                _mtf_ok, _mtf_reason = _check_mtf_gate(symbol, analysis, action)
                if not _mtf_ok:
                    log.warning(f"  🚫 {symbol}: Gate MTF — {_mtf_reason}")
                    self.errors.append((symbol, f"MTF_GATE: {_mtf_reason[:60]}"))
                    return False
            else:
                # Synthétiques : vérifier seulement que M5+M15 ne sont pas opposés
                _m5_s = analysis.get("tf_m5_dir", "NEUT")
                _m15_s = analysis.get("tf_m15_dir", "NEUT")
                _opp = "BEAR" if action == "BUY" else "BULL"
                if _m5_s == _opp and _m15_s == _opp:
                    log.warning(f"  🚫 {symbol}: Synth MTF — M5={_m5_s} M15={_m15_s} tous contre {action}")
                    self.errors.append((symbol, f"SYNTH_MTF: M5+M15 contre {action}"))
                    return False
                log.info(f"  [MTF] Synthétique {symbol}: gate assouplie M5={_m5_s} M15={_m15_s} — OK")

            # (garde Boom/Crash déjà appliquée en tête de fonction)

            # Apply spike anticipation if available
            entry_price = float(analysis.get("entry", 0))
            stop_loss = float(analysis.get("sl", 0))
            take_profit = float(analysis.get("tp", 0))
            verdict_str = analysis.get("opinion", "GOOD")
            rsi = float(analysis.get("rsi", 50))

            if self.spike_anticipator:
                anticipated = self.spike_anticipator.format_order_with_anticipation(
                    symbol=symbol,
                    action=action,
                    base_entry=entry_price,
                    base_sl=stop_loss,
                    base_tp=take_profit,
                    verdict_strength=verdict_str,
                    rsi=rsi,
                    volatility_regime=None,
                )

                if anticipated.get("anticipation_applied"):
                    entry_price = anticipated["entry"]
                    stop_loss = anticipated["sl"]
                    take_profit = anticipated["tp"]
                    log.info(
                        f"  [SPIKE] Anticipation: {anticipated['anticipation_distance_pips']:.1f} pips ahead"
                    )

            # DEDUP : Refuser si même ordre posté dans les 5 dernières secondes
            from time import time
            now = time()
            if symbol in self._recent_orders:
                last_order_time = self._recent_orders[symbol]
                if now - last_order_time < 5:  # moins de 5 sec
                    log.warning(f"  🚫 {symbol}: Ordre déjà placé y a {now - last_order_time:.1f}s — SKIP (dedup)")
                    self.errors.append((symbol, f"DEDUP: Ordre dans les 5s"))
                    return False

            payload = {
                "symbol": symbol,
                "action": action,
                "entry_price": entry_price,
                "stop_loss": stop_loss,
                "take_profit": take_profit,
                "lot": analysis.get("lot", 0.01),
                "source": "pipeline_hourly"
            }

            response = requests.post(
                f"{AI_SERVER_URL}/pending-order",
                json=payload,
                timeout=10
            )

            if response.status_code in [200, 201]:
                result = response.json()
                self._recent_orders[symbol] = now  # Enregistrer le timestamp
                log.info(f"  ✅ Ordre placé: {action} {symbol}")
                self.orders_placed.append((symbol, result))
                return True
            else:
                log.warning(f"  ❌ HTTP {response.status_code}: {response.text[:100]}")
                self.errors.append((symbol, f"HTTP_{response.status_code}"))
                return False
        except Exception as e:
            log.error(f"  ❌ Error: {e}")
            self.errors.append((symbol, str(e)))
            return False

    def build_report(self) -> str:
        """Construit le rapport texte et Word."""
        lines = [
            "PIPELINE HOURLY REPORT",
            f"Time: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')} UTC",
            "",
            f"Top-5 Analyzed: {len(self.top5_results)}",
            "Orders Placed: " + str(len(self.orders_placed)),
            "",
        ]

        if self.orders_placed:
            lines.append("✅ SUCCESSES:")
            for sym, result in self.orders_placed:
                lines.append(f"  • {sym} — Ticket: {result.get('order_ticket', '?')}")

        if self.errors:
            lines.append("\n❌ ERRORS:")
            for sym, err in self.errors:
                lines.append(f"  • {sym} — {err}")

        report_text = "\n".join(lines)

        # Generate Word document
        if DOCX_AVAILABLE:
            self._generate_word_report(report_text)

        return report_text

    def _generate_word_report(self, report_text: str) -> None:
        """Génère le rapport Word et l'envoie via PsychoBot."""
        try:
            doc = Document()
            doc.add_heading("PIPELINE HOURLY REPORT", 0)

            timestamp = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')
            doc.add_paragraph(f"Report Time: {timestamp}")
            doc.add_paragraph(f"Top-5 Analyzed: {len(self.top5_results)}")
            doc.add_paragraph(f"Orders Placed: {len(self.orders_placed)}")
            doc.add_paragraph("")

            if self.orders_placed:
                doc.add_heading("Successful Orders", level=1)
                for sym, result in self.orders_placed:
                    entry = result.get('entry', '?')
                    action = result.get('action', '?')
                    doc.add_paragraph(
                        f"✅ {sym} — {action} @ {entry}",
                        style='List Bullet'
                    )

            if self.errors:
                doc.add_heading("Rejected Orders", level=1)
                for sym, err in self.errors:
                    doc.add_paragraph(f"❌ {sym} — {err}", style='List Bullet')

            # Save to temp file
            report_file = Path("logs") / f"pipeline_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_file.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(report_file))

            log.info(f"  📄 Word report saved: {report_file}")

            # Send via PsychoBot /send-file endpoint
            self._send_word_report_via_psychobot(report_file)

        except Exception as e:
            log.warning(f"  ⚠️  Word report generation failed: {e}")

    def _send_word_report_via_psychobot(self, report_file: Path) -> None:
        """Envoie le rapport Word via PsychoBot Render (en arrière-plan si besoin)."""
        try:
            import subprocess
            import time
            # Delegate to send_tradingagents_report.py which has robust retry logic
            cmd = [
                sys.executable,
                "python/send_tradingagents_report.py",
                "--file", str(report_file),
                "--send-file"
            ]
            # Run async in background
            subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            log.info(f"  📤 [Background] Sending pipeline report to WhatsApp")

        except Exception as e:
            log.warning(f"  ⚠️  Failed to queue report send: {e}")

    def _queue_tradingagents_report(self, job_id: str, symbol: str) -> None:
        """Génère et envoie le rapport TradingAgents en arrière-plan."""
        try:
            import subprocess
            cmd = [
                sys.executable,
                "python/tradingagents_report_handler.py",
                "--job-id", job_id,
                "--symbol", symbol
            ]
            # Run async in background
            subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            log.info(f"  📄 [Background] Generating TradingAgents report (job_id={job_id})")
        except Exception as e:
            log.debug(f"  ⚠️  Failed to queue report: {e}")

    async def run_hourly_cycle(self):
        """Exécute un cycle complet."""
        log.info("=" * 70)
        log.info("🚀 PIPELINE HOURLY CYCLE")
        log.info("=" * 70)

        # Phase 1: Top-5
        self.top5_results = self.get_top5_signals()

        if not self.top5_results:
            log.warning("⚠️  Aucun signal valide")
            return

        # Phase 2-3: Pour chaque top-5
        for symbol, verdict, score in self.top5_results:
            log.info(f"\n>>> Traite {symbol}")

            # Analyse TradingAgents
            ta_analysis = self.analyze_with_trading_agents(symbol, verdict)

            if not ta_analysis:
                log.warning(f"  Analyse TA échouée — skip")
                self.errors.append((symbol, "TA_FAILED"))
                continue

            # Place l'ordre
            self.place_order_on_mt5(symbol, ta_analysis)

        # Rapport
        report = self.build_report()
        log.info("\n" + report)

        # Envoyer rapport via WhatsApp
        self.send_whatsapp_alert("PIPELINE", report)

        log.info("=" * 70)
        log.info("✅ Cycle terminé")

async def main():
    """Main loop — Exécute toutes les heures."""
    import argparse

    parser = argparse.ArgumentParser(description="Pipeline Hourly Autonomous")
    parser.add_argument("--once", action="store_true", help="Un seul cycle")
    args = parser.parse_args()

    pipeline = PipelineHourly()

    if args.once:
        await pipeline.run_hourly_cycle()
    else:
        import time
        while True:
            await pipeline.run_hourly_cycle()
            log.info("Prochain cycle dans 5 minutes...")
            time.sleep(300)

if __name__ == "__main__":
    asyncio.run(main())
