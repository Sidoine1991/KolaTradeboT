#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TradingAgents Report Handler
Récupère les rapports markdown depuis le job_id, les convertit en Word, et les envoie par WhatsApp.
"""

import sys
import logging
import requests
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Fix encoding for Windows
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler(LOG_DIR / "tradingagents_report.log", encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
log = logging.getLogger(__name__)

AI_SERVER_URL = "http://127.0.0.1:8000"


def fetch_report_markdown(job_id: str) -> Optional[str]:
    """Télécharge le rapport markdown depuis le job_id."""
    try:
        url = f"{AI_SERVER_URL}/api/jobs/{job_id}/report.md"
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            return response.text
        else:
            log.warning(f"⚠️  Failed to fetch report: HTTP {response.status_code}")
            return None
    except Exception as e:
        log.warning(f"⚠️  Fetch report error: {e}")
        return None


def markdown_to_word(markdown_content: str, output_path: Path) -> bool:
    """Convertit markdown en document Word basique."""
    if not DOCX_AVAILABLE:
        log.warning("⚠️  python-docx not installed, skipping Word conversion")
        return False

    try:
        doc = Document()
        doc.add_heading("TradingAgents Analysis Report", 0)

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
        doc.add_paragraph(f"Generated: {timestamp}")
        doc.add_paragraph("")

        # Parse markdown lines and add to Word
        lines = markdown_content.split('\n')
        current_list = []

        for line in lines:
            line = line.strip()

            # Headers
            if line.startswith('# '):
                if current_list:
                    # Add accumulated list
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                if current_list:
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                if current_list:
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                doc.add_heading(line[4:], level=3)

            # Bold text
            elif line.startswith('**') and line.endswith('**'):
                if current_list:
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                p = doc.add_paragraph(line[2:-2])
                p.runs[0].bold = True

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                current_list.append(line[2:])

            # Empty line
            elif not line:
                if current_list:
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                doc.add_paragraph("")

            # Normal text
            elif line:
                if current_list:
                    for item in current_list:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list = []
                doc.add_paragraph(line)

        # Flush remaining list
        if current_list:
            for item in current_list:
                doc.add_paragraph(item, style='List Bullet')

        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        log.info(f"✅ Word report saved: {output_path}")
        return True

    except Exception as e:
        log.warning(f"⚠️  Word conversion failed: {e}")
        return False


def handle_tradingagents_report(job_id: str, symbol: str) -> Optional[Path]:
    """
    Récupère le rapport TradingAgents, le convertit en Word, et l'envoie.

    Returns: Path au fichier Word généré, ou None si échec
    """
    log.info(f"🔄 Handling TradingAgents report for {symbol} (job_id: {job_id})")

    # Fetch markdown report
    markdown = fetch_report_markdown(job_id)
    if not markdown:
        log.warning(f"⚠️  No markdown report found for job_id: {job_id}")
        return None

    # Generate Word document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_path = LOG_DIR / f"tradingagents_report_{symbol}_{timestamp}.docx"

    if not markdown_to_word(markdown, word_path):
        log.warning(f"⚠️  Failed to convert markdown to Word")
        return None

    # Save markdown for reference
    md_path = LOG_DIR / f"tradingagents_report_{symbol}_{timestamp}.md"
    try:
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        log.info(f"✅ Markdown saved: {md_path}")
    except Exception as e:
        log.warning(f"⚠️  Failed to save markdown: {e}")

    # Send via send_tradingagents_report.py
    try:
        import subprocess
        cmd = [
            sys.executable,
            "python/send_tradingagents_report.py",
            "--file", str(word_path),
            "--send-file"
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0:
            log.info(f"✅ Report sent via WhatsApp")
        else:
            log.warning(f"⚠️  Report send failed: {result.stderr}")
    except Exception as e:
        log.warning(f"⚠️  Failed to send report: {e}")

    return word_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--job-id", required=True, help="TradingAgents job_id")
    parser.add_argument("--symbol", required=True, help="Trading symbol")
    args = parser.parse_args()

    result = handle_tradingagents_report(args.job_id, args.symbol)
    sys.exit(0 if result else 1)
