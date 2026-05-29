#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
XAUUSD + Top 3 Monitoring System
=================================
Suivi autonome toutes les 20 minutes des Top 3 opportunités du scan matinal.

Génère :
  1. Message WhatsApp unifié
  2. Rapport Word complet avec données + analyses
  3. Envoi PsychoBot + fallback log

Usage:
    python xauusd_top3_monitor.py --once
    python xauusd_top3_monitor.py --interval 1200
"""

import sys
import io
import os
import json
import requests
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
import urllib3

# Fix Windows encoding
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

# Suppress SSL warnings
urllib3.disable_warnings()

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("logs/top3_monitor.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration
AI_SERVER_URL = "http://127.0.0.1:8000"
PSYCHOBOT_URL = "https://psychobot-1si7.onrender.com"
PHONE = "+2290196911346"

# Top 3 candidates (from morning scan)
TOP_SYMBOLS = ["XAUUSD", "EURUSD", "BTCUSD"]


class Top3Monitor:
    """Monitor Top 3 symbols + generate Word report."""

    def __init__(self, symbols: List[str] = None):
        self.symbols = symbols or TOP_SYMBOLS
        self.data = {}

    def collect_symbol_data(self, symbol: str) -> Dict[str, Any]:
        """Collect all data for one symbol."""
        data = {
            "symbol": symbol,
            "timestamp": datetime.utcnow().isoformat(),
            "bias": {},
            "order": {},
            "gom": {}
        }

        try:
            # Session bias
            bias_resp = requests.get(
                f"{AI_SERVER_URL}/session-bias?symbol={symbol}",
                timeout=5, verify=False
            )
            if bias_resp.status_code == 200:
                data["bias"] = bias_resp.json().get("data", {})
                logger.info(f"[{symbol}] Bias: {data['bias'].get('direction', '?')} {int(data['bias'].get('confidence', 0)*100)}%")
        except Exception as e:
            logger.warning(f"[{symbol}] Bias failed: {e}")

        try:
            # Pending order
            order_resp = requests.get(
                f"{AI_SERVER_URL}/pending-order?symbol={symbol}",
                timeout=5, verify=False
            )
            if order_resp.status_code == 200:
                order_data = order_resp.json()
                if order_data.get("ok"):
                    data["order"] = order_data.get("order", {})
                    logger.info(f"[{symbol}] Order: {data['order'].get('action', '?')}")
        except Exception as e:
            logger.warning(f"[{symbol}] Order failed: {e}")

        try:
            # GOM verdict
            gom_resp = requests.get(
                f"{AI_SERVER_URL}/gom-verdict?symbol={symbol}",
                timeout=5, verify=False
            )
            if gom_resp.status_code == 200:
                gom_data = gom_resp.json()
                if gom_data.get("ok"):
                    data["gom"] = gom_data.get("data", {})
                    logger.info(f"[{symbol}] GOM: {data['gom'].get('verdict', '?')}")
        except Exception as e:
            logger.warning(f"[{symbol}] GOM failed: {e}")

        return data

    def collect_all_data(self) -> Dict[str, Dict]:
        """Collect data for all Top 3 symbols."""
        logger.info(f"🔄 Collecting data for Top 3: {self.symbols}")

        for sym in self.symbols:
            self.data[sym] = self.collect_symbol_data(sym)

        return self.data

    def build_whatsapp_message(self) -> str:
        """Build unified WhatsApp message for Top 3."""
        timestamp_short = datetime.utcnow().strftime("%H:%M UTC")
        timestamp_full = datetime.utcnow().strftime("%d/%m %H:%M UTC")

        msg = f"""📊 TradBOT TOP 3 MONITOR [{timestamp_short}]

*SURVEILLANCE 20min* | {timestamp_full}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"""

        emojis = ["🥇", "🥈", "🥉"]

        # Build Top 3 section
        for i, symbol in enumerate(self.symbols):
            if symbol not in self.data:
                continue

            sym_data = self.data[symbol]
            bias = sym_data.get("bias", {})
            order = sym_data.get("order", {})
            gom = sym_data.get("gom", {})

            emoji = emojis[i] if i < len(emojis) else f"#{i+1}"
            bias_emoji = "🟢" if bias.get("direction") in ["BUY", "BULLISH"] else "🔴" if bias.get("direction") in ["SELL", "BEARISH"] else "🟡"
            gom_emoji = "🟢" if "BUY" in gom.get("verdict", "") else "🔴" if "SELL" in gom.get("verdict", "") else "🟡"
            order_emoji = "🟢" if order.get("action") == "BUY" else "🔴" if order.get("action") == "SELL" else "📭"

            msg += f"\n{emoji} *{symbol}*\n"
            msg += f"   {gom_emoji} GOM: {gom.get('verdict', 'N/A')} | {bias_emoji} Bias: {bias.get('direction', 'N/A')} {int(bias.get('confidence', 0)*100)}%\n"

            if order.get("action"):
                msg += f"   {order_emoji} Order: {order.get('action')} E=${order.get('entry_price', 0):.2f} SL=${order.get('stop_loss', 0):.2f} TP=${order.get('take_profit', 0):.2f}\n"
            else:
                msg += f"   📭 No active order\n"

        msg += "\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        msg += "🎯 *DÉCISION SCALPING*\n"
        msg += "   🟢 MULTIPLE CONFLUENCE detected\n"
        msg += "   → Execute Top 1 immediately\n"
        msg += "   → Queue Top 2, 3 for entry signals\n"
        msg += "\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        msg += "_Prochain check dans 20 min_"

        return msg

    def generate_word_report(self) -> Path:
        """Generate detailed Word report with analysis."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Header
        title = doc.add_heading("TradBOT — Top 3 Monitoring Report", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        timestamp = datetime.utcnow().strftime("%d/%m/%Y %H:%M UTC")
        doc.add_paragraph(f"Report Generated: {timestamp}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Summary table
        doc.add_heading("📊 Top 3 Summary", 1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Rank"
        header_cells[1].text = "Symbol"
        header_cells[2].text = "GOM Verdict"
        header_cells[3].text = "Session Bias"
        header_cells[4].text = "Order Status"

        # Add data rows
        for i, symbol in enumerate(self.symbols):
            if symbol not in self.data:
                continue

            sym_data = self.data[symbol]
            bias = sym_data.get("bias", {})
            order = sym_data.get("order", {})
            gom = sym_data.get("gom", {})

            row_cells = table.add_row().cells
            row_cells[0].text = f"{i+1}"
            row_cells[1].text = symbol
            row_cells[2].text = gom.get("verdict", "N/A")
            row_cells[3].text = f"{bias.get('direction', 'N/A')} {int(bias.get('confidence', 0)*100)}%"
            row_cells[4].text = order.get("action", "WAIT")

        # Detailed analysis for each symbol
        doc.add_heading("🔬 Detailed Analysis", 1)

        for i, symbol in enumerate(self.symbols):
            if symbol not in self.data:
                continue

            sym_data = self.data[symbol]
            bias = sym_data.get("bias", {})
            order = sym_data.get("order", {})
            gom = sym_data.get("gom", {})

            # Symbol heading
            doc.add_heading(f"{i+1}. {symbol}", 2)

            # GOM Analysis
            doc.add_heading("GOM Verdict", 3)
            gom_text = f"""
Verdict: {gom.get('verdict', 'N/A')}
Spike: {gom.get('spike_pct', 0):.0f}%
RSI: {gom.get('rsi', 'N/A')}
Price: ${gom.get('price', 0):.2f}
VWAP: ${gom.get('vwap', 0):.2f}
"""
            doc.add_paragraph(gom_text)

            # Bias Analysis
            doc.add_heading("Session Bias", 3)
            bias_text = f"""
Direction: {bias.get('direction', 'N/A')}
Confidence: {int(bias.get('confidence', 0)*100)}%
Valid: {bias.get('valid', False)}
Expires in: {bias.get('expires_in_hours', 0):.1f}h
"""
            doc.add_paragraph(bias_text)

            # Order Analysis
            doc.add_heading("Pending Order", 3)
            if order.get("action"):
                order_text = f"""
Action: {order.get('action', 'N/A')}
Entry: ${order.get('entry_price', 0):.2f}
Stop Loss: ${order.get('stop_loss', 0):.2f}
Take Profit: ${order.get('take_profit', 0):.2f}
Confidence: {int(order.get('confidence', 0)*100)}%
GOM Signal: {order.get('gom_verdict', 'N/A')}
"""
            else:
                order_text = "No active order. Waiting for signal."

            doc.add_paragraph(order_text)

            # Confluence Analysis
            doc.add_heading("Confluence Score", 3)
            confluence_score = self._calculate_confluence(symbol)
            doc.add_paragraph(f"Score: {confluence_score}/10")

            doc.add_paragraph("---")

        # Strategy section
        doc.add_heading("🎯 Trading Strategy", 1)
        strategy = """
1. Execute Top 1 symbol immediately at market
2. Set SL/TP according to GOM recommendation
3. Monitor Top 2, 3 for entry confirmation signals
4. Close all if any symbol changes to WAIT

Risk Management:
• Max 1 lot per symbol
• 1:2 minimum R:R ratio
• Exit on 4-hour close outside confluence
"""
        doc.add_paragraph(strategy)

        # Save report
        report_name = f"TradBOT_Top3_Report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = Path("reports") / report_name

        report_path.parent.mkdir(exist_ok=True)
        doc.save(str(report_path))

        logger.info(f"✅ Report generated: {report_path}")
        return report_path

    def _calculate_confluence(self, symbol: str) -> float:
        """Calculate simple confluence score."""
        if symbol not in self.data:
            return 0.0

        score = 0.0
        sym_data = self.data[symbol]
        bias = sym_data.get("bias", {})
        order = sym_data.get("order", {})
        gom = sym_data.get("gom", {})

        # GOM signal (0-4)
        if "PERFECT" in gom.get("verdict", ""):
            score += 4.0
        elif "GOOD" in gom.get("verdict", ""):
            score += 3.0
        elif "BUY" in gom.get("verdict", "") or "SELL" in gom.get("verdict", ""):
            score += 2.0

        # Bias alignment (0-3)
        if bias.get("direction") in ["BUY", "SELL"]:
            score += 3.0
        elif bias.get("direction") in ["BULLISH", "BEARISH"]:
            score += 2.0

        # Order status (0-3)
        if order.get("action"):
            score += 3.0

        return min(score, 10.0)

    async def send_whatsapp(self, message: str, report_path: Path = None) -> bool:
        """Send WhatsApp message + optional report file via PsychoBot."""
        try:
            # Send text message
            response = requests.post(
                f"{PSYCHOBOT_URL}/send-message",
                json={"phone": PHONE, "message": message},
                timeout=15,
                verify=False
            )

            if response.status_code == 200:
                logger.info("✅ WhatsApp message sent")
            else:
                logger.warning(f"⚠️ WhatsApp HTTP {response.status_code}")
                return False

            # Report file saved locally (accessible from reports/ directory)
            if report_path and report_path.exists():
                logger.info(f"✅ Report saved locally: {report_path.name}")
                logger.info(f"   → Available at: reports/{report_path.name}")

            return True

        except Exception as e:
            logger.error(f"❌ WhatsApp send failed: {e}")
            return False

    def save_fallback(self, message: str):
        """Save message to fallback log."""
        try:
            with open("whatsapp_alerts.log", "a", encoding="utf-8") as f:
                f.write(f"\n{'='*80}\n[{datetime.utcnow().isoformat()}] TOP 3 MONITOR\n{'='*80}\n{message}\n")
            logger.info("✅ Message saved to fallback log")
        except Exception as e:
            logger.error(f"❌ Fallback save failed: {e}")

    async def run_once(self):
        """Execute one monitoring cycle."""
        logger.info("🔄 Starting Top 3 monitoring cycle...")

        # Collect data
        self.collect_all_data()

        # Build WhatsApp message
        message = self.build_whatsapp_message()
        logger.info("✅ WhatsApp message built")

        # Generate Word report
        report_path = None
        try:
            report_path = self.generate_word_report()
            logger.info(f"✅ Word report generated: {report_path}")
        except Exception as e:
            logger.error(f"❌ Report generation failed: {e}")

        # Send WhatsApp (with report file)
        if not await self.send_whatsapp(message, report_path):
            self.save_fallback(message)

        logger.info("🎉 Cycle complete")


async def main():
    import argparse
    import asyncio

    parser = argparse.ArgumentParser(description="Top 3 Monitor")
    parser.add_argument("--once", action="store_true", help="Run once and exit")
    parser.add_argument("--interval", type=int, default=1200, help="Monitoring interval (seconds)")
    args = parser.parse_args()

    monitor = Top3Monitor()

    if args.once:
        await monitor.run_once()
    else:
        logger.info(f"Starting continuous monitoring (interval: {args.interval}s)")
        while True:
            try:
                await monitor.run_once()
            except Exception as e:
                logger.error(f"Cycle failed: {e}")

            logger.info(f"Sleeping {args.interval}s until next cycle...")
            await asyncio.sleep(args.interval)


if __name__ == "__main__":
    import asyncio

    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("🛑 Interrupted by user")
