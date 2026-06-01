#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Morning Scanning System V2 — Dynamic Market Scanner
Scans ALL available symbols, filters open markets, analyzes Daily opportunities
"""

import sys
import io
import os
import json
import requests
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional

# Fix Windows encoding
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

class DynamicMorningScanner:
    def __init__(self, ai_server_url="http://127.0.0.1:8000"):
        self.ai_server = ai_server_url
        self.mcp_server = "http://127.0.0.1:3000"  # TradingView MCP
        self.timeout = 30

        # Categories to prioritize
        self.priority_categories = {
            "forex_majors": ["EURUSD", "GBPUSD", "USDJPY", "USDCHF", "AUDUSD", "NZDUSD", "USDCAD"],
            "forex_crosses": ["EURJPY", "GBPJPY", "EURGBP", "EURAUD", "GBPAUD", "AUDNZD"],
            "commodities": ["XAUUSD", "XAGUSD", "USOIL", "UKOIL", "NATGAS"],
            "indices": ["US30", "US500", "NAS100", "UK100", "GER40", "JP225"],
            "crypto": ["BTCUSD", "ETHUSD", "BNBUSD", "SOLUSD", "XRPUSD"],
            "synthetics": ["BOOM500", "CRASH500", "BOOM1000", "CRASH1000", "V75", "V100"]
        }

    def get_all_available_symbols(self) -> List[str]:
        """Get all available symbols from TradingView"""
        all_symbols = []

        try:
            # Try to get symbols from TradingView MCP
            response = requests.post(
                f"{self.mcp_server}/mcp",
                json={
                    "method": "symbol_search",
                    "params": {"query": "", "limit": 500}  # Empty query = all symbols
                },
                timeout=10
            )

            if response.status_code == 200:
                data = response.json()
                symbols = data.get("result", [])
                all_symbols = [s["symbol"] for s in symbols if "symbol" in s]
                print(f"✅ Found {len(all_symbols)} symbols from TradingView")
            else:
                print(f"⚠️ TradingView MCP unavailable, using priority lists")
                # Fallback to priority lists
                for category, syms in self.priority_categories.items():
                    all_symbols.extend(syms)

        except Exception as e:
            print(f"⚠️ Error fetching symbols: {e}, using priority lists")
            # Fallback to priority lists
            for category, syms in self.priority_categories.items():
                all_symbols.extend(syms)

        # Remove duplicates
        all_symbols = list(set(all_symbols))
        return all_symbols

    def check_market_open(self, symbol: str) -> bool:
        """Check if market is open for a symbol"""
        now = datetime.utcnow()
        weekday = now.weekday()  # 0=Monday, 6=Sunday
        hour = now.hour

        # Crypto - always open
        if any(crypto in symbol.upper() for crypto in ["BTC", "ETH", "BNB", "SOL", "XRP"]):
            return True

        # Synthetics - always open
        if any(synth in symbol.upper() for synth in ["BOOM", "CRASH", "V75", "V100", "V10"]):
            return True

        # Forex - Monday 00:00 UTC to Friday 22:00 UTC
        if any(fx in symbol.upper() for fx in ["EUR", "GBP", "USD", "JPY", "AUD", "NZD", "CAD", "CHF"]):
            if weekday < 5:  # Monday to Friday
                return True
            elif weekday == 5 and hour < 22:  # Friday until 22:00
                return True
            elif weekday == 6 and hour >= 22:  # Sunday from 22:00
                return True
            return False

        # Commodities (Gold, Silver, Oil) - Similar to Forex
        if any(comm in symbol.upper() for comm in ["XAU", "XAG", "OIL", "GAS"]):
            if weekday < 5:
                return True
            elif weekday == 5 and hour < 22:
                return True
            elif weekday == 6 and hour >= 22:
                return True
            return False

        # Indices - Market hours vary but generally weekdays
        if any(idx in symbol.upper() for idx in ["US30", "US500", "NAS", "UK100", "GER", "JP225"]):
            if weekday < 5 and 13 <= hour <= 21:  # Simplified: US market hours
                return True
            return False

        # Default: assume open during weekdays
        return weekday < 5

    def analyze_daily_opportunity(self, symbol: str) -> Dict:
        """Analyze Daily timeframe for strong opportunities"""
        try:
            # Get Daily data from TradingView
            response = requests.post(
                f"{self.mcp_server}/mcp",
                json={
                    "method": "tradbot_smc_analysis",
                    "params": {
                        "symbol": symbol,
                        "timeframe": "1D",
                        "lookback": 50  # 50 daily bars
                    }
                },
                timeout=15
            )

            if response.status_code == 200:
                analysis = response.json().get("result", {})

                # Calculate opportunity score based on SMC analysis
                score = 0.0
                signals = []

                # Check for Order Blocks
                if analysis.get("order_blocks"):
                    ob_count = len(analysis["order_blocks"])
                    score += min(ob_count * 1.5, 3.0)  # Max 3 points for OBs
                    signals.append(f"{ob_count} Order Blocks")

                # Check for Fair Value Gaps
                if analysis.get("fvg_zones"):
                    fvg_count = len(analysis["fvg_zones"])
                    score += min(fvg_count * 1.0, 2.0)  # Max 2 points for FVGs
                    signals.append(f"{fvg_count} FVGs")

                # Check for Structure Break
                if analysis.get("structure_break"):
                    score += 2.0
                    signals.append("Structure Break")

                # Check for Liquidity Sweep
                if analysis.get("liquidity_sweep"):
                    score += 1.5
                    signals.append("Liquidity Sweep")

                # Check for trend strength
                trend = analysis.get("trend", {})
                if trend.get("strength", 0) > 0.7:
                    score += 1.5
                    signals.append(f"Strong {trend.get('direction', 'N/A')} trend")

                return {
                    "score": min(score, 10.0),  # Cap at 10
                    "signals": signals,
                    "analysis": analysis
                }
            else:
                # Fallback to AI server analysis
                return self.analyze_via_ai_server(symbol)

        except Exception as e:
            print(f"⚠️ TradingView analysis failed for {symbol}: {e}")
            return self.analyze_via_ai_server(symbol)

    def analyze_via_ai_server(self, symbol: str) -> Dict:
        """Fallback analysis via AI server"""
        try:
            # Get GOM verdict
            gom_resp = requests.get(
                f"{self.ai_server}/gom-verdict?symbol={symbol}",
                timeout=self.timeout
            )
            gom = gom_resp.json() if gom_resp.status_code == 200 else {}

            # Get session bias
            bias_resp = requests.get(
                f"{self.ai_server}/session-bias?symbol={symbol}",
                timeout=self.timeout
            )
            bias = bias_resp.json() if bias_resp.status_code == 200 else {}

            # Calculate score from GOM
            score_buy = float(gom.get("score_buy", 0))
            score_sell = float(gom.get("score_sell", 0))
            score = max(score_buy, score_sell)

            signals = []
            if gom.get("verdict") and gom["verdict"] != "WAIT":
                signals.append(f"GOM {gom['verdict']}")
            if bias.get("direction") and bias["direction"] != "NEUTRAL":
                signals.append(f"Bias {bias['direction']} {int(bias.get('confidence', 0)*100)}%")

            return {
                "score": score,
                "signals": signals,
                "analysis": {"gom": gom, "bias": bias}
            }

        except Exception as e:
            print(f"✗ AI server analysis failed for {symbol}: {e}")
            return {"score": 0.0, "signals": [], "analysis": {}}

    def scan_all_markets(self) -> List[Dict]:
        """Scan all available markets for Daily opportunities"""
        print("🔍 Starting Dynamic Market Scan...")

        # Get all available symbols
        all_symbols = self.get_all_available_symbols()
        print(f"📊 Total symbols to scan: {len(all_symbols)}")

        # Filter open markets
        open_symbols = []
        for sym in all_symbols:
            if self.check_market_open(sym):
                open_symbols.append(sym)

        print(f"🏪 Open markets: {len(open_symbols)}/{len(all_symbols)}")

        # Analyze each open symbol
        opportunities = []
        for i, sym in enumerate(open_symbols):
            print(f"[{i+1}/{len(open_symbols)}] Analyzing {sym}...", end=" ")

            analysis = self.analyze_daily_opportunity(sym)

            if analysis["score"] >= 4.0:  # Minimum threshold
                opportunities.append({
                    "symbol": sym,
                    "score": analysis["score"],
                    "signals": analysis["signals"],
                    "analysis": analysis["analysis"]
                })
                print(f"✅ Score: {analysis['score']:.1f}/10")
            else:
                print(f"⊘ Score: {analysis['score']:.1f}/10 (below threshold)")

        # Sort by score and return top 3
        opportunities.sort(key=lambda x: x["score"], reverse=True)
        return opportunities[:3]

    def build_message(self, top3: List[Dict]) -> str:
        """Build WhatsApp message for top opportunities"""
        now = datetime.utcnow()
        time_str = now.strftime("%H:%M UTC")
        date_str = now.strftime("%d/%m/%Y")

        msg = f"📊 TRADBOT MORNING SCAN V2.0\n"
        msg += f"🌍 GLOBAL MARKET SCANNER\n"
        msg += f"📅 {date_str} | ⏰ {time_str}\n"
        msg += "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"

        msg += f"🎯 *TOP 3 DAILY OPPORTUNITIES*\n\n"

        emojis = ["🥇", "🥈", "🥉"]

        for i, opp in enumerate(top3):
            sym = opp["symbol"]
            score = opp["score"]
            signals = opp["signals"]

            msg += f"{emojis[i]} *{sym}* — Score: {score:.1f}/10\n"

            # Show signals
            if signals:
                for signal in signals[:3]:  # Max 3 signals
                    msg += f"   ✅ {signal}\n"

            # Add specific analysis details if available
            analysis = opp.get("analysis", {})
            if "gom" in analysis:
                gom = analysis["gom"]
                if gom.get("verdict") and gom["verdict"] != "WAIT":
                    msg += f"   📊 GOM: {gom['verdict']}\n"

            msg += "\n"

        msg += "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        msg += "📈 *TRADING STRATEGY*\n"
        msg += f"1️⃣ Focus on {top3[0]['symbol']} first (highest score)\n"
        msg += "2️⃣ Wait for Daily candle confirmation\n"
        msg += "3️⃣ Use H4/H1 for precise entry timing\n"
        msg += "4️⃣ Risk max 1% per trade\n\n"

        msg += "⏰ *NEXT ACTIONS*\n"
        msg += "• Monitor these symbols in Daily TF\n"
        msg += "• Wait for entry signals on lower TF\n"
        msg += "• Next scan: Tomorrow 06:00 UTC\n\n"

        msg += "🤖 TradBOT Dynamic Scanner v2.0"

        return msg

    def send_whatsapp(self, message: str, phone: str = "+2290196911346") -> bool:
        """Send message via PsychoBot"""
        try:
            response = requests.post(
                "https://psychobot-1si7.onrender.com/send-message",
                json={"phone": phone, "message": message},
                timeout=15,
                verify=False
            )
            return response.status_code == 200
        except Exception as e:
            print(f"❌ WhatsApp send failed: {e}")
            return False

    def save_report(self, top3: List[Dict], message: str):
        """Save detailed report as Word document"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('TradBOT Morning Scan Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Date and time
            now = datetime.utcnow()
            doc.add_paragraph(f"Generated: {now.strftime('%d/%m/%Y %H:%M UTC')}")

            # Executive Summary
            doc.add_heading('Executive Summary', 1)
            doc.add_paragraph(f"Top 3 Daily opportunities identified from global market scan.")

            # Detailed Analysis for each symbol
            for i, opp in enumerate(top3):
                doc.add_heading(f"{i+1}. {opp['symbol']}", 1)
                doc.add_paragraph(f"Score: {opp['score']:.1f}/10")

                # Signals
                doc.add_heading('Signals Detected:', 2)
                for signal in opp['signals']:
                    doc.add_paragraph(f"• {signal}", style='List Bullet')

                # Analysis details
                if 'analysis' in opp and opp['analysis']:
                    doc.add_heading('Technical Analysis:', 2)
                    doc.add_paragraph(json.dumps(opp['analysis'], indent=2))

            # Save document
            filename = f"D:\\Dev\\TradBOT\\reports\\morning_scan\\TradBOT_Morning_Scan_{now.strftime('%Y%m%d_%H%M')}.docx"
            os.makedirs(os.path.dirname(filename), exist_ok=True)
            doc.save(filename)

            print(f"✅ Report saved: {filename}")
            return filename

        except ImportError:
            print("⚠️ python-docx not installed, skipping Word report")
            return None
        except Exception as e:
            print(f"❌ Report save failed: {e}")
            return None


def main():
    scanner = DynamicMorningScanner()

    # Scan all markets
    top3 = scanner.scan_all_markets()

    if not top3:
        print("\n⚠️ No opportunities found above threshold")
        return

    print(f"\n✅ Found {len(top3)} top opportunities")

    # Build message
    message = scanner.build_message(top3)
    print("\n" + "="*50)
    print(message)
    print("="*50)

    # Save detailed report
    report_file = scanner.save_report(top3, message)

    # Send WhatsApp
    print("\n📱 Sending WhatsApp notification...")
    if scanner.send_whatsapp(message):
        print("✅ WhatsApp sent successfully")
    else:
        print("⚠️ WhatsApp send failed")

    print("\n✅ Morning scan complete!")


if __name__ == "__main__":
    main()