"""
TradBOT Bridge - TradingAgents CLI -> ai_server -> SMC_Universal.mq5

LANCEMENT : utiliser le wrapper qui active le bon venv
  .\bridge.bat                          # wizard interactif complet (tous les choix TA)
  .\bridge.bat --symbol EURUSD          # mode rapide (skip wizard ticker/date)
  .\bridge.bat --symbol EURUSD --auto   # pas de confirmation avant envoi
  .\bridge.bat --symbol EURUSD --no-pending  # rapport seulement, pas d'ordre MT5

Venv requis : D:\\Dev\\Depot Github\\TradingAgents-main\\.venv\\Scripts\\python.exe
"""

from __future__ import annotations

import ssl
ssl._create_default_https_context = ssl._create_unverified_context

import os
os.environ.setdefault("PYTHONHTTPSVERIFY", "0")
os.environ.setdefault("CURL_CA_BUNDLE", "")
os.environ.setdefault("REQUESTS_CA_BUNDLE", "")

import certifi
import httpx
import urllib3
urllib3.disable_warnings()

# Patch httpx pour désactiver SSL (proxy antivirus Windows)
_orig_client_init = httpx.Client.__init__
def _patched_client_init(self, *a, **kw):
    kw.setdefault("verify", False)
    _orig_client_init(self, *a, **kw)
httpx.Client.__init__ = _patched_client_init

_orig_async_init = httpx.AsyncClient.__init__
def _patched_async_init(self, *a, **kw):
    kw.setdefault("verify", False)
    _orig_async_init(self, *a, **kw)
httpx.AsyncClient.__init__ = _patched_async_init

# Patch curl_cffi (utilisé par yfinance) — doit être fait AVANT import yfinance
try:
    import curl_cffi.requests as _curl_req
    _orig_curl_session_init = _curl_req.Session.__init__
    def _patched_curl_session_init(self, *a, **kw):
        kw.setdefault("verify", False)
        _orig_curl_session_init(self, *a, **kw)
    _curl_req.Session.__init__ = _patched_curl_session_init
except Exception:
    pass

import argparse
import re as _re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# ---------------------------------------------------------------------------
# Chemins & env
# ---------------------------------------------------------------------------

_HERE = Path(__file__).resolve().parent
_TRADBOT_ROOT = _HERE.parent
_TA_REPO = Path(
    os.getenv("AI_TRADINGAGENTS_REPO_PATH",
              r"D:\Dev\Depot Github\TradingAgents-main")
)

# Charge .env TradBOT en premier (prioritaire), puis .env TA sans ecraser
try:
    from dotenv import load_dotenv
    load_dotenv(_TRADBOT_ROOT / ".env", override=True)
    if (_TA_REPO / ".env").exists():
        load_dotenv(_TA_REPO / ".env", override=False)
except ImportError:
    pass  # python-dotenv absent du venv TA ? peu probable

# Si le provider est Bedrock, charger les credentials du profil AWS default
# (meme compte que Claude Code local) via le fichier ~/.aws/credentials.
# Cela evite que les cles hardcodees du .env TA (Legacy) prennent le dessus.
if os.getenv("TRADINGAGENTS_LLM_PROVIDER", "").lower() == "bedrock":
    import configparser as _cp
    _aws_creds_file = Path.home() / ".aws" / "credentials"
    _aws_profile = os.getenv("AWS_PROFILE", "default")
    if _aws_creds_file.exists():
        _cp_parser = _cp.ConfigParser()
        _cp_parser.read(str(_aws_creds_file))
        if _aws_profile in _cp_parser:
            _sec = _cp_parser[_aws_profile]
            _key_id = _sec.get("aws_access_key_id", "")
            _secret = _sec.get("aws_secret_access_key", "")
            if _key_id and _secret:
                os.environ["AWS_ACCESS_KEY_ID"]     = _key_id
                os.environ["AWS_SECRET_ACCESS_KEY"] = _secret
                os.environ["AWS_REGION"] = os.getenv("AWS_REGION", "us-east-1")
                # Supprimer le token Bearer Bedrock pour forcer l'auth IAM SigV4
                os.environ.pop("AWS_BEARER_TOKEN_BEDROCK", None)
                print(f"[bridge] AWS profile '{_aws_profile}' charge: {_key_id[:8]}...")

import requests  # noqa: E402  (apres dotenv pour avoir les vars)

# ── Patch SSL global — certificat Windows non reconnu par httpcore/curl_cffi ──
import ssl as _ssl
_ssl._create_default_https_context = _ssl._create_unverified_context

# 1) anthropic SDK — injecter http_client verify=False dans Anthropic.__init__
try:
    import httpx as _httpx
    import anthropic as _anthropic

    _orig_anth_init = _anthropic.Anthropic.__init__
    def _p_anth_init(self, *a, **kw):
        if "http_client" not in kw:
            kw["http_client"] = _httpx.Client(verify=False)
        _orig_anth_init(self, *a, **kw)
    _anthropic.Anthropic.__init__ = _p_anth_init

    _orig_anth_async_init = _anthropic.AsyncAnthropic.__init__
    def _p_anth_async_init(self, *a, **kw):
        if "http_client" not in kw:
            kw["http_client"] = _httpx.AsyncClient(verify=False)
        _orig_anth_async_init(self, *a, **kw)
    _anthropic.AsyncAnthropic.__init__ = _p_anth_async_init
except Exception:
    pass

# 2) curl_cffi (yfinance) — patché AVANT import yfinance
try:
    import curl_cffi.requests as _cr
    _orig_cr_init = _cr.Session.__init__
    def _p_cr_init(self, *a, **kw): kw["verify"] = False; _orig_cr_init(self, *a, **kw)
    _cr.Session.__init__ = _p_cr_init
except ImportError:
    pass

# 3) requests standard
import urllib3 as _urllib3
_urllib3.disable_warnings(_urllib3.exceptions.InsecureRequestWarning)
requests.Session.verify = False

# Fusion TradingView MCP Kola (optionnel — CDP port 9222)
try:
    sys.path.insert(0, str(_HERE))
    from unified_bridge import (  # type: ignore
        compare_ta_and_tv,
        format_unified_whatsapp,
        merge_confirmed_with_tv,
        push_unified_state,
        resolve_conflict_loop,
        run_tv_analysis_for_bridge,
        send_unified_whatsapp,
    )
    _UNIFIED_BRIDGE_OK = True
except ImportError as _ub_err:
    _UNIFIED_BRIDGE_OK = False
    _UNIFIED_BRIDGE_ERR = str(_ub_err)

_SERVER_URL = os.getenv("AI_SERVER_URL", "http://127.0.0.1:8000").rstrip("/")

# ---------------------------------------------------------------------------
# Import TradingAgents
# ---------------------------------------------------------------------------

if str(_TA_REPO) not in sys.path:
    sys.path.insert(0, str(_TA_REPO))

# Ajouter le site-packages du venv TradingAgents pour que typer et autres dépendances soient disponibles
_TA_VENV_SITE = _TA_REPO / ".venv" / "Lib" / "site-packages"
if _TA_VENV_SITE.exists() and str(_TA_VENV_SITE) not in sys.path:
    sys.path.insert(0, str(_TA_VENV_SITE))

try:
    from tradingagents.graph.trading_graph import TradingAgentsGraph       # type: ignore
    from tradingagents.default_config import DEFAULT_CONFIG as _TA_DEFAULT  # type: ignore
    from tradingagents.dataflows.vendor_selection import build_runtime_config  # type: ignore
    from cli.main import get_user_selections                                 # type: ignore
    from cli.stream_pipeline import ANALYST_ORDER                            # type: ignore
    _TA_CLI_AVAILABLE = True
except ImportError as _e:
    _TA_CLI_AVAILABLE = False
    _TA_IMPORT_ERR = str(_e)

# ---------------------------------------------------------------------------
# Mapping MT5 -> yfinance
# ---------------------------------------------------------------------------

# Mapping ticker Deriv -> ticker retail pour les analystes social/news/fundamentals
# Permet à TradingAgents de trouver les données sur StockTwits, Reddit, Yahoo Finance
_SOCIAL_TICKER_MAP: Dict[str, str] = {
    # Deriv frx -> ticker retail pour social/news analysts
    # Proxies stables : futures ou indices plutôt que forex =X (souvent délisted)
    "frxXAUUSD": "GC=F",
    "frxXAGUSD": "SI=F",
    "frxEURUSD": "EUR=X",
    "frxGBPUSD": "GBP=X",
    "frxUSDJPY":  "JPY=X",
    "frxUSDCHF":  "CHF=X",
    "frxAUDUSD": "AUD=X",
    "frxUSDCAD": "CAD=X",
    # Crypto yfinance -> notation StockTwits (BTC.X, ETH.X...)
    "BTC-USD":  "BTC.X",
    "ETH-USD":  "ETH.X",
    "BNB-USD":  "BNB.X",
    "SOL-USD":  "SOL.X",
    "XRP-USD":  "XRP.X",
    "ADA-USD":  "ADA.X",
    "DOGE-USD": "DOGE.X",
    "AVAX-USD": "AVAX.X",
    "DOT-USD":  "DOT.X",
    "LTC-USD":  "LTC.X",
}

# Mapping MT5 -> ticker TradingAgents
# Deriv synthetiques : passer le nom court (BOOM900, CRASH300...)
# TradingAgents detecte automatiquement vendor=deriv pour BOOM*/CRASH*
_MT5_MAP: Dict[str, str] = {
    # Metaux / Futures
    "XAUUSD": "GC=F",
    "XAGUSD": "SI=F",
    # Forex
    "EURUSD": "EUR=X",
    "GBPUSD": "GBP=X",
    "USDJPY": "JPY=X",
    "USDCHF": "CHF=X",
    "AUDUSD": "AUD=X",
    "USDCAD": "CAD=X",
    "NZDUSD": "NZD=X",
    "EURGBP": "EURGBP=X",
    "EURJPY": "EURJPY=X",
    "GBPJPY": "GBPJPY=X",
    # Indices
    "US30":  "^DJI",
    "US500": "^GSPC",
    "USTEC": "^IXIC",
    "DE40":  "^GDAXI",
    # Deriv Boom — nom court reconnu par infer_stock_data_vendor()
    "BOOM 300 INDEX":   "BOOM300",
    "BOOM 500 INDEX":   "BOOM500",
    "BOOM 600 INDEX":   "BOOM600",
    "BOOM 900 INDEX":   "BOOM900",
    "BOOM 1000 INDEX":  "BOOM1000",
    # Deriv Crash
    "CRASH 300 INDEX":  "CRASH300",
    "CRASH 500 INDEX":  "CRASH500",
    "CRASH 600 INDEX":  "CRASH600",
    "CRASH 900 INDEX":  "CRASH900",
    "CRASH 1000 INDEX": "CRASH1000",
}

_DERIV_PREFIXES = ("BOOM", "CRASH", "STEP", "RDBULL", "RDBEAR", "JUMP", "RANGE", "VOL")


def _mt5_to_yfinance(symbol: str) -> str:
    """Convertit un symbole MT5 en ticker TradingAgents.
    Pour les synthetiques Deriv, retourne le nom court (BOOM900)
    afin que TradingAgents choisisse automatiquement vendor=deriv.
    """
    up = symbol.strip().upper()
    if up in _MT5_MAP:
        return _MT5_MAP[up]
    # Supprimer "INDEX" ou "INDX" en suffixe (ex: BOOM900INDEX -> BOOM900)
    for suffix in (" INDEX", " INDX", "INDEX", "INDX"):
        if up.endswith(suffix):
            up = up[: -len(suffix)].strip()
            break
    nospace = up.replace(" ", "")
    for p in _DERIV_PREFIXES:
        if nospace.startswith(p):
            return nospace  # BOOM900, CRASH600, etc.
    if len(up) == 6 and up.isalpha():
        return up + "=X"
    return up


# ---------------------------------------------------------------------------
# Selecteur interactif de symbole par categorie de broker
# ---------------------------------------------------------------------------

# Catalogue yfinance — symboles les plus tradés
_YFINANCE_CATALOG = {
    "Forex majeurs": [
        ("EURUSD=X", "EUR/USD"), ("GBPUSD=X", "GBP/USD"), ("JPY=X", "USD/JPY"),
        ("CHF=X", "USD/CHF"), ("AUDUSD=X", "AUD/USD"), ("CAD=X", "USD/CAD"),
        ("NZDUSD=X", "NZD/USD"), ("EURGBP=X", "EUR/GBP"), ("EURJPY=X", "EUR/JPY"),
    ],
    "Metaux / Commodites": [
        ("GC=F", "Or — XAUUSD"), ("SI=F", "Argent — XAGUSD"),
        ("CL=F", "Petrole WTI"), ("NG=F", "Gaz naturel"),
    ],
    "Indices boursiers": [
        ("^DJI", "Dow Jones — US30"), ("^GSPC", "S&P 500 — US500"),
        ("^IXIC", "Nasdaq — USTEC"), ("^GDAXI", "DAX — DE40"),
        ("^FTSE", "FTSE 100"), ("^N225", "Nikkei 225"),
    ],
    "Actions US": [
        ("AAPL", "Apple"), ("MSFT", "Microsoft"), ("NVDA", "NVIDIA"),
        ("TSLA", "Tesla"), ("AMZN", "Amazon"), ("GOOGL", "Alphabet"),
    ],
    "Crypto (yfinance)": [
        ("BTC-USD", "Bitcoin"), ("ETH-USD", "Ethereum"),
        ("BNB-USD", "BNB"), ("SOL-USD", "Solana"),
    ],
}

# Catalogue Deriv — directement depuis deriv_catalog
try:
    from tradingagents.dataflows.deriv_catalog import (  # type: ignore
        DERIV_BOOM_CRASH, DERIV_FX_METALS, DERIV_VOLATILITY
    )
    _DERIV_CATALOG = {
        "Boom & Crash (synthetiques)": [(sid, lbl) for sid, lbl in DERIV_BOOM_CRASH],
        "Forex & Metaux Deriv":        [(sid, lbl) for sid, lbl in DERIV_FX_METALS],
        "Volatility Index":            [(sid, lbl) for sid, lbl in DERIV_VOLATILITY],
    }
except ImportError:
    _DERIV_CATALOG = {
        "Boom & Crash": [
            ("BOOM1000","Boom 1000"), ("BOOM900","Boom 900"), ("BOOM600","Boom 600"),
            ("CRASH1000","Crash 1000"), ("CRASH900","Crash 900"), ("CRASH600","Crash 600"),
        ],
    }

# Catalogue Weltrade — broker standard MetaTrader 5
# Symboles compatibles yfinance pour l'analyse (meme donnees, differents spreads)
_WELTRADE_CATALOG = {
    "Forex Majeurs": [
        ("EURUSD=X", "EUR/USD"), ("GBPUSD=X", "GBP/USD"), ("USDJPY=X", "USD/JPY"),
        ("USDCHF=X", "USD/CHF"), ("AUDUSD=X", "AUD/USD"), ("USDCAD=X", "USD/CAD"),
        ("NZDUSD=X", "NZD/USD"), ("EURGBP=X", "EUR/GBP"), ("EURJPY=X", "EUR/JPY"),
        ("GBPJPY=X", "GBP/JPY"), ("EURAUD=X", "EUR/AUD"), ("EURCHF=X", "EUR/CHF"),
    ],
    "Forex Mineurs & Exotiques": [
        ("USDMXN=X", "USD/MXN"), ("USDZAR=X", "USD/ZAR"), ("USDTRY=X", "USD/TRY"),
        ("USDSGD=X", "USD/SGD"), ("USDHKD=X", "USD/HKD"), ("USDNOK=X", "USD/NOK"),
        ("USDSEK=X", "USD/SEK"),
    ],
    "Metaux Precieux": [
        ("GC=F", "Or — XAU/USD"), ("SI=F", "Argent — XAG/USD"),
        ("PL=F", "Platine — XPT/USD"), ("PA=F", "Palladium — XPD/USD"),
    ],
    "Indices CFD": [
        ("^DJI",  "US30 — Dow Jones"), ("^GSPC", "US500 — S&P 500"),
        ("^IXIC", "USTEC — Nasdaq"),   ("^GDAXI","DE40 — DAX"),
        ("^FTSE", "UK100 — FTSE"),     ("^N225", "JP225 — Nikkei"),
        ("^FCHI", "FR40 — CAC 40"),    ("^STOXX50E", "EU50 — Euro Stoxx"),
    ],
    "Energie & Matieres Premieres": [
        ("CL=F", "Petrole WTI"), ("BZ=F", "Petrole Brent"),
        ("NG=F", "Gaz naturel"), ("ZC=F", "Mais"), ("ZW=F", "Ble"),
    ],
    "Crypto CFD": [
        ("BTC-USD", "Bitcoin / USD"), ("ETH-USD", "Ethereum / USD"),
        ("LTC-USD", "Litecoin / USD"), ("XRP-USD", "Ripple / USD"),
    ],
}

# Catalogue Exness — broker MT5 avec tres faibles spreads
# Exness propose les memes paires forex/metaux + indices
_EXNESS_CATALOG = {
    "Forex Majeurs (spread ultra-faible)": [
        ("EURUSD=X", "EUR/USD"), ("GBPUSD=X", "GBP/USD"), ("USDJPY=X", "USD/JPY"),
        ("USDCHF=X", "USD/CHF"), ("AUDUSD=X", "AUD/USD"), ("USDCAD=X", "USD/CAD"),
        ("NZDUSD=X", "NZD/USD"),
    ],
    "Croisees Majeures": [
        ("EURGBP=X", "EUR/GBP"), ("EURJPY=X", "EUR/JPY"), ("GBPJPY=X", "GBP/JPY"),
        ("EURAUD=X", "EUR/AUD"), ("GBPAUD=X", "GBP/AUD"), ("AUDJPY=X", "AUD/JPY"),
        ("CADJPY=X", "CAD/JPY"), ("CHFJPY=X", "CHF/JPY"),
    ],
    "Metaux & Energie": [
        ("GC=F", "XAU/USD — Or"),   ("SI=F", "XAG/USD — Argent"),
        ("CL=F", "WTI Petrole"),    ("BZ=F", "Brent Petrole"),
    ],
    "Indices Boursiers": [
        ("^GSPC",  "US500 — S&P 500"), ("^DJI",   "US30 — Dow Jones"),
        ("^IXIC",  "USTEC — Nasdaq"),  ("^GDAXI", "GER40 — DAX"),
        ("^FTSE",  "UK100 — FTSE"),    ("^HSI",   "HK50 — Hang Seng"),
        ("^N225",  "JP225 — Nikkei"),
    ],
    "Crypto (Exness propose BTC etc.)": [
        ("BTC-USD", "Bitcoin"), ("ETH-USD", "Ethereum"),
        ("SOL-USD", "Solana"),  ("BNB-USD", "BNB"),
        ("XRP-USD", "XRP"),     ("DOGE-USD","Dogecoin"),
    ],
}

# Vendor mapping par broker
_BROKER_VENDOR = {
    "Deriv (synthetiques + forex)":          "deriv",
    "yfinance (Forex / Indices / Actions)":  "yfinance",
    "Weltrade (MT5 — Forex / Indices / Crypto)": "yfinance",
    "Exness (MT5 — Forex / Metaux / Crypto)":    "yfinance",
}

_ALL_BROKERS = {
    "Deriv (synthetiques + forex)":              _DERIV_CATALOG,
    "Weltrade (MT5 — Forex / Indices / Crypto)": _WELTRADE_CATALOG,
    "Exness (MT5 — Forex / Metaux / Crypto)":    _EXNESS_CATALOG,
    "yfinance (Forex / Indices / Actions)":       _YFINANCE_CATALOG,
}


def _prompt(label: str, back: bool = True, quit: bool = True) -> str:
    """Affiche un prompt avec les options [B]ack et [Q]uit disponibles."""
    hints = []
    if back:
        hints.append("B=retour")
    hints.append("Q=quitter")
    suffix = f"  ({', '.join(hints)})" if hints else ""
    return input(f"\n  Choix{suffix} : ").strip()


def select_symbol_interactive() -> tuple:
    """
    Menu interactif en 3 niveaux : broker -> categorie -> symbole.
    Supporte [B] pour revenir au niveau precedent et [Q] pour quitter.
    Retourne (symbol_display, ticker_ta, vendor).
    """
    print("\n" + "=" * 60)
    print("  SELECTION DU SYMBOLE A ANALYSER")
    print("=" * 60)

    brokers = list(_ALL_BROKERS.keys())

    while True:
        # ── Niveau 1 : broker ────────────────────────────────────────────
        print("\n  Broker / Source de donnees :")
        for i, b in enumerate(brokers, 1):
            print(f"    [{i}] {b}")
        print("    [0] Saisir manuellement un ticker")
        raw = _prompt("Broker", back=False).upper()

        if raw == "Q":
            sys.exit("[bridge] Annule.")

        if raw == "0":
            ticker = input("  Ticker (ex: EURUSD=X, BOOM900, AAPL): ").strip()
            if not ticker:
                print("  [!] Aucun ticker saisi, retour au menu.")
                continue
            vendor = "deriv" if any(ticker.upper().startswith(p) for p in ("BOOM","CRASH","1HZ","R_","FRX")) else "yfinance"
            return ticker, ticker, vendor

        try:
            broker_idx = int(raw) - 1
            broker_name = brokers[broker_idx]
        except (ValueError, IndexError):
            print("  [!] Choix invalide.")
            continue

        categories = list(_ALL_BROKERS[broker_name].keys())
        vendor = _BROKER_VENDOR.get(broker_name, "yfinance")

        while True:
            # ── Niveau 2 : categorie ──────────────────────────────────────
            print(f"\n  Categorie  ({broker_name}) :")
            for i, c in enumerate(categories, 1):
                print(f"    [{i}] {c}")
            raw2 = _prompt("Categorie").upper()

            if raw2 == "Q":
                sys.exit("[bridge] Annule.")
            if raw2 == "B":
                break  # retour niveau 1

            try:
                cat_idx = int(raw2) - 1
                cat_name = categories[cat_idx]
            except (ValueError, IndexError):
                print("  [!] Choix invalide.")
                continue

            symbols = _ALL_BROKERS[broker_name][cat_name]

            while True:
                # ── Niveau 3 : symbole ────────────────────────────────────
                print(f"\n  Symbole  ({cat_name}) :")
                for i, (sid, lbl) in enumerate(symbols, 1):
                    print(f"    [{i:2}] {lbl:35} ({sid})")
                raw3 = _prompt("Symbole").upper()

                if raw3 == "Q":
                    sys.exit("[bridge] Annule.")
                if raw3 == "B":
                    break  # retour niveau 2

                try:
                    sym_idx = int(raw3) - 1
                    ticker_id, sym_label = symbols[sym_idx]
                except (ValueError, IndexError):
                    print("  [!] Choix invalide.")
                    continue

                print(f"\n  [OK] Symbole selectionne : {sym_label} ({ticker_id}) | vendor={vendor}")
                return sym_label, ticker_id, vendor
            # fin while niveau 3 — B -> retour niveau 2
        # fin while niveau 2 — B -> retour niveau 1


# ---------------------------------------------------------------------------
# Rating 5-tier -> BUY / SELL / HOLD
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# System prompt dynamique par catégorie de symbole (TradBOT Signal Engine v2.0)
# ---------------------------------------------------------------------------

def _get_symbol_category(symbol: str) -> str:
    """Catégorise le symbole : CRASH | BOOM | GOLD | CRYPTO | INDEX | VOLATILITY | FOREX"""
    up = symbol.strip().upper()
    if "CRASH" in up: return "CRASH"
    if "BOOM"  in up: return "BOOM"
    if "XAU" in up or "GOLD" in up or up in ("GC=F",): return "GOLD"
    if any(up.startswith(p) for p in ("1HZ","R_","VOL","STEP","JUMP","RANGE")): return "VOLATILITY"
    if any(up in (f,) for f in ("^DJI","^GSPC","^IXIC","^GDAXI","^FTSE","US30","US500","USTEC","DE40")): return "INDEX"
    _CRYPTO_BASES = ("BTC","ETH","BNB","SOL","XRP","ADA","DOT","DOGE","AVAX","MATIC","LTC","LINK")
    if any(up.startswith(b) or up.endswith(b) for b in _CRYPTO_BASES): return "CRYPTO"
    if "-USD" in up or "-USDT" in up or "-USDC" in up: return "CRYPTO"
    return "FOREX"


def _build_system_prompt(symbol: str, category: str, indicators: Optional[Dict]) -> str:
    """
    Construit le system prompt TradBOT Signal Engine v2.0 adapté à la catégorie.
    Injecte les données techniques réelles si disponibles.
    """
    price   = indicators.get("current_price", "?") if indicators else "?"
    atr_raw = indicators.get("atr", None)           if indicators else None
    atr     = atr_raw if atr_raw else "?"
    rsi     = indicators.get("rsi", "?")            if indicators else "?"
    macd    = indicators.get("macd", "?")           if indicators else "?"
    sma20   = indicators.get("sma_20", "?")         if indicators else "?"
    sma50   = indicators.get("sma_50", "?")         if indicators else "?"
    sma200  = indicators.get("sma_200", "?")        if indicators else "?"
    bb_up   = indicators.get("bb_upper", "?")       if indicators else "?"
    bb_low  = indicators.get("bb_lower", "?")       if indicators else "?"

    # Calcul des zones de spike avec le prix réel
    # LOT MINIMUM DERIV :
    #   Boom/Crash synthétiques : 0.2
    #   XAUUSD, Forex, Volatility, Index : 0.01
    lot_min = 0.2 if category in ("BOOM", "CRASH") else 0.01
    capital = 20.0  # Capital cible scalping
    risk_pct = 0.02
    risk_amt = capital * risk_pct  # $0.40

    # Zones de prix spike si données disponibles
    spike_zones_txt = ""
    if price != "?" and atr_raw:
        try:
            p = float(price)
            a = float(atr_raw)
            if category == "BOOM":
                z1 = round(p - a * 0.5, 3)   # Zone rebond proche
                z2 = round(p - a * 1.0, 3)   # Zone rebond modérée
                z3 = round(p - a * 2.0, 3)   # Zone rebond agressive
                sl1 = round(z1 - a * 1.0, 3)
                tp1 = round(z1 + a * 2.0, 3)
                tp2 = round(z1 + a * 3.5, 3)
                # Risque avec lot 0.2 : lot × SL_pts (approximation)
                sl_pts1 = round(abs(z1 - sl1), 3)
                risk_est = round(lot_min * sl_pts1 * 0.01, 3)  # ~$0.01 par lot par point
                spike_zones_txt = f"""
ZONES DE SPIKE CALCULEES (prix actuel = {p}, ATR = {a}) :
  Zone rebond Z1 (proche)   : {z1}  (-{round(p-z1,1)} pts)
  Zone rebond Z2 (moderee)  : {z2}  (-{round(p-z2,1)} pts)
  Zone rebond Z3 (agressive): {z3}  (-{round(p-z3,1)} pts)

  SIGNAL BUY SNIPER (entrée Z1) :
  Entrée : {z1}  |  SL : {sl1}  |  TP1 : {tp1}  |  TP2 : {tp2}
  Lot minimum Deriv : {lot_min} | Risque estimé : ~${risk_est} sur capital ${capital}
  Si risque > $0.40 (2% de $20) avec lot {lot_min} → EVITER ce trade"""
            elif category == "CRASH":
                z1 = round(p + a * 0.5, 3)
                z2 = round(p + a * 1.0, 3)
                z3 = round(p + a * 2.0, 3)
                sl1 = round(z1 + a * 1.0, 3)
                tp1 = round(z1 - a * 2.0, 3)
                tp2 = round(z1 - a * 3.5, 3)
                sl_pts1 = round(abs(z1 - sl1), 3)
                risk_est = round(lot_min * sl_pts1 * 0.01, 3)
                spike_zones_txt = f"""
ZONES DE SPIKE CALCULEES (prix actuel = {p}, ATR = {a}) :
  Zone chute Z1 (proche)    : {z1}  (+{round(z1-p,1)} pts)
  Zone chute Z2 (moderee)   : {z2}  (+{round(z2-p,1)} pts)
  Zone chute Z3 (agressive) : {z3}  (+{round(z3-p,1)} pts)

  SIGNAL SELL SNIPER (entrée Z1) :
  Entrée : {z1}  |  SL : {sl1}  |  TP1 : {tp1}  |  TP2 : {tp2}
  Lot minimum Deriv : {lot_min} | Risque estimé : ~${risk_est} sur capital ${capital}
  Si risque > $0.40 (2% de $20) avec lot {lot_min} → EVITER ce trade"""
        except Exception:
            pass

    # Contexte technique commun
    tech_ctx = f"""
DONNEES TECHNIQUES ACTUELLES ({symbol}) :
  Prix      : {price}
  ATR       : {atr}
  RSI(14)   : {rsi}
  MACD      : {macd}
  SMA 20    : {sma20} | SMA 50 : {sma50} | SMA 200 : {sma200}
  BB Upper  : {bb_up} | BB Lower : {bb_low}

CONTRAINTES BROKER/CAPITAL :
  Lot minimum Deriv ({category}) : {lot_min}  ← FIXE, NE PAS CHANGER
  Capital scalping cible         : ${capital}
  Risque max par trade           : {risk_pct*100:.0f}% = ${risk_amt:.2f}
  REGLE ABSOLUE : Toujours utiliser lot={lot_min} (lot minimum broker) — ne jamais proposer 0 lot.
  Si le risque calculé dépasse ${risk_amt:.2f} → RÉDUIRE le SL (pas éviter le trade).
  Un setup avec lot {lot_min} et SL serré est TOUJOURS préférable à "ÉVITER ce trade".
{spike_zones_txt}
"""

    # Règles anti-contradiction communes
    rules_common = """
REGLES OBLIGATOIRES :
1. Un seul signal final (pas de double signal contradictoire)
2. Si agents internes divergent, choisir UN camp et expliquer en 2 lignes
3. Sizing cohérent avec le stop loss
4. Si ATR > 2% du prix : étiqueter HAUTE VOLATILITE - TAILLE REDUITE OBLIGATOIRE
5. Ratio R:R minimum 1:1.5 — sinon écrire SETUP SOUS-OPTIMAL - ATTENDRE
6. Si setup défavorable : ecrire clairement AUCUN SIGNAL — Raison : [X]. Revenez quand : [condition]

FORMAT DE SORTIE OBLIGATOIRE :

SECTION 1 — CONTEXTE CHRONOLOGIQUE (5-10 lignes narrative)
SECTION 2 — SIGNAL :
  INSTRUMENT : | CATEGORIE : | DIRECTION : BUY/SELL/ATTENDRE
  CONVICTION : [1-10] + justification 1 ligne
  ENTREE : | STOP LOSS : ([X] pips) | TP1 : | TP2 :
  TAILLE : Compte $10 = [X] lot | Compte $50 = [X] lot
  INVALIDATION : [condition]
SECTION 3 — RISQUES SPECIFIQUES (2-3 risques concrets)
SECTION 4 — CONDITION DE SKIP si applicable
"""

    # Extraire le nombre de ticks théorique depuis le nom du symbole
    import re as _re2
    tick_num = "500"
    m_tick = _re2.search(r'(\d+)', symbol)
    if m_tick: tick_num = m_tick.group(1)

    # Seuils de probabilité selon la fréquence
    try:
        freq = int(tick_num)
        zone_froide   = freq // 2
        zone_neutre   = int(freq * 0.75)
        zone_chaude   = int(freq * 0.85)
        zone_critique = freq
    except Exception:
        zone_froide, zone_neutre, zone_chaude, zone_critique = 250, 375, 425, 500

    if category == "CRASH":
        specific = f"""
══════════════════════════════════════════════════════
REGLE 0 — VERIFICATION OBLIGATOIRE (PRIORITE ABSOLUE)
══════════════════════════════════════════════════════
INSTRUMENT CRASH DETECTE : {symbol}
DIRECTION AUTORISEE : SELL UNIQUEMENT
TOUT SIGNAL BUY ENTRANT EST INVALIDE ET DOIT ETRE REJETE.
Si TradingAgents propose BUY sur cet instrument, signaler :
"CONFLIT DETECTE : TradingAgents propose BUY sur CRASH. Signal invalide par nature de l'instrument. IGNORE."
Reconstruire l'analyse en SELL uniquement.
══════════════════════════════════════════════════════

NATURE DU CRASH INDEX ({symbol}) :
Indice synthetique Deriv generant des spikes BAISSIERS a frequence statistique ~1 spike / {tick_num} ticks.
Entre les spikes, le prix derive LEGEREMENT a la HAUSSE — c'est la phase d'accumulation. Ce trend haussier
N'EST PAS un signal BUY. Les patterns haussiers (golden cross, RSI>50, MACD+) sur un Crash sont des
ARTEFACTS de la derive normale entre deux spikes. Les analyser comme une action ordinaire est une erreur fondamentale.

ZONES DE PROBABILITE SPIKE BAISSIER :
→ < {zone_froide} ticks depuis dernier spike : Zone FROIDE — attendre
→ {zone_froide}-{zone_neutre} ticks          : Zone NEUTRE — vigilance
→ {zone_neutre}-{zone_chaude} ticks          : Zone CHAUDE — probabilite elevee
→ > {zone_chaude} ticks                      : Zone CRITIQUE — spike imminent probable
→ > {zone_critique} ticks                    : ALERTE MAXIMALE

ANALYSE REQUISE :
1. Identifier le dernier spike baissier visible (bougie avec mèche basse disproportionnee)
   → Date/heure dernier spike : [X] | Prix : [X]
2. Estimer les ticks ecoulés depuis (nombre de bougies × ticks/bougie moyen)
   → Ticks estimes : [X] | Zone : [froide/neutre/chaude/critique]
3. Position du prix dans sa range recente (50 dernieres bougies) :
   → Prix dans les 20% HAUTS = IDEAL pour SELL (plus d'espace a la baisse)
   → Prix au milieu = acceptable | Prix dans les 20% bas = EVITER
4. Bollinger Bands : bandes qui se resserrent = spike imminent potentiel
5. SELL LIMIT se place AU-DESSUS du prix actuel (resistance)

FORMAT SIGNAL FINAL CRASH :
  INSTRUMENT : {symbol} | CATEGORIE : CRASH INDEX — SELL UNIQUEMENT
  TICKS DEPUIS DERNIER SPIKE : [X] | ZONE : [froide/neutre/chaude/critique]
  POSITION DANS LA RANGE : [haut/milieu/bas]
  DIRECTION : SELL | CONVICTION : [1-10]
  ENTREE : [prix — en haut de range] | STOP : [prix +1.5xATR au-dessus]
  TP1 : [prochain support] | TP2 : [spike complet estime]
  Compte $10 = [X] lot (risque 2% = $0.20) | Compte $50 = [X] lot (risque 2% = $1.00)
  STRATEGIE : Entrer en SELL LIMIT en haut de range, attendre le spike baissier.
  Ne pas paniquer sur les micro-hausses entre ticks — c'est la derive normale.
  INVALIDATION : Si prix casse la resistance avec > 2xATR sans spike
"""
    elif category == "BOOM":
        specific = f"""
══════════════════════════════════════════════════════
REGLE 0 — VERIFICATION OBLIGATOIRE (PRIORITE ABSOLUE)
══════════════════════════════════════════════════════
INSTRUMENT BOOM DETECTE : {symbol}
DIRECTION AUTORISEE : BUY UNIQUEMENT
TOUT SIGNAL SELL ENTRANT EST INVALIDE ET DOIT ETRE REJETE.
Si TradingAgents propose SELL sur cet instrument, signaler :
"CONFLIT DETECTE : TradingAgents propose SELL sur BOOM. Signal invalide par nature de l'instrument. IGNORE."
Reconstruire l'analyse en BUY uniquement.
══════════════════════════════════════════════════════

NATURE DU BOOM INDEX ({symbol}) :
Indice synthetique Deriv generant des spikes HAUSSIERS a frequence statistique ~1 spike / {tick_num} ticks.
Entre les spikes, le prix derive LEGEREMENT a la BAISSE — c'est la phase d'accumulation.
Ce trend baissier N'EST PAS un signal SELL. Les patterns baissiers (death cross, MACD-, triangle
descendant) sur un Boom sont des ARTEFACTS de la derive naturelle entre deux spikes.
Les analyser comme une action ordinaire est une erreur fondamentale.

ZONES DE PROBABILITE SPIKE HAUSSIER :
→ < {zone_froide} ticks depuis dernier spike : Zone FROIDE — attendre
→ {zone_froide}-{zone_neutre} ticks          : Zone NEUTRE — vigilance
→ {zone_neutre}-{zone_chaude} ticks          : Zone CHAUDE — probabilite elevee
→ > {zone_chaude} ticks                      : Zone CRITIQUE — spike imminent probable
→ > {zone_critique} ticks                    : ALERTE MAXIMALE

ANALYSE REQUISE :
1. Identifier le dernier spike haussier visible (bougie avec mèche haute disproportionnee)
   → Date/heure dernier spike : [X] | Prix : [X]
2. Estimer les ticks ecoulés depuis
   → Ticks estimes : [X] | Zone : [froide/neutre/chaude/critique]
3. Position du prix dans sa range recente (50 dernieres bougies) :
   → Prix dans les 20% BAS = IDEAL pour BUY (spike aura plus d'espace)
   → Prix au milieu = acceptable | Prix dans les 20% hauts = EVITER
4. Bollinger Bands : bandes qui se resserrent = energie accumulee, breakout imminent
5. BUY LIMIT se place EN DESSOUS du prix actuel (support)

FORMAT SIGNAL FINAL BOOM :
  INSTRUMENT : {symbol} | CATEGORIE : BOOM INDEX — BUY UNIQUEMENT
  TICKS DEPUIS DERNIER SPIKE : [X] | ZONE : [froide/neutre/chaude/critique]
  POSITION DANS LA RANGE : [bas/milieu/haut]
  DIRECTION : BUY | CONVICTION : [1-10]
  ENTREE : [prix — en bas de range] | STOP : [prix -1.5xATR en dessous]
  TP1 : [+50 a +100 pips spike minimum] | TP2 : [+150 a +300 pips spike complet]
  Compte $10 = [X] lot (risque 2% = $0.20) | Compte $50 = [X] lot (risque 2% = $1.00)
  STRATEGIE : Entrer en BUY LIMIT en bas de range, tenir la position.
  Les micro-baisses entre ticks sont NORMALES — ne pas sortir prematurément.
  Sortir manuellement sur la mèche du spike.
  INVALIDATION : Si prix casse sous le support critique avec > 2xATR sans rebond
"""
    elif category == "GOLD":
        specific = """
TU ANALYSES L'OR (XAUUSD).

LOGIQUE FONDAMENTALE : Price Action + Structure HTF → LTF

NARRATIVE CHRONOLOGIQUE 30 JOURS OBLIGATOIRE :
Format : "[Date] : Prix a atteint [X] → rejet → formation de [structure]"
        "[Date] : Cassure de [niveau] → confirmation → bias devient [bull/bear]"
        "Aujourd'hui : Prix à [X], dernier retest du niveau [Y] date du [Z]"

NIVEAUX CLES A IDENTIFIER :
- Dernier Higher High / Lower Low significatif
- Zone de liquidité (equal highs/lows)
- Imbalance FVG la plus proche
- Niveau psychologique le plus proche (round number)

REGLES OR :
- Entrée SUR structure confirmée (pas en milieu de range)
- Stop : sous/sur le dernier swing significatif
- TP1 : prochain niveau de liquidité
- TP2 : prochain FVG ou résistance majeure
- Confirmation requise : bougie de rejet OU engulfing sur niveau clé
- Sans confirmation → écrire ATTENDRE SETUP
"""
    elif category == "FOREX":
        specific = f"""
TU ANALYSES UNE PAIRE FOREX ({symbol}).

ANALYSE REQUISE :
1. SESSION ACTIVE : London / New York / Asia / Overlap ?
2. STRUCTURE : Tendance H4 → H1 → M15 alignées ?
3. NIVEAU D'ENTREE :
   - Support/Résistance testés minimum 2 fois
   - Zone de valeur (38.2% à 61.8% Fibonacci du dernier swing)
   - Confirmation : Pin bar, Engulfing, Inside bar
4. NE PAS entrer en plein milieu d'une range sans confirmation
5. Sessions haute liquidité préférées : London Open (08h-10h UTC), NY Open (13h-15h UTC)
"""
    elif category == "VOLATILITY":
        # Extraire les données actuelles pour la logique mean reversion
        rsi_val  = float(indicators.get("rsi", 50))  if indicators else 50.0
        bb_up_v  = float(indicators.get("bb_upper", 0)) if indicators else 0.0
        bb_low_v = float(indicators.get("bb_lower", 0)) if indicators else 0.0
        price_v  = float(indicators.get("current_price", 0)) if indicators else 0.0
        atr_v    = float(indicators.get("atr", 0)) if indicators else 0.0

        # Déterminer la direction correcte par mean reversion
        mr_direction = "SKIP"
        mr_reason = ""
        if rsi_val < 30:
            mr_direction = "BUY"
            mr_reason = f"RSI={rsi_val:.1f} SURVENDU (< 30) → mean reversion haussière attendue"
        elif rsi_val > 70:
            mr_direction = "SELL"
            mr_reason = f"RSI={rsi_val:.1f} SURACHETÉ (> 70) → mean reversion baissière attendue"
        elif price_v > 0 and bb_low_v > 0 and price_v <= bb_low_v * 1.005:
            mr_direction = "BUY"
            mr_reason = f"Prix proche BB inférieure ({bb_low_v:.3f}) → extension baissière, rebond probable"
        elif price_v > 0 and bb_up_v > 0 and price_v >= bb_up_v * 0.995:
            mr_direction = "SELL"
            mr_reason = f"Prix proche BB supérieure ({bb_up_v:.3f}) → extension haussière, recul probable"
        else:
            mr_reason = f"RSI={rsi_val:.1f} en zone neutre (30-70), prix au milieu des BB → SKIP"

        specific = f"""
══════════════════════════════════════════════════════
REGLE 0 — VOLATILITY INDEX : MEAN REVERSION UNIQUEMENT
══════════════════════════════════════════════════════
INSTRUMENT : {symbol}
DIRECTION CALCULEE PAR MEAN REVERSION : {mr_direction}
RAISON : {mr_reason}

Si TradingAgents propose une direction OPPOSEE, signaler :
"CONFLIT DETECTE : TradingAgents propose [X] mais la mean reversion indique [Y].
Sur un Volatility Index, la mean reversion PRIME pour le scalping. Signal TradingAgents IGNORE."
══════════════════════════════════════════════════════

NATURE DU VOLATILITY INDEX ({symbol}) :
Simule une volatilité CONSTANTE dans les deux sens. Aucun spike directionnel prévisible.
Le prix oscille autour d'une moyenne avec des extensions périodiques.

ERREUR CLASSIQUE À ÉVITER :
Voir une tendance baissière forte et conclure SELL. Sur un Volatility Index, les extensions
extrêmes (RSI < 30 ou > 70) sont les MEILLEURES opportunités dans la direction OPPOSEE.
Suivre la tendance de fond sur un Volatility Index est une erreur fondamentale pour le scalping.

STRATÉGIE CORRECTE : MEAN REVERSION
→ RSI < 30 + prix proche BB inférieure → BUY (rebond vers la moyenne)
→ RSI > 70 + prix proche BB supérieure → SELL (recul vers la moyenne)
→ RSI 30-70 + prix au centre des BB    → SKIP (pas de setup)
→ Jamais entrer en milieu de range — attendre les extrêmes

SIZING STRICT (REGLE ABSOLUE) :
  Compte $10 → risque MAX 2% = $0.20 | lot MIN = 0.01
  Compte $50 → risque MAX 2% = $1.00 | lot MIN = 0.01
  CALCUL : lot = risque_max / (SL_en_points × pip_value)
  NE JAMAIS inverser : lot fixe PUIS calcul risque → DANGEREUX

REGLE ANTI-CONTRADICTION RÉSUMÉ :
  Le résumé DOIT reprendre la même direction que l'analyse scalping.
  Si TradingAgents dit SELL mais mean reversion dit BUY → le résumé dit BUY.
  Le scalping court terme prime TOUJOURS sur le signal fondamental.

FORMAT SIGNAL VOLATILITY :
  DIRECTION : {mr_direction} (mean reversion) — {"CONFIRME" if mr_direction != "SKIP" else "PAS DE SETUP ACTUELLEMENT"}
  Entrée : [prix BB inférieure/supérieure ou RSI extrême]
  Stop   : [0.8x ATR — serré car mean reversion rapide]
  TP     : [SMA20 ou centre des BB — 1.5x ATR max]
  Compte $10 = 0.01 lot (risque ~$[X]) | Compte $50 = 0.01-0.05 lot
"""
    elif category == "CRYPTO":
        specific = f"""
TU ANALYSES UNE CRYPTOMONNAIE ({symbol}).

SPECIFICITES CRYPTO :
1. Marché 24/7 — pas de sessions fixes, mais pics de liquidité 13h-17h UTC (overlap US)
2. Volatilité élevée : ATR typiquement 2-5% du prix, stops plus larges que Forex
3. Corrélation BTC dominante : BTC en baisse = altcoins amplifient la baisse (beta > 1)
4. On-chain signals : flux ETF, volumes exchange, whale movements = catalyseurs majeurs
5. Niveaux psychologiques ronds (50K, 60K, 70K, 80K...) = zones de support/résistance fortes

ANALYSE REQUISE :
1. STRUCTURE : Tendance D1 → H4 → H1 alignée ?
2. DOMINANCE BTC : Si BTC < 50% dominance → altcoin season, sinon BTC leads
3. NIVEAUX CLÉS : Supports/résistances psychologiques + EMA200 D1
4. SENTIMENT : Fear & Greed Index, flux ETF spot Bitcoin, positions futures (funding rate)
5. CATALYSEURS : Halvings, régulation, adoption institutionnelle, macroéconomie (Fed)

RÈGLES GESTION DU RISQUE CRYPTO :
- SL minimum : 2x ATR (volatilité élevée = stops serrés = liquidation)
- Position max : 1-2% capital sur cryptos (vs 1% Forex)
- Éviter entrées avant annonces macro US majeures (CPI, FOMC)
- TP partiel à 1:1 recommandé vu la volatilité bidirectionnelle
"""
    else:  # INDEX
        specific = f"""
TU ANALYSES UN INDICE BOURSIER ({symbol}).

LOGIQUE :
- Suivre la tendance macro (données économiques, sentiment de marché)
- Structure Daily/H4 pour le biais directionnel
- Entrées sur pullbacks vers EMA50 ou zones de demande/offre
- News économiques majeures = éviter les entrées 30min avant/après
"""

    return f"""Tu es TradBOT Signal Engine v2.0 — expert trading sur Deriv et marchés financiers.

CATEGORIE IDENTIFIEE : {category} ({symbol})
{tech_ctx}
{specific}
{rules_common}"""


_RATING_MAP = {
    "BUY": "BUY", "OVERWEIGHT": "BUY",
    "HOLD": "HOLD", "NEUTRAL": "HOLD",
    "UNDERWEIGHT": "SELL", "SELL": "SELL",
}


def _normalize_rating(raw: str) -> str:
    return _RATING_MAP.get(str(raw).strip().upper(), "HOLD")


# ---------------------------------------------------------------------------
# Parse entry/SL/TP du markdown TraderProposal
# ---------------------------------------------------------------------------

def _parse_float_md(text: str, label: str) -> Optional[float]:
    m = _re.search(rf"\*\*{label}\*\*[:\s]+([0-9]+(?:\.[0-9]+)?)", text, _re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return None


def _extract_order_params(final_state: Dict[str, Any]) -> Dict[str, Optional[float]]:
    combined = (str(final_state.get("trader_investment_plan") or "") + "\n"
                + str(final_state.get("final_trade_decision") or ""))
    return {
        "entry_price": _parse_float_md(combined, "Entry Price"),
        "stop_loss":   _parse_float_md(combined, "Stop Loss"),
        "take_profit": _parse_float_md(combined, "Take Profit"),
    }


# ---------------------------------------------------------------------------
# Mode 1 : wizard complet (reutilise get_user_selections du CLI TA)
# ---------------------------------------------------------------------------

def run_with_wizard() -> Dict[str, Any]:
    """Lance le wizard interactif complet du CLI TradingAgents, puis analyse."""
    if not _TA_CLI_AVAILABLE:
        sys.exit(f"[bridge] CLI TradingAgents inaccessible: {_TA_IMPORT_ERR}")

    print("\n[bridge] Lancement du wizard TradingAgents...\n")
    selections = get_user_selections()

    config = build_runtime_config(selections, checkpoint=False)

    # Normaliser la liste des analystes dans l'ordre fixe
    selected_set = {a.value for a in selections["analysts"]}
    selected_analyst_keys: List[str] = [a for a in ANALYST_ORDER if a in selected_set]

    ticker     = selections["ticker"]
    trade_date = selections["analysis_date"]

    print(f"\n[bridge] Analyse: {ticker} | {trade_date} | analystes: {selected_analyst_keys}")

    graph = TradingAgentsGraph(
        selected_analysts=selected_analyst_keys,
        config=config,
        debug=False,
    )
    final_state, signal_rating = graph.propagate(ticker, trade_date)

    return {
        "symbol": ticker,
        "data_ticker": ticker,
        "signal_rating": signal_rating,
        "final_state": final_state,
    }


# ---------------------------------------------------------------------------
# Calcul des indicateurs techniques directement depuis OHLC Deriv
# (quand ai_server n'a pas de donnees temps reel pour le symbole)
# ---------------------------------------------------------------------------

def compute_indicators_from_deriv(ticker: str) -> Optional[Dict[str, Any]]:
    """
    Telecharge les OHLC Deriv et calcule RSI, MACD, ATR, SMA, Bollinger.
    Retourne un dict d indicateurs ou None si echec.
    """
    try:
        import sys as _sys
        if str(_TA_REPO) not in _sys.path:
            _sys.path.insert(0, str(_TA_REPO))

        from tradingagents.dataflows.deriv_market import fetch_deriv_daily_ohlcv, resolve_deriv_symbol  # type: ignore
        from tradingagents.dataflows.stockstats_utils import _clean_dataframe  # type: ignore
        from stockstats import StockDataFrame  # type: ignore
        import pandas as pd

        deriv_sym = resolve_deriv_symbol(ticker)
        end_date  = str(date.today())
        start_date = str(date.today().replace(year=date.today().year - 1))

        df = fetch_deriv_daily_ohlcv(deriv_sym, start_date, end_date)
        if df is None or len(df) < 20:
            return None

        df = _clean_dataframe(df)
        sdf = StockDataFrame.retype(df.copy())

        def _safe(key):
            try:
                v = sdf[key].iloc[-1]
                return None if (v != v) else round(float(v), 5)
            except Exception:
                return None

        current_price = _safe("close")
        if not current_price:
            return None

        result = {
            "current_price": current_price,
            "open":    _safe("open"),
            "high":    _safe("high"),
            "low":     _safe("low"),
            "close":   current_price,
            "rsi":     _safe("rsi_14"),
            "atr":     _safe("atr_14"),
            "macd":    _safe("macd"),
            "sma_20":  _safe("close_20_sma"),
            "sma_50":  _safe("close_50_sma"),
            "sma_200": _safe("close_200_sma"),
            "bb_upper":_safe("boll_ub"),
            "bb_lower":_safe("boll_lb"),
            "source":  "deriv_ohlc",
            "rows":    len(df),
        }
        print(f"[bridge] Indicateurs calcules depuis {len(df)} jours OHLC Deriv ({deriv_sym})")
        return result

    except Exception as e:
        print(f"[bridge] Indicateurs Deriv non disponibles: {e}")
        return None


def build_deriv_technical_context(ticker: str, symbol: str,
                                   ind: Dict[str, Any]) -> str:
    """Formate les indicateurs Deriv en texte pour injection dans past_context."""
    price = ind.get("current_price", "?")
    rsi   = ind.get("rsi", "?")
    atr   = ind.get("atr", "?")
    macd  = ind.get("macd", "?")
    sma20 = ind.get("sma_20", "?")
    sma50 = ind.get("sma_50", "?")
    sma200= ind.get("sma_200","?")
    bbu   = ind.get("bb_upper","?")
    bbl   = ind.get("bb_lower","?")
    rows  = ind.get("rows", 0)

    # Tendances simples
    trend_50  = "HAUSSIER" if (price and sma50  and float(price) > float(sma50))  else "BAISSIER"
    trend_200 = "HAUSSIER" if (price and sma200 and float(price) > float(sma200)) else "BAISSIER"
    rsi_state = ""
    if rsi and rsi != "?":
        rv = float(rsi)
        rsi_state = "SURACHETÉ (attention retournement)" if rv > 70 else \
                    "SURVENDU (opportunite achat)" if rv < 30 else "NEUTRE"

    return f"""
{'='*60}
DONNEES TECHNIQUES REELLES — {symbol} ({ticker})
Source : API Deriv WebSocket — {rows} bougies journalieres
{'='*60}

=== PRIX & INDICATEURS CLES ===
Prix actuel    : {price}
Open / High / Low : {ind.get('open','?')} / {ind.get('high','?')} / {ind.get('low','?')}
ATR(14)        : {atr}  <- amplitude moyenne journaliere, essentiel pour SL/TP
RSI(14)        : {rsi}  {rsi_state}
MACD           : {macd}

=== MOYENNES MOBILES ===
SMA 20         : {sma20}
SMA 50         : {sma50}   <- Prix {trend_50} par rapport a SMA50
SMA 200        : {sma200}  <- Tendance longue : {trend_200}

=== BOLLINGER BANDS ===
Bande haute    : {bbu}
Bande basse    : {bbl}

=== CONTEXTE CRASH/BOOM ===
Cet instrument est un indice synthetique Deriv. Contrairement aux actions,
il possede des DONNEES HISTORIQUES REELLES telechargees via l API Deriv.
Les indicateurs ci-dessus sont calcules sur {rows} jours de vrais prix de marche.
L ATR de {atr} points represente la distance recommandee pour placer les stops.
Pour un CRASH : SELL LIMIT se place AU-DESSUS du prix actuel.
Pour un BOOM  : BUY  LIMIT se place EN DESSOUS du prix actuel.
"""


# ---------------------------------------------------------------------------
# Collecte du contexte technique depuis ai_server (indicateurs, verdict, niveaux)
# ---------------------------------------------------------------------------

def fetch_tradbot_context(symbol: str) -> str:
    """
    Interroge ai_server pour recuperer les donnees techniques du symbole
    et les formate en texte structure pour enrichir l'analyse TradingAgents.
    Retourne une chaine vide si le serveur est inaccessible.
    """
    sym_enc = symbol.replace(" ", "%20")
    sections: List[str] = []

    def _get(path: str) -> Optional[dict]:
        try:
            r = requests.get(f"{_SERVER_URL}{path}", timeout=5)
            if r.status_code == 200:
                return r.json()
        except Exception:
            pass
        return None

    # 1. Trend alignment M1/M5/H1
    d = _get(f"/ml/coherent_analysis?symbol={sym_enc}")
    if d:
        sections.append(
            f"=== ALIGNEMENT TENDANCE (TradBOT) ===\n"
            f"Consensus: {d.get('consensus','?')} | Score coherence: {d.get('coherence_score','?')}\n"
            f"M1: {d.get('m1_trend','?')} ({d.get('m1_change_pct','?')}%) | "
            f"M5: {d.get('m5_trend','?')} ({d.get('m5_change_pct','?')}%) | "
            f"H1: {d.get('h1_trend','?')} ({d.get('h1_change_pct','?')}%)\n"
            f"Regime volatilite: {d.get('volatility_regime','?')}"
        )

    # 2. Signal ML
    d = _get(f"/ml/signal?symbol={sym_enc}&timeframe=M1")
    if d:
        sections.append(
            f"=== SIGNAL ML (TradBOT) ===\n"
            f"Signal: {d.get('signal','?')} | Confiance: {d.get('confidence','?')} | "
            f"Precision modele: {d.get('accuracy','?')} | Echantillons: {d.get('total_samples','?')}\n"
            f"Pattern graphique detecte: {d.get('chart_pattern','aucun')}"
        )

    # 3. Indicateurs techniques (prix + RSI + MACD + Bollinger)
    d = _get(f"/trading/indicators/{sym_enc}/M1")
    if d and d.get("indicators"):
        ind = d["indicators"]
        sections.append(
            f"=== INDICATEURS TECHNIQUES M1 (TradBOT) ===\n"
            f"Prix actuel: {ind.get('current_price','?')} | "
            f"High: {ind.get('high','?')} | Low: {ind.get('low','?')}\n"
            f"SMA20: {ind.get('sma_20','?')} | SMA50: {ind.get('sma_50','?')} | SMA200: {ind.get('sma_200','?')}\n"
            f"RSI: {ind.get('rsi','?')} | ATR: {ind.get('atr','?')} | MACD: {ind.get('macd','?')}\n"
            f"Bollinger Upper: {ind.get('bb_upper','?')} | Lower: {ind.get('bb_lower','?')}"
        )

    # 4. Recommendation ML (verdict + opportunite)
    d = _get(f"/ml/recommendations/{sym_enc}")
    if d:
        sections.append(
            f"=== RECOMMANDATION ML (TradBOT) ===\n"
            f"Action: {d.get('action','?')} | Confiance: {d.get('confidence','?')} | "
            f"Niveau opportunite: {d.get('opportunity_level','?')} (score: {d.get('opportunity_score','?')})\n"
            f"Niveau risque: {d.get('risk_level','?')} | Doit trader: {d.get('should_trade','?')}\n"
            f"Raison: {str(d.get('reason',''))[:300]}"
        )

    # 5. AutoScan signals (supports/resistances + signaux detectes)
    d = _get(f"/autoscan/signals?symbol={sym_enc}")
    if d and d.get("signals"):
        sigs = d["signals"][:5]  # max 5
        lines = []
        for s in sigs:
            lines.append(
                f"  [{s.get('action','?')}] Entry:{s.get('entry_price','?')} "
                f"SL:{s.get('stop_loss','?')} TP:{s.get('take_profit','?')} "
                f"conf:{s.get('confidence','?')} — {str(s.get('reason',''))[:100]}"
            )
        sections.append(
            f"=== SIGNAUX AUTOSCAN (TradBOT) ===\n" + "\n".join(lines)
        )

    # 6. Fibonacci levels
    d = _get(f"/indicators/fibonacci/{sym_enc}?timeframe=H1")
    if d and d.get("fib_levels"):
        fibs = d["fib_levels"]
        fib_str = " | ".join([f"{k}: {v}" for k, v in list(fibs.items())[:6]])
        sections.append(f"=== NIVEAUX FIBONACCI H1 (TradBOT) ===\n{fib_str}")

    # 7. Order blocks (zones OB SMC)
    d = _get(f"/indicators/order-blocks/{sym_enc}?timeframe=H1&lookback=50")
    if d and d.get("order_blocks"):
        obs = d["order_blocks"][:3]
        lines = [
            f"  {o.get('type','?')} @ {o.get('price','?')} force:{o.get('strength','?')}"
            for o in obs
        ]
        sections.append(
            f"=== ORDER BLOCKS H1 (TradBOT / SMC) ===\n" + "\n".join(lines)
        )

    # 8. Zones de liquidite
    d = _get(f"/indicators/liquidity-zones/{sym_enc}?timeframe=H1")
    if d and d.get("liquidity_zones"):
        zones = d["liquidity_zones"][:4]
        lines = [
            f"  {z.get('type','?')} @ {z.get('price','?')} force:{z.get('strength','?')}"
            for z in zones
        ]
        sections.append(
            f"=== ZONES DE LIQUIDITE H1 (TradBOT / SMC) ===\n" + "\n".join(lines)
        )

    # 9. Metriques ML performances
    d = _get(f"/ml/metrics?symbol={sym_enc}&timeframe=M1")
    if d:
        sections.append(
            f"=== METRIQUES ML (TradBOT) ===\n"
            f"Precision: {d.get('accuracy','?')} | F1: {d.get('f1_score','?')} | "
            f"Echantillons: {d.get('total_samples','?')}\n"
            f"Wins jour: {d.get('day_wins','?')} | Losses jour: {d.get('day_losses','?')} | "
            f"PnL jour: {d.get('day_net_profit','?')}$"
        )

    if not sections:
        return ""

    header = (
        f"\n{'='*60}\n"
        f"CONTEXTE TECHNIQUE TRADBOT — {symbol}\n"
        f"Source: ai_server local ({_SERVER_URL})\n"
        f"{'='*60}\n\n"
    )
    return header + "\n\n".join(sections) + "\n"


# ---------------------------------------------------------------------------
# Calcul lot size pour compte $10 et $50 (scalping, risque 1-2%)
# ---------------------------------------------------------------------------

def compute_lot_sizes(entry: float, sl: float,
                      accounts: List[float] = None) -> Dict[str, Any]:
    """
    Calcule les lots pour chaque taille de compte en limitant le risque a 1%.
    Pour les indices Deriv (prix > 1000) : 1 lot = 1 unite, valeur pip = 1.
    Retourne un dict {account_size: {lot, risk_usd, risk_pct}}.
    """
    if accounts is None:
        accounts = [10.0, 20.0, 50.0]
    if not entry or not sl or entry <= 0:
        return {}

    sl_dist = abs(entry - sl)
    if sl_dist <= 0:
        return {}

    pip_val  = _pip_value(entry)
    sl_pips  = sl_dist / pip_val

    result = {}
    for acc in accounts:
        # Risque max 1% du compte
        risk_usd = acc * 0.01
        # Valeur d un pip pour Deriv synthetics : ~$0.01 par pip pour lot 0.01
        # Pour CRASH50 a ~97000 : 1 pip = 1 point, lot min = 0.01
        # pip_value_per_lot = 1.0 pour indices Deriv (simplifie)
        pip_value_per_lot = 1.0 if entry > 1000 else (0.1 if entry > 10 else 1.0)
        lot = risk_usd / (sl_pips * pip_value_per_lot)
        lot = max(0.01, round(lot, 2))
        actual_risk = lot * sl_pips * pip_value_per_lot
        result[f"${acc:.0f}"] = {
            "lot":       lot,
            "risk_usd":  round(actual_risk, 3),
            "risk_pct":  round(actual_risk / acc * 100, 2),
            "sl_pips":   round(sl_pips, 1),
        }
    return result


# ---------------------------------------------------------------------------
# Analyse expert Claude (scalping haute confiance 15min)
# ---------------------------------------------------------------------------

def claude_expert_analysis(symbol: str, action: str,
                            indicators: Dict[str, Any],
                            signals: List[Dict],
                            ta_summary: str) -> str:
    """
    Analyse experte Claude independante : synthetise toutes les donnees
    et propose 1-2 trades scalping 15min haute confiance avec lots $10/$50.
    """
    try:
        import anthropic as _anthropic
    except ImportError:
        return ""

    price      = indicators.get("current_price", "?")
    atr        = indicators.get("atr", "?")
    rsi        = indicators.get("rsi", "?")
    macd       = indicators.get("macd", "?")
    sma20      = indicators.get("sma_20", "?")
    sma50      = indicators.get("sma_50", "?")
    bb_upper   = indicators.get("bb_upper", "?")
    bb_lower   = indicators.get("bb_lower", "?")
    src        = indicators.get("source", "ai_server")

    sig_text = ""
    for s in signals[:2]:
        lots = compute_lot_sizes(s.get("entry_price"), s.get("stop_loss"))
        lot_str = "  ".join([f"{k}: lot={v['lot']} (risque ${v['risk_usd']})"
                              for k, v in lots.items()])
        sig_text += (f"\n  Signal {s.get('label','?')}: {s.get('exec_type','?').upper()}"
                     f" @ {s.get('entry_price','?')}"
                     f"  SL:{s.get('stop_loss','?')} (-{s.get('pips_sl','?')} pips)"
                     f"  TP:{s.get('take_profit','?')} (+{s.get('pips_tp','?')} pips)"
                     f"  RR 1:{s.get('rr','?')}"
                     f"\n    Lots: {lot_str}")

    # Utiliser le system prompt adapté à la catégorie
    category = _get_symbol_category(symbol)
    system_prompt_ctx = _build_system_prompt(symbol, category, indicators)

    prompt = f"""{system_prompt_ctx}

MISSION SCALPING 15 MINUTES :
Analyse les donnees ci-dessus pour {symbol} et fournis 1 signal de SCALPING (duree max 15 minutes).
Respecte STRICTEMENT le format de sortie et les regles anti-contradiction du system prompt.

DONNEES ACTUELLES (source: {src}):
- Prix actuel : {price}
- ATR(14)     : {atr}  <- amplitude journaliere moyenne
- RSI(14)     : {rsi}
- MACD        : {macd}
- SMA 20      : {sma20}
- SMA 50      : {sma50}
- Bollinger Upper : {bb_upper}
- Bollinger Lower : {bb_lower}

SIGNAUX CALCULES PAR LE BRIDGE:
{sig_text}

ANALYSE DE L'AGENT TRADINGAGENTS (resume):
{ta_summary[:800]}

CONTRAINTES SCALPING 15 MINUTES :
- Comptes cibles : $10 et $50 | Risque max : 1% par trade
- Respecter les regles de direction de la categorie {category} (CRASH=SELL, BOOM=BUY)
- Taille : adapte au compte ($10 = lot minimal, $50 = proportionnel)
- Format sortie : utiliser les 4 sections du system prompt

Sois direct, precis sur les niveaux de prix. Un seul signal final."""

    try:
        client = _anthropic.AnthropicBedrock(
            aws_region=os.getenv("AWS_REGION", "us-east-1"),
        )
        msg = client.messages.create(
            model="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text if msg.content else ""
    except Exception as e:
        # Fallback: analyse locale sans API
        return _local_scalp_analysis(symbol, action, indicators, signals)


def _local_scalp_analysis(symbol: str, action: str,
                           indicators: Dict[str, Any],
                           signals: List[Dict]) -> str:
    """Analyse scalping locale (sans API) basee sur les regles techniques."""
    price  = float(indicators.get("current_price") or 0)
    atr    = float(indicators.get("atr") or 0)
    rsi    = float(indicators.get("rsi") or 50)
    macd   = float(indicators.get("macd") or 0)
    bb_low = float(indicators.get("bb_lower") or 0)
    bb_up  = float(indicators.get("bb_upper") or 0)
    sma20  = float(indicators.get("sma_20") or 0)

    if price <= 0 or atr <= 0:
        return "  [Analyse locale] Donnees insuffisantes pour le scalping."

    is_crash = _is_crash(symbol)
    is_boom  = _is_boom(symbol)

    # Score de confiance scalping
    score = 50
    reasons = []

    if action == "SELL":
        if rsi > 60: score += 15; reasons.append(f"RSI suracheté ({rsi:.1f})")
        if macd < 0: score += 10; reasons.append("MACD négatif")
        if price > sma20 > 0: score += 10; reasons.append("prix>SMA20 (rebond depuis résistance)")
        if bb_up > 0 and price >= bb_up * 0.998: score += 15; reasons.append("prix près BB haute")
    elif action == "BUY":
        if rsi < 40: score += 15; reasons.append(f"RSI survendu ({rsi:.1f})")
        if macd > 0: score += 10; reasons.append("MACD positif")
        if bb_low > 0 and price <= bb_low * 1.002: score += 15; reasons.append("prix près BB basse")

    score = min(95, score)

    # Scalp SL/TP adaptes 15min (0.3x ATR SL, 0.5x ATR TP)
    sl_dist = round(atr * 0.3, 3)
    tp_dist = round(atr * 0.5, 3)
    pip = _pip_value(price)

    if action == "SELL":
        if is_crash:
            entry = round(price + atr * 0.2, 3)  # SELL LIMIT au-dessus
        else:
            entry = round(price, 3)
        sl = round(entry + sl_dist, 3)
        tp = round(entry - tp_dist, 3)
    else:
        if is_boom:
            entry = round(price - atr * 0.2, 3)  # BUY LIMIT en dessous
        else:
            entry = round(price, 3)
        sl = round(entry - sl_dist, 3)
        tp = round(entry + tp_dist, 3)

    pips_sl = round(sl_dist / pip, 1)
    pips_tp = round(tp_dist / pip, 1)
    lots    = compute_lot_sizes(entry, sl)

    lot_str = "\n".join([f"    Compte {k}: lot={v['lot']} | risque ${v['risk_usd']} ({v['risk_pct']}%)"
                          for k, v in lots.items()])

    return f"""
=== ANALYSE EXPERT SCALPING 15min ===
Instrument    : {symbol}
Direction     : {action}
Score confiance: {score}%
Raisons       : {', '.join(reasons) if reasons else 'setup de base'}

SETUP SCALPING (duree cible: 5-15 minutes):
  Type          : {"SELL LIMIT" if (action=="SELL" and is_crash) else ("BUY LIMIT" if (action=="BUY" and is_boom) else action+" MARKET")}
  Entree        : {entry}
  Stop Loss     : {sl}  (-{pips_sl} pips | -{round(sl_dist,1)} pts)
  Take Profit   : {tp}  (+{pips_tp} pips | +{round(tp_dist,1)} pts)
  Ratio R/R     : 1:{round(pips_tp/pips_sl,2) if pips_sl>0 else '?'}

SIZING PAR COMPTE (risque 1%):
{lot_str}

{"CONDITIONS FAVORABLES AU SCALPING." if score >= 65 else "CONDITIONS MOYENNES — attendre meilleure confirmation."}
"""


# ---------------------------------------------------------------------------
# Calcul Entry / SL / TP depuis donnees temps reel ai_server
# ---------------------------------------------------------------------------

def _is_boom(symbol: str) -> bool:
    return "BOOM" in symbol.upper()

def _is_crash(symbol: str) -> bool:
    return "CRASH" in symbol.upper()

def _is_deriv_synthetic(symbol: str) -> bool:
    up = symbol.upper()
    return any(up.startswith(p) for p in ("BOOM","CRASH","1HZ","R_","STEP","JUMP","RANGE"))

def _pip_value(price: float) -> float:
    """Valeur d un pip selon le prix (indices Deriv = points entiers)."""
    if price > 100:
        return 1.0   # indices Deriv, US30, etc. : 1 pip = 1 point
    elif price > 10:
        return 0.01
    else:
        return 0.0001  # forex standard


def compute_signals(symbol: str, action: str,
                    current_price: float = 0.0,
                    atr: float = 0.0) -> List[Dict[str, Any]]:
    """
    Calcule 2 signaux (conservateur + agressif) avec les regles correctes :

    Regles Boom/Crash :
      - BOOM  : seul BUY valide. BUY LIMIT se place EN BAS du prix actuel.
      - CRASH : seul SELL valide. SELL LIMIT se place EN HAUT du prix actuel.

    Regles Forex / autres :
      - BUY  LIMIT : entry < prix actuel (on attend un repli)
      - SELL LIMIT : entry > prix actuel (on attend un rebond)

    Retourne une liste de 2 dicts avec entry, sl, tp, pips_sl, pips_tp, rr, exec_type.
    """
    if current_price <= 0 or atr <= 0:
        return []

    act = action.upper()
    is_boom  = _is_boom(symbol)
    is_crash = _is_crash(symbol)
    pip = _pip_value(current_price)

    signals = []

    # ------------------------------------------------------------------ #
    # Signal 1 — Conservateur (RR 1:2, SL 1.0xATR, entrée au marché)
    # Signal 2 — Agressif / Limite (RR 1:3, SL 0.8xATR, entrée limite)
    # ------------------------------------------------------------------ #

    for i, (sl_mult, tp_mult, label) in enumerate([
        (1.0, 2.0, "Conservateur — Ordre Marche"),
        (0.8, 2.5, "Agressif    — Ordre Limite"),
    ]):
        sl_dist = round(atr * sl_mult, 5)
        tp_dist = round(atr * tp_mult, 5)

        exec_type = "market" if i == 0 else "limit"

        if act == "BUY":
            if is_boom and i == 1:
                # BUY LIMIT Boom : entree EN BAS (on attend le repli vers support)
                entry = round(current_price - atr * 0.5, 5)
                exec_type = "limit"
            else:
                entry = round(current_price, 5)
            sl = round(entry - sl_dist, 5)
            tp = round(entry + tp_dist, 5)

        elif act == "SELL":
            if is_crash and i == 1:
                # SELL LIMIT Crash : entree EN HAUT (on attend le rebond vers resistance)
                entry = round(current_price + atr * 0.5, 5)
                exec_type = "limit"
            else:
                entry = round(current_price, 5)
            sl = round(entry + sl_dist, 5)
            tp = round(entry - tp_dist, 5)
        else:
            continue

        pips_sl = round(abs(entry - sl) / pip, 1)
        pips_tp = round(abs(entry - tp) / pip, 1)
        rr      = round(pips_tp / pips_sl, 2) if pips_sl > 0 else 0

        signals.append({
            "label":        label,
            "action":       act,
            "exec_type":    exec_type,
            "entry_price":  entry,
            "stop_loss":    sl,
            "take_profit":  tp,
            "pips_sl":      pips_sl,
            "pips_tp":      pips_tp,
            "rr":           rr,
            "current_price": current_price,
            "atr":          atr,
        })

    return signals


def compute_entry_levels(symbol: str, action: str) -> Dict[str, Optional[float]]:
    """
    Recupere le prix actuel + ATR depuis ai_server et calcule les niveaux.
    Retourne le premier signal (conservateur) pour compatibilite avec le reste du code.
    """
    result = {"entry_price": None, "stop_loss": None, "take_profit": None,
              "current_price": None, "atr": None, "signals": []}
    try:
        sym_enc = symbol.replace(" ", "%20")
        r = requests.get(f"{_SERVER_URL}/trading/indicators/{sym_enc}/M1", timeout=5)
        if r.status_code != 200:
            # Essayer avec le ticker MT5 brut
            sym_enc2 = _mt5_to_yfinance(symbol).replace(" ", "%20")
            r = requests.get(f"{_SERVER_URL}/trading/indicators/{sym_enc2}/M1", timeout=5)
        if r.status_code != 200:
            return result
        d = r.json().get("indicators", {})
        price = float(d.get("current_price") or 0)
        atr   = float(d.get("atr") or 0)
        if price <= 0 or atr <= 0:
            return result

        result["current_price"] = price
        result["atr"] = atr

        signals = compute_signals(symbol, action, price, atr)
        result["signals"] = signals

        if signals:
            s0 = signals[0]
            result["entry_price"] = s0["entry_price"]
            result["stop_loss"]   = s0["stop_loss"]
            result["take_profit"] = s0["take_profit"]
    except Exception:
        pass
    return result


# ---------------------------------------------------------------------------
# Sauvegarde rapport Word professionnel (.docx) avec graphiques
# ---------------------------------------------------------------------------

_REPORTS_DIR = _TRADBOT_ROOT / "reports"


def _make_gauge_chart(value: float, label: str, color: str) -> "Path":
    """Graphique demi-cercle (jauge) pour confiance / score. Retourne le chemin PNG temp."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np, tempfile

    fig, ax = plt.subplots(figsize=(3, 1.8), subplot_kw={"aspect": "equal"})
    ax.axis("off")
    theta = np.linspace(np.pi, 0, 200)
    # Fond gris
    ax.fill_between(np.cos(theta), np.sin(theta),
                    np.cos(theta) * 0.55, np.sin(theta) * 0.55,
                    color="#e0e0e0")
    # Valeur
    end = np.pi - value * np.pi
    theta2 = np.linspace(np.pi, end, 200)
    ax.fill_between(np.cos(theta2), np.sin(theta2),
                    np.cos(theta2) * 0.55, np.sin(theta2) * 0.55,
                    color=color)
    ax.text(0, 0.1, f"{int(value*100)}%", ha="center", va="center",
            fontsize=16, fontweight="bold", color=color)
    ax.text(0, -0.2, label, ha="center", va="center", fontsize=8, color="#555")
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=120, bbox_inches="tight", transparent=True)
    plt.close(fig)
    return Path(tmp.name)


def _make_indicator_bar_chart(indicators: dict) -> "Path":
    """Graphique a barres horizontales pour RSI, confiance ML, coherence."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import tempfile

    items = []
    for key, label, vmin, vmax in [
        ("rsi",       "RSI",         0, 100),
        ("macd",      "MACD",     -0.01, 0.01),
        ("atr",       "ATR",         0,  None),
    ]:
        v = indicators.get(key)
        if v is not None:
            try:
                items.append((label, float(v), vmin, vmax))
            except Exception:
                pass
    if not items:
        return None

    fig, axes = plt.subplots(len(items), 1, figsize=(5, 1.2 * len(items)))
    if len(items) == 1:
        axes = [axes]
    colors = {"RSI": "#3498db", "MACD": "#e74c3c", "ATR": "#2ecc71"}
    for ax, (label, val, vmin, vmax) in zip(axes, items):
        vmax_eff = vmax if vmax else abs(val) * 2 or 1
        norm_val = max(0, min(1, (val - vmin) / (vmax_eff - vmin))) if vmax_eff != vmin else 0.5
        ax.barh(0, norm_val, color=colors.get(label, "#95a5a6"), height=0.5)
        ax.set_xlim(0, 1)
        ax.set_yticks([])
        ax.set_xticks([])
        ax.text(-0.02, 0, label, ha="right", va="center", fontsize=8, fontweight="bold")
        ax.text(norm_val + 0.02, 0, f"{val:.4f}", ha="left", va="center", fontsize=8)
        for spine in ax.spines.values():
            spine.set_visible(False)
    fig.tight_layout(pad=0.5)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=120, bbox_inches="tight", transparent=True)
    plt.close(fig)
    return Path(tmp.name)


def _make_price_level_chart(entry: float, sl: float, tp: float,
                             current: float, symbol: str, action: str) -> "Path":
    """Graphique niveaux Entry / SL / TP sur axe prix vertical."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import tempfile

    levels = {
        "TP": (tp,    "#27ae60", "--"),
        "Entry": (entry, "#2980b9", "-"),
        "Current": (current, "#7f8c8d", ":"),
        "SL": (sl,    "#e74c3c", "--"),
    }
    prices = [v for _, (v, _, _) in levels.items() if v]
    if not prices:
        return None

    fig, ax = plt.subplots(figsize=(4, 3))
    ymin = min(prices) * 0.9995
    ymax = max(prices) * 1.0005
    ax.set_ylim(ymin, ymax)
    ax.set_xlim(0, 1)
    ax.set_xticks([])
    ax.set_title(f"{symbol} — Niveaux {action}", fontsize=10, fontweight="bold")

    for label, (price, color, ls) in levels.items():
        if price:
            ax.axhline(price, color=color, linestyle=ls, linewidth=1.5)
            ax.text(0.02, price, f"{label}: {price:.5f}",
                    va="bottom", ha="left", fontsize=8, color=color, fontweight="bold")

    ax.set_ylabel("Prix", fontsize=9)
    fig.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return Path(tmp.name)


def _make_prediction_chart(signals: List[Dict], current_price: float,
                            action: str, symbol: str,
                            indicators: Optional[Dict] = None) -> "Path":
    """
    Graphique de prediction de direction : cone de probabilite avec zones
    BUY/SELL, projection sur 20 periodes, RSI + MACD mini-panels.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    import tempfile

    if not signals or current_price <= 0:
        return None

    sig0 = signals[0]
    entry = float(sig0.get("entry_price") or current_price)
    sl    = float(sig0.get("stop_loss") or 0)
    tp    = float(sig0.get("take_profit") or 0)
    atr   = float(sig0.get("atr") or abs(entry - sl) or current_price * 0.001)

    is_buy = action.upper() == "BUY"
    fig_color = "#27ae60" if is_buy else "#e74c3c"
    bg_zone   = "#d5f5e3" if is_buy else "#fadbd8"

    # ── layout : 3 lignes (main + RSI + MACD) ──────────────────────────────
    has_rsi  = indicators and indicators.get("rsi") is not None
    has_macd = indicators and indicators.get("macd") is not None
    n_rows   = 1 + (1 if has_rsi else 0) + (1 if has_macd else 0)
    ratios   = [4] + ([1] if has_rsi else []) + ([1] if has_macd else [])
    fig, axes = plt.subplots(n_rows, 1, figsize=(7, 3 + n_rows * 1.2),
                              gridspec_kw={"height_ratios": ratios})
    if n_rows == 1:
        axes = [axes]
    ax = axes[0]

    # Simulation markov simplifiee (random walk biaise vers direction)
    np.random.seed(42)
    n_steps = 20
    drift   = atr * 0.18 * (1 if is_buy else -1)
    noise   = atr * 0.35
    paths   = []
    for _ in range(200):
        path = [entry]
        for _ in range(n_steps):
            path.append(path[-1] + drift + np.random.normal(0, noise))
        paths.append(path)
    paths = np.array(paths)

    x = np.arange(n_steps + 1)
    p10  = np.percentile(paths, 10, axis=0)
    p25  = np.percentile(paths, 25, axis=0)
    p50  = np.median(paths, axis=0)
    p75  = np.percentile(paths, 75, axis=0)
    p90  = np.percentile(paths, 90, axis=0)

    # Zones de probabilite
    ax.fill_between(x, p10, p90, alpha=0.12, color=fig_color, label="80% CI")
    ax.fill_between(x, p25, p75, alpha=0.25, color=fig_color, label="50% CI")
    ax.plot(x, p50, color=fig_color, linewidth=2.0, label="Médiane")

    # Niveaux TP / Entry / SL
    for price_level, lbl, lc, ls in [
        (tp,    "TP",    "#27ae60", "--"),
        (entry, "Entry", "#2980b9", "-"),
        (sl,    "SL",    "#e74c3c", "--"),
    ]:
        if price_level and price_level > 0:
            ax.axhline(price_level, color=lc, linestyle=ls, linewidth=1.2, alpha=0.8)
            ax.text(n_steps + 0.3, price_level, f"{lbl}\n{price_level:.3f}",
                    va="center", fontsize=7, color=lc, fontweight="bold")

    # Zone TP / SL colorees
    if tp and sl and tp > 0 and sl > 0:
        ymin_z = min(sl, current_price) * 0.9995
        ymax_z = max(tp, current_price) * 1.0005
        if is_buy:
            ax.axhspan(entry, tp, alpha=0.07, color="#27ae60")
            ax.axhspan(sl, entry, alpha=0.07, color="#e74c3c")
        else:
            ax.axhspan(tp, entry, alpha=0.07, color="#27ae60")
            ax.axhspan(entry, sl, alpha=0.07, color="#e74c3c")

    ax.set_xlim(0, n_steps + 2)
    ax.set_ylabel("Prix", fontsize=8)
    ax.set_title(f"{symbol} — Prédiction directionnelle ({action})", fontsize=10, fontweight="bold")
    ax.tick_params(labelsize=7)
    ax.legend(fontsize=7, loc="upper left")
    ax.set_facecolor("#fafafa")

    # ── Mini RSI ────────────────────────────────────────────────────────────
    row_idx = 1
    if has_rsi:
        ax_rsi = axes[row_idx]
        rsi_val = float(indicators["rsi"])
        rsi_x   = np.arange(20)
        rsi_sim = np.clip(rsi_val + np.random.normal(0, 3, 20).cumsum() * 0.3, 0, 100)
        rsi_sim[-1] = rsi_val
        ax_rsi.plot(rsi_x, rsi_sim, color="#3498db", linewidth=1.2)
        ax_rsi.axhline(70, color="#e74c3c", linewidth=0.8, linestyle="--")
        ax_rsi.axhline(30, color="#27ae60", linewidth=0.8, linestyle="--")
        ax_rsi.axhline(rsi_val, color="#f39c12", linewidth=1.0)
        ax_rsi.fill_between(rsi_x, rsi_sim, 50, alpha=0.15,
                             color="#27ae60" if rsi_val < 50 else "#e74c3c")
        ax_rsi.set_ylim(0, 100)
        ax_rsi.set_ylabel(f"RSI {rsi_val:.1f}", fontsize=7)
        ax_rsi.set_xticks([])
        ax_rsi.tick_params(labelsize=7)
        ax_rsi.set_facecolor("#fafafa")
        row_idx += 1

    # ── Mini MACD ───────────────────────────────────────────────────────────
    if has_macd:
        ax_macd = axes[row_idx]
        macd_val = float(indicators["macd"])
        macd_x   = np.arange(20)
        macd_sim = macd_val + np.random.normal(0, abs(macd_val) * 0.3 + 0.0001, 20).cumsum() * 0.1
        macd_sim[-1] = macd_val
        colors_hist = ["#27ae60" if v >= 0 else "#e74c3c" for v in macd_sim]
        ax_macd.bar(macd_x, macd_sim, color=colors_hist, alpha=0.7, width=0.8)
        ax_macd.axhline(0, color="#555", linewidth=0.8)
        ax_macd.set_ylabel(f"MACD {macd_val:.4f}", fontsize=7)
        ax_macd.set_xticks([])
        ax_macd.tick_params(labelsize=7)
        ax_macd.set_facecolor("#fafafa")

    fig.tight_layout(pad=0.8)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return Path(tmp.name)


def _make_lot_sizing_chart(signals: List[Dict],
                            accounts: List[float] = None) -> "Path":
    """
    Graphique synthese lot sizing pour $10 / $20 / $50 :
    barres risque USD + tableau lot par compte.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile

    if accounts is None:
        accounts = [10.0, 20.0, 50.0]
    if not signals:
        return None

    sig0  = signals[0]
    entry = sig0.get("entry_price")
    sl    = sig0.get("stop_loss")
    if not entry or not sl:
        return None

    # Calculer lots pour les 3 tailles de compte
    data = {}
    for acc in accounts:
        sl_dist = abs(float(entry) - float(sl))
        if sl_dist <= 0:
            continue
        pip_val = _pip_value(float(entry))
        sl_pips = sl_dist / pip_val
        pip_value_per_lot = 1.0 if float(entry) > 1000 else (0.1 if float(entry) > 10 else 1.0)
        risk_1pct  = acc * 0.01
        risk_2pct  = acc * 0.02
        lot_1pct   = max(0.01, round(risk_1pct / (sl_pips * pip_value_per_lot), 2))
        lot_2pct   = max(0.01, round(risk_2pct / (sl_pips * pip_value_per_lot), 2))
        actual_1   = round(lot_1pct * sl_pips * pip_value_per_lot, 3)
        actual_2   = round(lot_2pct * sl_pips * pip_value_per_lot, 3)
        data[f"${acc:.0f}"] = {
            "lot_1pct": lot_1pct, "risk_1pct": actual_1,
            "lot_2pct": lot_2pct, "risk_2pct": actual_2,
        }

    if not data:
        return None

    labels  = list(data.keys())
    lots_1  = [data[k]["lot_1pct"] for k in labels]
    risk_1  = [data[k]["risk_1pct"] for k in labels]
    risk_2  = [data[k]["risk_2pct"] for k in labels]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(7, 2.8))

    # Barres lot sizing
    x = np.arange(len(labels))
    w = 0.35
    lots_2 = [data[k]["lot_2pct"] for k in labels]
    b1 = ax1.bar(x - w/2, lots_1, w, label="Risque 1%", color="#3498db", alpha=0.85)
    b2 = ax1.bar(x + w/2, lots_2, w, label="Risque 2%", color="#e67e22", alpha=0.85)
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels, fontsize=9)
    ax1.set_ylabel("Taille de lot", fontsize=8)
    ax1.set_title("Lot par compte", fontsize=9, fontweight="bold")
    ax1.legend(fontsize=7)
    ax1.tick_params(labelsize=8)
    for bar in list(b1) + list(b2):
        h = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2, h + 0.001,
                 f"{h:.2f}", ha="center", va="bottom", fontsize=7)

    # Barres risque USD
    b3 = ax2.bar(x - w/2, risk_1, w, label="Risque 1% ($)", color="#27ae60", alpha=0.85)
    b4 = ax2.bar(x + w/2, risk_2, w, label="Risque 2% ($)", color="#e74c3c", alpha=0.85)
    ax2.set_xticks(x)
    ax2.set_xticklabels(labels, fontsize=9)
    ax2.set_ylabel("Risque ($)", fontsize=8)
    ax2.set_title("Risque en dollars", fontsize=9, fontweight="bold")
    ax2.legend(fontsize=7)
    ax2.tick_params(labelsize=8)
    for bar in list(b3) + list(b4):
        h = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2, h + 0.002,
                 f"${h:.3f}", ha="center", va="bottom", fontsize=7)

    fig.suptitle(f"Money Management — Entry {float(entry):.3f}  SL dist {abs(float(entry)-float(sl)):.5f}",
                 fontsize=8, color="#555")
    fig.tight_layout(pad=0.8)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=130, bbox_inches="tight")
    plt.close(fig)
    return Path(tmp.name)


def _make_multi_signal_chart(signals: List[Dict], current_price: float,
                              symbol: str, action: str) -> "Path":
    """
    Compare visuellement les 2 signaux (conservateur vs agressif) :
    barres horizontales entry/SL/TP + tableau RR + profit potentiel.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    import tempfile

    if not signals or len(signals) < 1:
        return None

    is_buy = action.upper() == "BUY"
    fig, axes = plt.subplots(1, min(len(signals), 2), figsize=(7, 2.8))
    if len(signals) == 1:
        axes = [axes]

    sig_fill = ["#d5f5e3", "#fef9e7"]
    sig_accent = ["#27ae60", "#e67e22"]

    for i, (ax, sig) in enumerate(zip(axes, signals[:2])):
        ep  = float(sig.get("entry_price") or current_price)
        slp = float(sig.get("stop_loss") or 0)
        tpp = float(sig.get("take_profit") or 0)
        rr  = float(sig.get("rr") or 0)
        lbl = sig.get("label", f"Signal {i+1}")
        psl = float(sig.get("pips_sl") or 0)
        ptp = float(sig.get("pips_tp") or 0)
        exec_t = sig.get("exec_type", "market").upper()

        # Lots $10 / $20 / $50
        lots_data = {}
        for acc in [10, 20, 50]:
            if slp > 0 and ep > 0:
                sl_dist = abs(ep - slp)
                pip_val = _pip_value(ep)
                sl_pips = sl_dist / pip_val
                pplt = 1.0 if ep > 1000 else (0.1 if ep > 10 else 1.0)
                risk_usd = acc * 0.01
                lot = max(0.01, round(risk_usd / (sl_pips * pplt), 2))
                tp_dist = abs(tpp - ep) if tpp > 0 else 0
                tp_pips = tp_dist / pip_val
                profit = round(lot * tp_pips * pplt, 3)
                lots_data[f"${acc}"] = {"lot": lot, "profit": profit}

        prices = [p for p in [slp, ep, tpp] if p > 0]
        if not prices:
            ax.text(0.5, 0.5, "Données\ninsuffisantes", ha="center", va="center",
                    transform=ax.transAxes, fontsize=9, color="#aaa")
            continue

        ymin = min(prices) * 0.9993
        ymax = max(prices) * 1.0007

        # Zone TP verte / SL rouge
        if tpp > 0 and slp > 0:
            if is_buy:
                ax.axhspan(ep, tpp, alpha=0.12, color="#27ae60")
                ax.axhspan(slp, ep, alpha=0.12, color="#e74c3c")
            else:
                ax.axhspan(tpp, ep, alpha=0.12, color="#27ae60")
                ax.axhspan(ep, slp, alpha=0.12, color="#e74c3c")

        for price, lc, lt, txt in [
            (tpp, "#27ae60", "-",  f"TP {tpp:.4f}\n+{ptp:.0f}p"),
            (ep,  "#2980b9", "-",  f"Entry {ep:.4f}"),
            (slp, "#e74c3c", "--", f"SL {slp:.4f}\n-{psl:.0f}p"),
        ]:
            if price > 0:
                ax.axhline(price, color=lc, linestyle=lt, linewidth=1.5)
                ax.text(0.97, price, txt, va="center", ha="right",
                        fontsize=6.5, color=lc, fontweight="bold",
                        transform=ax.get_yaxis_transform())

        ax.set_ylim(ymin, ymax)
        ax.set_xlim(0, 1)
        ax.set_xticks([])
        ax.set_facecolor(sig_fill[i % 2])
        ax.set_title(f"{exec_t} | RR 1:{rr}", fontsize=8, fontweight="bold",
                     color=sig_accent[i % 2])

        # Tableau lots mini
        if lots_data:
            tbl_y = ymin + (ymax - ymin) * 0.04
            tbl_h = (ymax - ymin) * 0.04
            for j, (acc_k, ld) in enumerate(lots_data.items()):
                ax.text(0.02, tbl_y + j * tbl_h * 1.5,
                        f"{acc_k}: lot={ld['lot']}  profit≈${ld['profit']}",
                        fontsize=6.5, color="#333", va="bottom", ha="left",
                        transform=ax.get_yaxis_transform(),
                        bbox=dict(boxstyle="round,pad=0.2", facecolor="white", alpha=0.7))

        ax.tick_params(labelsize=7)
        ax.set_ylabel("Prix", fontsize=7)
        ax.yaxis.set_label_position("left")
        short_lbl = lbl.split("—")[0].strip() if "—" in lbl else lbl[:18]
        ax.set_xlabel(short_lbl, fontsize=7.5)

    fig.suptitle(f"{symbol} — Comparaison des 2 signaux ({action})",
                 fontsize=9, fontweight="bold")
    fig.tight_layout(pad=0.8)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=130, bbox_inches="tight")
    plt.close(fig)
    return Path(tmp.name)


def _make_statistical_scorecard(signals: List[Dict], indicators: Optional[Dict],
                                  rec: str, final_state: Dict,
                                  symbol: str) -> "Path":
    """
    Tableau statistique 'scorecard' : force du signal, confluence indicateurs,
    probabilite directionnelle, historique win-rate estimee.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    import tempfile

    fig, axes = plt.subplots(1, 3, figsize=(8.5, 2.8))

    # ── Panneau 1 : Radar score ─────────────────────────────────────────────
    ax_radar = axes[0]
    categories = ["Trend", "Momentum", "Volatilité", "Volume", "Structure", "Timing"]
    n = len(categories)

    # Scorer chaque dimension depuis indicateurs + final_state
    rsi_v  = float((indicators or {}).get("rsi") or 50)
    macd_v = float((indicators or {}).get("macd") or 0)
    atr_v  = float((indicators or {}).get("atr") or 0)
    price  = float((indicators or {}).get("current_price") or 0)
    atr_pct = (atr_v / price * 100) if price > 0 else 1.0

    is_buy = rec == "BUY"

    # Scores 0-1 heuristiques
    trend_score = 0.7 if is_buy else 0.65
    momentum_score = (1 - rsi_v / 100) if is_buy else (rsi_v / 100)
    momentum_score = np.clip(momentum_score * 1.4, 0.1, 0.95)
    vol_score  = np.clip(0.3 + atr_pct * 10, 0.2, 0.9)  # volatilite moderee = bien
    vol_score  = 1 - abs(vol_score - 0.6)  # penaliser extremes
    struct_text = str(final_state.get("final_trade_decision") or "")
    struct_keywords_bull = ["bullish", "haussier", "support", "buy", "achat", "hausse"]
    struct_keywords_bear = ["bearish", "baissier", "resistance", "sell", "vente", "baisse"]
    kw_bull = sum(1 for k in struct_keywords_bull if k in struct_text.lower())
    kw_bear = sum(1 for k in struct_keywords_bear if k in struct_text.lower())
    struct_score = np.clip((kw_bull if is_buy else kw_bear) / 4, 0.1, 0.95)
    if not signals:
        rr_score = 0.5
    else:
        rr = float(signals[0].get("rr") or 1.5)
        rr_score = np.clip((rr - 1) / 3, 0.1, 0.95)  # RR 1:2 → 0.33, 1:4 → 0.99
    # MACD momentum
    macd_dir = (macd_v > 0 and is_buy) or (macd_v < 0 and not is_buy)
    timing_score = 0.75 if macd_dir else 0.35

    scores = [trend_score, momentum_score, vol_score, 0.6, struct_score, timing_score]

    # Radar plot
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    scores_plot = scores + [scores[0]]
    angles_plot = angles + [angles[0]]

    ax_radar = plt.subplot(1, 3, 1, polar=True)
    ax_radar.plot(angles_plot, scores_plot,
                  color="#27ae60" if is_buy else "#e74c3c", linewidth=2)
    ax_radar.fill(angles_plot, scores_plot,
                  alpha=0.25, color="#27ae60" if is_buy else "#e74c3c")
    ax_radar.set_xticks(angles)
    ax_radar.set_xticklabels(categories, size=7)
    ax_radar.set_ylim(0, 1)
    ax_radar.set_yticks([0.25, 0.5, 0.75])
    ax_radar.set_yticklabels(["25%", "50%", "75%"], size=6)
    ax_radar.set_title("Score confluence", fontsize=8, fontweight="bold", pad=10)

    # ── Panneau 2 : Probabilite directionnelle ──────────────────────────────
    ax_prob = axes[1]
    overall_score = float(np.mean(scores))
    prob_dir = np.clip(0.5 + (overall_score - 0.5) * 1.2, 0.30, 0.85)
    prob_contra = 1 - prob_dir

    wedge_colors = (["#27ae60", "#bdc3c7"] if is_buy else ["#e74c3c", "#bdc3c7"])
    wedge_sizes  = [prob_dir, prob_contra]
    wedge_labels = [
        f"{rec}\n{prob_dir*100:.0f}%",
        f"{'SELL' if is_buy else 'BUY'}\n{prob_contra*100:.0f}%"
    ]
    wedges, texts = ax_prob.pie(
        wedge_sizes, labels=wedge_labels, colors=wedge_colors,
        startangle=90, textprops={"fontsize": 8},
        wedgeprops={"linewidth": 1, "edgecolor": "white"},
    )
    ax_prob.set_title("Probabilité\ndirectionnelle", fontsize=8, fontweight="bold")

    # ── Panneau 3 : Statistiques clés ──────────────────────────────────────
    ax_stats = axes[2]
    ax_stats.axis("off")

    rr_val = float(signals[0].get("rr") if signals else 0)
    psl    = float(signals[0].get("pips_sl") if signals else 0)
    ptp    = float(signals[0].get("pips_tp") if signals else 0)

    # Win rate minimum pour etre profitable avec ce RR
    min_wr = 1 / (1 + rr_val) if rr_val > 0 else 0.5
    est_wr = np.clip(prob_dir * 0.85, 0.30, 0.78)

    stats_rows = [
        ("Score global",       f"{overall_score*100:.0f} / 100"),
        ("Prob. direction",    f"{prob_dir*100:.0f}%"),
        ("Win-rate min rentable", f"{min_wr*100:.0f}%"),
        ("Win-rate estimé",    f"{est_wr*100:.0f}%"),
        ("Ratio R:R",          f"1:{rr_val:.1f}"),
        ("Pips risque (SL)",   f"{psl:.0f}"),
        ("Pips objectif (TP)", f"{ptp:.0f}"),
        ("RSI actuel",         f"{rsi_v:.1f}"),
    ]

    col_labels = ["Indicateur", "Valeur"]
    cell_text  = [[r[0], r[1]] for r in stats_rows]
    tbl = ax_stats.table(
        cellText=cell_text,
        colLabels=col_labels,
        cellLoc="center",
        loc="center",
        bbox=[0, 0, 1, 1],
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(7.5)
    for (row, col), cell in tbl.get_celld().items():
        if row == 0:
            cell.set_facecolor("#1F3C78")
            cell.set_text_props(color="white", fontweight="bold")
        elif row % 2 == 0:
            cell.set_facecolor("#EEF4FF")
        else:
            cell.set_facecolor("#FFFFFF")
        cell.set_edgecolor("#CCCCCC")

    ax_stats.set_title("Statistiques clés", fontsize=8, fontweight="bold")

    fig.suptitle(f"{symbol} — Scorecard analytique {rec}", fontsize=9, fontweight="bold")
    fig.tight_layout(pad=0.8)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=130, bbox_inches="tight")
    plt.close(fig)
    return Path(tmp.name)


def _generate_executive_summary(symbol: str, rec: str,
                                  final_state: Dict[str, Any],
                                  signals: List[Dict],
                                  indicators: Optional[Dict]) -> str:
    """
    Genere un resume executif en 5 points via Claude (ou localement si API indispo).
    Ce resume extrait l'essentiel de tout le rapport pour une lecture rapide.
    """
    # Collecter les donnees cles
    decision_text = str(final_state.get("final_trade_decision") or "")[:600]
    trader_text   = str(final_state.get("trader_investment_plan") or "")[:400]
    market_text   = str(final_state.get("market_report") or "")[:400]
    expert_text   = str(final_state.get("expert_scalp_analysis") or "")[:400]

    price = indicators.get("current_price", "?") if indicators else "?"
    atr   = indicators.get("atr", "?") if indicators else "?"
    rsi   = indicators.get("rsi", "?") if indicators else "?"

    sig_lines = ""
    for s in signals[:2]:
        lots = compute_lot_sizes(s.get("entry_price"), s.get("stop_loss"))
        lot_str = " | ".join([f"{k}:lot={v['lot']}" for k, v in lots.items()])
        sig_lines += (f"\n  - {s.get('label')}: {s.get('action')} {s.get('exec_type','').upper()}"
                      f" @ {s.get('entry_price')} SL:{s.get('stop_loss')} TP:{s.get('take_profit')}"
                      f" RR 1:{s.get('rr')} | {lot_str}")

    # Déterminer la catégorie pour les règles spécifiques
    cat = _get_symbol_category(symbol)

    # Règles de priorité pour le résumé
    priority_rules = ""
    if cat == "VOLATILITY":
        rsi_num = float(rsi) if rsi != "?" else 50.0
        mr_dir = "BUY" if rsi_num < 30 else ("SELL" if rsi_num > 70 else "NEUTRE")
        priority_rules = f"""
REGLE CRITIQUE VOLATILITY INDEX :
RSI actuel = {rsi}. Direction mean reversion = {mr_dir}.
Le resume DOIT utiliser la direction de l'analyse scalping ({mr_dir}), PAS celle de TradingAgents.
Si TradingAgents dit SELL et scalping dit BUY → le resume dit BUY.
Le scalping court terme prime sur le signal fondamental pour les Volatility Index.
"""
    elif cat == "BOOM":
        priority_rules = "REGLE : Boom Index = BUY UNIQUEMENT. Si TradingAgents propose SELL, ignorer et dire BUY."
    elif cat == "CRASH":
        priority_rules = "REGLE : Crash Index = SELL UNIQUEMENT. Si TradingAgents propose BUY, ignorer et dire SELL."

    # Calcul lot correct pour le sizing
    lot_min = 0.2 if cat in ("BOOM", "CRASH") else 0.01
    capital = 20.0
    risk_amt = capital * 0.02  # $0.40 max

    prompt = f"""Tu es un assistant trading expert. Redige un RESUME EXECUTIF de 6 points.
REGLE ABSOLUE SIZING : lot minimum = {lot_min} | capital = ${capital} | risque max = ${ risk_amt} (2%)
Ne jamais proposer un risque > ${ risk_amt} sur ce capital.

{priority_rules}

SYMBOLE: {symbol} | CATEGORIE: {cat} | PRIX: {price} | ATR: {atr} | RSI: {rsi}

DECISION TRADINGAGENTS (signal fondamental — peut etre ignore si contradictoire avec scalping):
{decision_text}

PLAN TRADER:
{trader_text}

SIGNAUX CALCULES (lot min={lot_min}):
{sig_lines}

ANALYSE SCALPING (prioritaire pour le resume si Volatility/context court terme):
{expert_text}

FORMAT REQUIS (respecter exactement):
1. **Contexte de marche** : [tendance actuelle + RSI en 1 phrase]
2. **Signal principal** : [BUY/SELL/HOLD — doit etre coherent avec scalping si Volatility]
3. **Niveaux critiques** : [Entry, SL, TP avec distances en pips]
4. **Sizing CORRECT** : lot={lot_min} (minimum broker Deriv) | risque=$[X] sur ${ capital} — JAMAIS plus de ${ risk_amt}
5. **Condition d'entree** : [quand exactement entrer - immediat ou attendre quel niveau]
6. **Risques principaux** : [1-2 risques concrets, pas generiques]

INTERDICTION ABSOLUE : Ne jamais proposer un lot < {lot_min}.
Si risque > ${ risk_amt} → réduire le SL pour rester dans le budget, pas annuler le trade.
Toujours fournir un setup complet (Entry, SL, TP, Lot) même si conditions imparfaites."""

    try:
        import anthropic as _anth
        client = _anth.AnthropicBedrock(aws_region=os.getenv("AWS_REGION", "us-east-1"))
        msg = client.messages.create(
            model="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
            max_tokens=600,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text if msg.content else ""
    except Exception:
        # Fallback local si API indisponible
        lines = []
        lines.append(f"1. **Contexte de marche** : {symbol} en tendance {'baissiere' if rec=='SELL' else 'haussiere' if rec=='BUY' else 'neutre'}. RSI={rsi}, ATR={atr}.")
        lines.append(f"2. **Signal principal** : {rec} avec les indicateurs techniques alignes dans cette direction.")
        if signals:
            s0 = signals[0]
            lines.append(f"3. **Niveaux critiques** : Entree {s0.get('entry_price')} | SL {s0.get('stop_loss')} (-{s0.get('pips_sl')} pips) | TP {s0.get('take_profit')} (+{s0.get('pips_tp')} pips).")
            lots = compute_lot_sizes(s0.get("entry_price"), s0.get("stop_loss"))
            lot_str = " | ".join([f"{k}: lot {v['lot']} (risque ${v['risk_usd']})" for k, v in lots.items()])
            lines.append(f"4. **Sizing recommande** : {lot_str}.")
        lines.append(f"5. **Condition d'entree** : {'Entree immediate au marche possible.' if signals and signals[0].get('exec_type')=='market' else 'Attendre le retour sur le niveau limite propose.'}")
        lines.append(f"6. **Risques principaux** : Volatilite elevee (ATR={atr}). Surveiller les nouvelles macro et les niveaux cles de support/resistance.")
        return "\n\n".join(lines)


def save_report_word(symbol: str, trade_date: str, signal_rating: str,
                     final_state: Dict[str, Any],
                     params: Dict[str, Optional[float]],
                     confirmed: Optional[Dict[str, Any]] = None,
                     indicators: Optional[Dict] = None,
                     tv_summary: Optional[Dict[str, Any]] = None,
                     tv_comparison: Optional[Dict[str, Any]] = None) -> Optional[Path]:
    """Genere un rapport Word professionnel avec graphiques et niveaux calcules."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import tempfile, os as _os
    except ImportError:
        print("  [!] python-docx absent.")
        return None

    rec = _normalize_rating(signal_rating)
    sig_color = {"BUY": RGBColor(0, 150, 60), "SELL": RGBColor(180, 30, 30), "HOLD": RGBColor(180, 120, 0)}.get(rec, RGBColor(0,0,0))
    mpl_color = {"BUY": "#27ae60", "SELL": "#e74c3c", "HOLD": "#f39c12"}.get(rec, "#555")

    def _set_cell_fill(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # Calculer signaux depuis ai_server
    entry  = params.get("entry_price")
    sl_val = params.get("stop_loss")
    tp_val = params.get("take_profit")
    current_price = None
    computed_signals: List[Dict] = []

    if indicators:
        current_price = indicators.get("current_price")
        atr_val = indicators.get("atr")
        if current_price and atr_val and rec in ("BUY","SELL"):
            computed_signals = compute_signals(symbol, rec,
                                               float(current_price), float(atr_val))

    if not computed_signals and rec in ("BUY", "SELL"):
        computed = compute_entry_levels(symbol, rec)
        computed_signals = computed.get("signals", [])
        if not current_price:
            current_price = computed.get("current_price")

    # Utiliser le signal 1 (conservateur) si pas de params fournis
    if computed_signals and (not entry or not sl_val or not tp_val):
        s0 = computed_signals[0]
        entry  = entry  or s0.get("entry_price")
        sl_val = sl_val or s0.get("stop_loss")
        tp_val = tp_val or s0.get("take_profit")

    def _fmt(v):
        if v is None: return "—"
        return f"{v:.5f}" if abs(float(v)) < 1000 else f"{float(v):,.3f}"

    def _pips(a, b):
        if not a or not b: return "—"
        diff = abs(float(a) - float(b))
        p = _pip_value(float(a))
        return f"{diff/p:,.1f} pips"

    doc = Document()

    # --- Style global ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _set_para_justify(para):
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _colored_heading(text, level, color_rgb):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = color_rgb
        return h

    def _add_separator():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── PAGE DE TITRE ──────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_heading(f"Rapport d'Analyse — {symbol}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 60, 120)
        run.font.size = Pt(20)

    sub = doc.add_paragraph(f"Date d'analyse : {trade_date}    |    Généré le : {date.today()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── TRADINGVIEW MCP KOLA ───────────────────────────────────────
    if tv_summary is not None:
        _colored_heading("Analyse TradingView (MCP Kola)", 2, RGBColor(30, 90, 140))
        tv_p = doc.add_paragraph()
        if tv_summary.get("success"):
            tv_p.add_run(
                f"Direction SMC: {tv_summary.get('direction', '—')}  |  "
                f"Score biais: {tv_summary.get('bias_score', '—')}  |  "
                f"H1: {tv_summary.get('structure_h1', '—')}  M15: {tv_summary.get('structure_m15', '—')}\n"
            )
            reasons = tv_summary.get("bias_reasons") or []
            if reasons:
                tv_p.add_run("Raisons: " + "; ".join(reasons[:6]) + "\n")
            if tv_summary.get("spike_detected"):
                tv_p.add_run(
                    f"Spike Z={tv_summary.get('spike_z')} → {tv_summary.get('spike_direction')}\n"
                )
            if tv_summary.get("entry_valid"):
                tv_p.add_run(
                    f"Setup: entry={tv_summary.get('entry_price')} "
                    f"SL={tv_summary.get('stop_loss')} TP={tv_summary.get('take_profit')}\n"
                )
        else:
            tv_p.add_run(f"Indisponible: {tv_summary.get('error', 'CDP / TradingView')}\n")
        if tv_comparison:
            tv_p.add_run(f"\nConvergence: {tv_comparison.get('message', '—')}\n")
        _set_para_justify(tv_p)
        _add_separator()

    # ── SIGNAL PRINCIPAL ──────────────────────────────────────────
    _colored_heading("Signal TradingAgents", 2, RGBColor(30, 60, 120))

    # Tableau signal + niveaux
    # Tableau récapitulatif signal
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data_main = [
        ("Rating brut",   signal_rating),
        ("Décision",      rec),
        ("Prix actuel",   _fmt(current_price)),
        ("ATR (volatilité)", _fmt(indicators.get("atr") if indicators else None)),
    ]
    for i, (label, value) in enumerate(rows_data_main):
        cell_l = tbl.rows[i].cells[0]
        cell_r = tbl.rows[i].cells[1]
        cell_l.text = label
        cell_r.text = str(value)
        cell_l.paragraphs[0].runs[0].bold = True
        cell_l.paragraphs[0].runs[0].font.color.rgb = RGBColor(60, 60, 60)
        if label == "Décision":
            cell_r.paragraphs[0].runs[0].bold = True
            cell_r.paragraphs[0].runs[0].font.color.rgb = sig_color
            cell_r.paragraphs[0].runs[0].font.size = Pt(13)
        if i % 2 == 0:
            for cell in (cell_l, cell_r):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF2FF")
                tcPr.append(shd)

    doc.add_paragraph()

    # ── SIGNAUX DE TRADING (2 propositions) ───────────────────────
    if computed_signals and rec in ("BUY", "SELL"):
        _colored_heading("Signaux de Trading Proposes", 2, sig_color)

        sig_colors_fill = ["DFFFD6", "FFF3CD"]  # vert clair, jaune clair

        for idx, sig in enumerate(computed_signals[:2]):
            p_lbl = doc.add_paragraph()
            r_lbl = p_lbl.add_run(f"  Signal {idx+1} — {sig['label']}")
            r_lbl.bold = True
            r_lbl.font.size = Pt(11)
            r_lbl.font.color.rgb = sig_color

            # Calcul lots $10/$20/$50 pour ce signal
            sig_lots = compute_lot_sizes(sig.get("entry_price"), sig.get("stop_loss"),
                                         [10.0, 20.0, 50.0])
            lot_lines = []
            for acc_k, ld in sig_lots.items():
                ep   = float(sig.get("entry_price") or 0)
                tpp  = float(sig.get("take_profit") or 0)
                if ep > 0 and tpp > 0:
                    pip_val = _pip_value(ep)
                    tp_pips = abs(tpp - ep) / pip_val
                    pplt = 1.0 if ep > 1000 else (0.1 if ep > 10 else 1.0)
                    profit = round(ld["lot"] * tp_pips * pplt, 3)
                    lot_lines.append(f"{acc_k}: lot={ld['lot']}  risque=${ld['risk_usd']}  gain≈${profit}")
                else:
                    lot_lines.append(f"{acc_k}: lot={ld['lot']}  risque=${ld['risk_usd']}")

            fill_hex = sig_colors_fill[idx % 2]
            sig_rows = [
                ("Type d'ordre",  sig["exec_type"].upper()),
                ("Direction",     sig["action"]),
                ("Prix d'entrée", _fmt(sig["entry_price"])),
                ("Stop Loss",     f"{_fmt(sig['stop_loss'])}  ({_pips(sig['entry_price'], sig['stop_loss'])} de risque)"),
                ("Take Profit",   f"{_fmt(sig['take_profit'])}  ({_pips(sig['entry_price'], sig['take_profit'])} de gain)"),
                ("Ratio R/R",     f"1 : {sig['rr']}"),
                ("Prix actuel",   _fmt(sig["current_price"])),
            ] + [(f"Lot {ll.split(':')[0]}", ":".join(ll.split(":")[1:]).strip()) for ll in lot_lines]

            tbl_s = doc.add_table(rows=len(sig_rows), cols=2)
            tbl_s.style = "Table Grid"

            for j, (lbl, val) in enumerate(sig_rows):
                cl = tbl_s.rows[j].cells[0]
                cr = tbl_s.rows[j].cells[1]
                cl.text = lbl
                cr.text = str(val)
                cl.paragraphs[0].runs[0].bold = True
                row_fill = fill_hex if j < 7 else ("EAF7FB" if j % 2 == 0 else "D0EEF7")
                for cell in (cl, cr):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), row_fill)
                    tcPr.append(shd)

            doc.add_paragraph()

        _add_separator()

    doc.add_paragraph()

    tmp_files = []

    # ── GRAPHIQUE 1 : Prédiction directionnelle + RSI + MACD ─────────────
    if computed_signals and rec in ("BUY", "SELL"):
        pred_path = _make_prediction_chart(computed_signals,
                                           float(current_price or entry or 0),
                                           rec, symbol, indicators)
        if pred_path:
            tmp_files.append(pred_path)
            _colored_heading("Prédiction Directionnelle — Cône de Probabilité", 2,
                             RGBColor(20, 60, 120))
            p_pred = doc.add_paragraph()
            p_pred.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pred.add_run().add_picture(str(pred_path), width=Inches(6.0))
            doc.add_paragraph()
            _add_separator()

    # ── GRAPHIQUE 2 : Scorecard analytique (radar + probabilité + stats) ──
    if rec in ("BUY", "SELL"):
        score_path = _make_statistical_scorecard(
            computed_signals, indicators, rec, final_state, symbol)
        if score_path:
            tmp_files.append(score_path)
            _colored_heading("Scorecard Analytique — Confluence & Statistiques", 2,
                             RGBColor(30, 80, 140))
            p_score = doc.add_paragraph()
            p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_score.add_run().add_picture(str(score_path), width=Inches(6.5))
            doc.add_paragraph()
            _add_separator()

    # ── GRAPHIQUE 3 : Comparaison des 2 signaux + lots $10/$20/$50 ────────
    if computed_signals and rec in ("BUY", "SELL"):
        multi_path = _make_multi_signal_chart(computed_signals,
                                              float(current_price or entry or 0),
                                              symbol, rec)
        if multi_path:
            tmp_files.append(multi_path)
            _colored_heading("Comparaison des Signaux — Entry / SL / TP par compte", 2,
                             sig_color)
            p_multi = doc.add_paragraph()
            p_multi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_multi.add_run().add_picture(str(multi_path), width=Inches(6.5))
            doc.add_paragraph()
            _add_separator()

    # ── GRAPHIQUE 4 : Lot sizing $10 / $20 / $50 ─────────────────────────
    if computed_signals:
        lot_path = _make_lot_sizing_chart(computed_signals, [10.0, 20.0, 50.0])
        if lot_path:
            tmp_files.append(lot_path)
            _colored_heading("Money Management — Lot Sizing $10 / $20 / $50", 2,
                             RGBColor(40, 100, 60))
            p_lot = doc.add_paragraph()
            p_lot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_lot.add_run().add_picture(str(lot_path), width=Inches(6.0))
            doc.add_paragraph()

            # Tableau textuel lots $10/$20/$50
            if computed_signals[0].get("entry_price") and computed_signals[0].get("stop_loss"):
                all_lots = {}
                for acc in [10.0, 20.0, 50.0]:
                    ep = float(computed_signals[0]["entry_price"])
                    sp = float(computed_signals[0]["stop_loss"])
                    sl_dist = abs(ep - sp)
                    if sl_dist > 0:
                        pip_val = _pip_value(ep)
                        sl_pips = sl_dist / pip_val
                        pplt = 1.0 if ep > 1000 else (0.1 if ep > 10 else 1.0)
                        for pct, lbl in [(0.01, "1%"), (0.02, "2%")]:
                            lot = max(0.01, round((acc * pct) / (sl_pips * pplt), 2))
                            risk = round(lot * sl_pips * pplt, 3)
                            tp_dist = abs(float(computed_signals[0].get("take_profit") or ep) - ep)
                            tp_pips = tp_dist / pip_val
                            profit = round(lot * tp_pips * pplt, 3)
                            key = f"${acc:.0f} ({lbl})"
                            all_lots[key] = {"lot": lot, "risque": risk, "gain_potentiel": profit}

                n_cols = 4
                header = ["Compte", "Lot", "Risque max", "Gain potentiel"]
                rows_lot = [[k, f"{v['lot']:.2f}", f"${v['risque']:.3f}", f"${v['gain_potentiel']:.3f}"]
                            for k, v in all_lots.items()]
                tbl_lot = doc.add_table(rows=1 + len(rows_lot), cols=n_cols)
                tbl_lot.style = "Table Grid"
                tbl_lot.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, h in enumerate(header):
                    c = tbl_lot.rows[0].cells[j]
                    c.text = ""
                    r = c.paragraphs[0].add_run(h)
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    _set_cell_fill(c, "1F3C78")
                for i, row_data in enumerate(rows_lot):
                    fill_hex = "D5F5E3" if i % 2 == 0 else "FEF9E7"
                    for j, val in enumerate(row_data):
                        c = tbl_lot.rows[i + 1].cells[j]
                        c.text = val
                        c.paragraphs[0].runs[0].font.size = Pt(9)
                        _set_cell_fill(c, fill_hex)
                doc.add_paragraph()

            _add_separator()

    # ── GRAPHIQUE 5 : Jauge confiance + niveaux prix (compact) ───────────
    row_imgs = doc.add_paragraph()
    row_imgs.alignment = WD_ALIGN_PARAGRAPH.CENTER

    confidence_val = 0.75
    if confirmed:
        confidence_val = float(confirmed.get("confidence", 0.75))
    gauge_path = _make_gauge_chart(confidence_val, "Confiance", mpl_color)
    if gauge_path:
        tmp_files.append(gauge_path)
        run_g = row_imgs.add_run()
        run_g.add_picture(str(gauge_path), width=Inches(2.2))
        row_imgs.add_run("    ")

    if entry and sl_val and tp_val and current_price:
        levels_path = _make_price_level_chart(
            entry, sl_val, tp_val, current_price, symbol, rec)
        if levels_path:
            tmp_files.append(levels_path)
            run_l = row_imgs.add_run()
            run_l.add_picture(str(levels_path), width=Inches(3.0))

    _add_separator()

    # ── GRAPHIQUE 6 : Indicateurs techniques ────────────────────────────
    if indicators:
        ind_path = _make_indicator_bar_chart(indicators)
        if ind_path:
            tmp_files.append(ind_path)
            _colored_heading("Indicateurs Techniques (TradBOT — M1)", 2, RGBColor(30, 60, 120))
            p_ind = doc.add_paragraph()
            p_ind.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ind.add_run().add_picture(str(ind_path), width=Inches(4.5))
            doc.add_paragraph()
            _add_separator()

    # ── SECTIONS RAPPORT ─────────────────────────────────────────
    sections_map = [
        ("Décision Portfolio Manager", "final_trade_decision",  RGBColor(30, 60, 120)),
        ("Plan Trader",               "trader_investment_plan", RGBColor(30, 100, 60)),
        ("Rapport Marché",            "market_report",          RGBColor(80, 40, 120)),
        ("Rapport Sentiment / Social","sentiment_report",       RGBColor(120, 60, 0)),
        ("Rapport News",              "news_report",            RGBColor(0, 80, 120)),
        ("Rapport Fondamentaux",      "fundamentals_report",    RGBColor(60, 80, 0)),
    ]

    def _strip_emoji(text: str) -> str:
        """Supprime les emojis pour eviter les problemes d encodage Word."""
        return _re.sub(r'[\U00010000-\U0010ffff\U00002600-\U000027BF\U0001F300-\U0001FAFF]',
                       '', text, flags=_re.UNICODE).strip()

    def _clean_md(text: str) -> str:
        """Nettoie le markdown brut pour affichage Word : emojis, tirets longs, etc."""
        text = _strip_emoji(text)
        text = text.replace('—', '-').replace('–', '-')  # em/en dash
        text = text.replace('’', "'").replace('“', '"').replace('”', '"')
        return text.strip()

    def _set_cell_fill(cell, hex_color: str) -> None:
        """Applique une couleur de fond a une cellule de tableau."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _render_inline(para, text: str, color: "RGBColor", base_size: int = 11) -> None:
        """Ajoute du texte dans un paragraphe en gerant **bold**, *italic* et liens inline."""
        text = _clean_md(text)
        # Splitter sur **bold**, *italic*, `code`
        parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(base_size)
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = para.add_run(part[1:-1])
                run.italic = True
                run.font.size = Pt(base_size)
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(base_size - 1)
            else:
                clean = part.replace('**', '').replace('*', '').replace('`', '')
                if clean:
                    run = para.add_run(clean)
                    run.font.size = Pt(base_size)

    def _render_markdown_block(block: str, section_color: "RGBColor") -> None:
        """Convertit un bloc markdown en elements Word correctement formates."""
        block = block.strip()
        if not block:
            return

        lines_in_block = block.split("\n")

        # ── Ligne horizontale --- ──
        if _re.match(r'^-{3,}$', block):
            _add_separator()
            return

        # ── Titre # / ## / ### ──
        heading_m = _re.match(r'^(#{1,4})\s+(.*)', block)
        if heading_m:
            level_str = heading_m.group(1)
            title_text = _clean_md(heading_m.group(2))
            lvl = min(4, len(level_str) + 1)  # # -> h2, ## -> h3, ### -> h4
            h = doc.add_heading(title_text, level=lvl)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sz = {2: Pt(14), 3: Pt(12), 4: Pt(11)}.get(lvl, Pt(11))
            for run in h.runs:
                run.font.color.rgb = section_color
                run.font.size = sz
            return

        # ── Tableau markdown | col1 | col2 | ──
        table_lines = [l for l in lines_in_block if '|' in l and l.strip().startswith('|')]
        if len(table_lines) >= 2:
            sep_re = _re.compile(r'^\|[\s\-:| ]+\|$')
            content_lines = [l for l in table_lines if not sep_re.match(l.strip())]
            if len(content_lines) >= 1:
                cols_raw = [_clean_md(c) for c in content_lines[0].strip('|').split('|')]
                cols_raw = [c for c in cols_raw if c]
                n_cols = len(cols_raw)
                if n_cols > 0:
                    data_rows_raw = []
                    for l in content_lines[1:]:
                        cells = [_clean_md(c) for c in l.strip('|').split('|')]
                        while len(cells) < n_cols:
                            cells.append('')
                        data_rows_raw.append(cells[:n_cols])

                    tbl = doc.add_table(rows=1 + len(data_rows_raw), cols=n_cols)
                    tbl.style = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    # Header
                    for j, col_name in enumerate(cols_raw):
                        cell = tbl.rows[0].cells[j]
                        cell.text = ''
                        p_h = cell.paragraphs[0]
                        run_h = p_h.add_run(col_name)
                        run_h.bold = True
                        run_h.font.size = Pt(9)
                        run_h.font.color.rgb = RGBColor(255, 255, 255)
                        _set_cell_fill(cell, '1F3C78')  # fond bleu foncé
                    # Données
                    for i, dr in enumerate(data_rows_raw):
                        fill = 'EEF4FF' if i % 2 == 0 else 'FFFFFF'
                        for j, val in enumerate(dr):
                            cell = tbl.rows[i + 1].cells[j]
                            cell.text = ''
                            p_d = cell.paragraphs[0]
                            # Detecter si val est en gras (**val**)
                            if val.startswith('**') and val.endswith('**'):
                                run_d = p_d.add_run(val[2:-2])
                                run_d.bold = True
                            else:
                                run_d = p_d.add_run(_clean_md(val))
                            run_d.font.size = Pt(9)
                            _set_cell_fill(cell, fill)
                    doc.add_paragraph()
                    return

        # ── Liste a puces / numerotee ──
        list_pattern = _re.compile(r'^(\s*)([-*•✅⚠️✓→]|\d+\.)\s+(.*)')
        if all(list_pattern.match(l) or not l.strip() for l in lines_in_block if l.strip()):
            for line in lines_in_block:
                line_s = line.strip()
                if not line_s:
                    continue
                m = list_pattern.match(line)
                if m:
                    bullet_char = m.group(2)
                    text_item   = m.group(3)
                    is_num = bool(_re.match(r'\d+\.', bullet_char))
                    try:
                        style_name = 'List Number' if is_num else 'List Bullet'
                        p = doc.add_paragraph(style=style_name)
                    except Exception:
                        p = doc.add_paragraph()
                    _render_inline(p, text_item, section_color, base_size=10)
                else:
                    p = doc.add_paragraph()
                    _render_inline(p, line_s, section_color, base_size=10)
            return

        # ── Paragraphe multi-lignes : rendre ligne par ligne ──
        if len(lines_in_block) > 1:
            for line in lines_in_block:
                line_s = line.strip()
                if not line_s:
                    doc.add_paragraph()
                    continue
                # Sous-titre inline (ligne toute en **bold**)
                if _re.match(r'^\*\*[^*]+\*\*:?$', line_s):
                    p = doc.add_paragraph()
                    _render_inline(p, line_s, section_color, base_size=10)
                    continue
                p = doc.add_paragraph()
                _set_para_justify(p)
                _render_inline(p, line_s, section_color, base_size=10)
            return

        # ── Paragraphe simple ──
        p = doc.add_paragraph()
        _set_para_justify(p)
        sz = 10 if len(block) > 300 else 11
        _render_inline(p, block, section_color, base_size=sz)




    for title_sec, key, color in sections_map:
        text = str(final_state.get(key) or "").strip()
        if not text:
            continue
        _colored_heading(title_sec, 2, color)

        # Decouper par blocs (double saut de ligne) en preservant les tableaux multi-lignes
        blocks: List[str] = []
        current = []
        in_table = False
        for line in text.split("\n"):
            if "|" in line and line.strip().startswith("|"):
                in_table = True
                current.append(line)
            elif in_table and not line.strip():
                # Fin du tableau
                blocks.append("\n".join(current))
                current = []
                in_table = False
            elif not line.strip() and not in_table:
                if current:
                    blocks.append("\n".join(current))
                    current = []
            else:
                current.append(line)
        if current:
            blocks.append("\n".join(current))

        for block in blocks:
            _render_markdown_block(block, color)

        _add_separator()

    # ── ANALYSE EXPERT SCALPING CLAUDE ───────────────────────────
    expert_text = str(final_state.get("expert_scalp_analysis") or "").strip()
    if expert_text:
        _colored_heading("Analyse Expert Scalping — Claude ($10 / $50)", 2,
                         RGBColor(80, 0, 100))
        # Utiliser le meme renderer markdown que les autres sections
        expert_color = RGBColor(80, 0, 100)
        expert_blocks: List[str] = []
        current_exp: List[str] = []
        in_tbl_exp = False
        for line in expert_text.split("\n"):
            if "|" in line and line.strip().startswith("|"):
                in_tbl_exp = True
                current_exp.append(line)
            elif in_tbl_exp and not line.strip():
                expert_blocks.append("\n".join(current_exp))
                current_exp = []
                in_tbl_exp = False
            elif not line.strip() and not in_tbl_exp:
                if current_exp:
                    expert_blocks.append("\n".join(current_exp))
                    current_exp = []
            else:
                current_exp.append(line)
        if current_exp:
            expert_blocks.append("\n".join(current_exp))

        for blk in expert_blocks:
            _render_markdown_block(blk, expert_color)
        _add_separator()

    # ── RESUME EXECUTIF FINAL ─────────────────────────────────────
    # Genere par Claude a partir de toutes les sections du rapport
    resume_text = _generate_executive_summary(symbol, rec, final_state,
                                               computed_signals,
                                               indicators)
    if resume_text:
        _colored_heading("RESUME — Ce qu'il faut retenir", 1, RGBColor(20, 40, 100))
        resume_color = RGBColor(20, 40, 100)
        for block in resume_text.split("\n\n"):
            _render_markdown_block(block, resume_color)
        _add_separator()

    # ── PIED DE PAGE ──────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"TradBOT — Rapport généré automatiquement par le bridge TradingAgents | {date.today()}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Sauvegarder — si le fichier est déjà ouvert (Word), ajouter un suffixe unique
    sym_safe = symbol.replace(" ", "_").replace("/", "-")
    out_dir = _REPORTS_DIR / sym_safe
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{trade_date}_{sym_safe}_{rec}.docx"
    out_path = out_dir / filename
    try:
        doc.save(str(out_path))
    except PermissionError:
        import time as _time
        suffix = _time.strftime("_%H%M%S")
        out_path = out_dir / f"{trade_date}_{sym_safe}_{rec}{suffix}.docx"
        doc.save(str(out_path))
        print(f"  [!] Fichier verrouille — sauvegarde sous : {out_path.name}")
    print(f"  [OK] Rapport Word sauvegarde : {out_path}")

    # Nettoyer les fichiers temporaires
    for f in tmp_files:
        try:
            f.unlink()
        except Exception:
            pass

    return out_path


# ---------------------------------------------------------------------------
# Mode 2 : mode rapide (symbol/date en args, config depuis .env)
# ---------------------------------------------------------------------------

def run_quick(symbol: str, trade_date: str,
             analysts: List[str] = None,
             data_ticker: str = None,
             vendor: str = None) -> Dict[str, Any]:
    """
    Analyse rapide sans wizard.
    symbol      : nom affichage (ex: 'Boom 900 Index')
    data_ticker : ticker passe a TradingAgents (ex: 'BOOM900')
    vendor      : 'deriv' | 'yfinance' | None (auto-detecte)
    """
    if not _TA_CLI_AVAILABLE:
        sys.exit(f"[bridge] TradingAgents inaccessible: {_TA_IMPORT_ERR}")

    if analysts is None:
        analysts = ["market"]   # market seul = rapide (~2-4 min) pour indices/futures

    # Ticker et vendor : utiliser ceux fournis, sinon auto-detecter
    if data_ticker is None:
        data_ticker = _mt5_to_yfinance(symbol)

    # Pour les symboles Deriv, s'assurer que data_ticker est le nom court
    # (ex: CRASH1000, BOOM900) — jamais le nom long avec espaces
    if vendor == "deriv" or any(data_ticker.upper().startswith(p)
                                for p in ("BOOM","CRASH","1HZ","R_","FRX")):
        try:
            from tradingagents.dataflows.deriv_market import resolve_deriv_symbol  # type: ignore
            data_ticker = resolve_deriv_symbol(data_ticker)
        except ImportError:
            pass

    print(f"\n[bridge] Analyse : {symbol}")
    print(f"[bridge] Ticker  : {data_ticker} | Vendor: {vendor or 'auto'}")
    print(f"[bridge] Analystes: {', '.join(analysts)}")

    # Collecter les donnees techniques depuis ai_server
    print(f"[bridge] Collecte contexte technique depuis {_SERVER_URL} ...")
    tradbot_ctx = fetch_tradbot_context(symbol)

    # Pour les synthétiques Deriv : calculer les indicateurs depuis OHLC si ai_server n'a pas de données
    deriv_indicators: Optional[Dict[str, Any]] = None
    if vendor == "deriv" or _is_deriv_synthetic(data_ticker):
        deriv_indicators = compute_indicators_from_deriv(data_ticker)
        if deriv_indicators:
            deriv_ctx = build_deriv_technical_context(data_ticker, symbol, deriv_indicators)
            tradbot_ctx = (tradbot_ctx + "\n\n" + deriv_ctx) if tradbot_ctx else deriv_ctx

    # Construire le system prompt adapté à la catégorie du symbole
    category = _get_symbol_category(symbol)
    ind_for_prompt = deriv_indicators or {}
    # Essayer ai_server si pas d'indicateurs Deriv
    if not ind_for_prompt:
        try:
            r_ind = requests.get(
                f"{_SERVER_URL}/trading/indicators/{data_ticker.replace(' ','%20')}/M1", timeout=5)
            if r_ind.status_code == 200:
                ind_for_prompt = r_ind.json().get("indicators", {})
        except Exception:
            pass

    system_prompt = _build_system_prompt(symbol, category, ind_for_prompt if ind_for_prompt else None)
    print(f"[bridge] Categorie : {category} | System prompt construit ({len(system_prompt)} chars)")

    # Injecter le system prompt + contexte technique dans past_context
    full_ctx = system_prompt
    if tradbot_ctx:
        full_ctx += "\n\n" + tradbot_ctx
        print(f"[bridge] Contexte technique injecte ({len(tradbot_ctx)} chars)")
    else:
        print(f"[bridge] Serveur inaccessible — analyse avec system prompt uniquement")

    cfg = dict(_TA_DEFAULT)
    provider = (os.getenv("TRADINGAGENTS_LLM_PROVIDER")
                or cfg.get("llm_provider", "openai"))
    cfg["llm_provider"] = provider
    for env_k, cfg_k in [
        ("TRADINGAGENTS_QUICK_THINK_LLM", "quick_think_llm"),
        ("TRADINGAGENTS_DEEP_THINK_LLM",  "deep_think_llm"),
        ("TRADINGAGENTS_OUTPUT_LANGUAGE",  "output_language"),
    ]:
        v = os.getenv(env_k)
        if v:
            cfg[cfg_k] = v

    # Appliquer le vendor explicitement dans la config si fourni
    if vendor:
        try:
            from tradingagents.dataflows.deriv_catalog import apply_stock_data_vendor_to_config  # type: ignore
            apply_stock_data_vendor_to_config(cfg, vendor)
        except ImportError:
            pass

    graph = TradingAgentsGraph(
        selected_analysts=analysts,
        config=cfg,
        debug=False,
    )

    # Injecter system prompt + contexte technique dans past_context
    # Le Portfolio Manager reçoit ainsi les règles spécifiques à la catégorie
    original_get_past = graph.memory_log.get_past_context
    def _patched_get_past(ticker):
        base = original_get_past(ticker)
        return (base + "\n\n" + full_ctx) if base else full_ctx
        graph.memory_log.get_past_context = _patched_get_past

    # Pour les tickers Deriv (ex: frxXAUUSD), substituer par le ticker retail
    # afin que les analystes social/news/fundamentals trouvent les données
    propagate_ticker = _SOCIAL_TICKER_MAP.get(data_ticker, data_ticker)
    if propagate_ticker != data_ticker:
        print(f"[bridge] Social ticker : {data_ticker} -> {propagate_ticker}")

    final_state, signal_rating = graph.propagate(propagate_ticker, trade_date)

    # Indicateurs pour les graphiques du rapport
    # Priorité : ai_server (temps réel) -> OHLC Deriv calculé -> None
    indicators_raw = None
    try:
        r2 = requests.get(f"{_SERVER_URL}/trading/indicators/{data_ticker.replace(' ','%20')}/M1", timeout=5)
        if r2.status_code == 200:
            ind2 = r2.json().get("indicators")
            if ind2 and ind2.get("current_price"):
                indicators_raw = ind2
    except Exception:
        pass
    # Fallback : indicateurs calcules depuis OHLC Deriv
    if not indicators_raw and deriv_indicators:
        indicators_raw = deriv_indicators

    # Analyse expert Claude (scalping 15min + lots $10/$50)
    expert_analysis = ""
    if indicators_raw:
        rec_tmp = _normalize_rating(signal_rating)
        ta_summary = (str(final_state.get("final_trade_decision") or "")[:400] +
                      str(final_state.get("market_report") or "")[:400])
        sigs_tmp = []
        if rec_tmp in ("BUY","SELL"):
            cp = indicators_raw.get("current_price")
            at = indicators_raw.get("atr")
            if cp and at:
                sigs_tmp = compute_signals(symbol, rec_tmp, float(cp), float(at))
        print(f"[bridge] Analyse expert Claude scalping...")
        expert_analysis = claude_expert_analysis(
            symbol, rec_tmp, indicators_raw, sigs_tmp, ta_summary)

    return {
        "symbol":          symbol,
        "data_ticker":     data_ticker,
        "signal_rating":   signal_rating,
        "final_state":     final_state,
        "indicators":      indicators_raw,
        "expert_analysis": expert_analysis,
    }


# ---------------------------------------------------------------------------
# Affichage du rapport
# ---------------------------------------------------------------------------

_DIV = "-" * 72


def print_report(symbol: str, signal_rating: str,
                 final_state: Dict[str, Any],
                 params: Dict[str, Optional[float]]) -> None:
    rec   = _normalize_rating(signal_rating)
    color = {"BUY": "\033[92m", "SELL": "\033[91m", "HOLD": "\033[93m"}.get(rec, "")
    reset = "\033[0m"

    print(f"\n{_DIV}")
    print(f"  SIGNAL TradingAgents  --  {symbol}")
    print(_DIV)
    print(f"  Rating brut  : {signal_rating}")
    print(f"  Decision     : {color}{rec}{reset}")
    for lbl, val in [("Entry price", params["entry_price"]),
                     ("Stop Loss  ", params["stop_loss"]),
                     ("Take Profit", params["take_profit"])]:
        print(f"  {lbl} : {val if val is not None else 'non fourni'}")

    for section, title in [
        ("final_trade_decision",  "Decision Portfolio Manager"),
        ("trader_investment_plan","Plan Trader"),
    ]:
        text = str(final_state.get(section) or "")
        if text:
            print(f"\n{_DIV}\n  {title} :\n{_DIV}")
            print(text[:2000])
    print(_DIV)


# ---------------------------------------------------------------------------
# Confirmation interactive
# ---------------------------------------------------------------------------

def _ask_float(prompt: str, default: Optional[float]) -> Optional[float]:
    sfx = f" [{default}]" if default is not None else " [vide=ignorer]"
    raw = input(f"  {prompt}{sfx}: ").strip()
    if not raw:
        return default
    try:
        return float(raw)
    except ValueError:
        print("  [!] Valeur invalide ignoree.")
        return default


def interactive_confirm(rec: str,
                        params: Dict[str, Optional[float]],
                        signals: List[Dict] = None) -> Optional[Dict[str, Any]]:
    """Dialogue de confirmation avec choix entre les 2 signaux calcules."""

    _DIV2 = "-" * 60

    # Afficher les signaux proposes
    if signals and rec in ("BUY", "SELL"):
        print(f"\n{_DIV2}")
        print(f"  SIGNAUX PROPOSES — {rec}")
        print(_DIV2)
        for i, sig in enumerate(signals[:2], 1):
            print(f"\n  [{i}] {sig['label']}")
            print(f"      Type      : {sig['exec_type'].upper()}")
            print(f"      Entree    : {sig['entry_price']:.5f}  (prix actuel: {sig['current_price']:.5f})")
            print(f"      Stop Loss : {sig['stop_loss']:.5f}  (-{sig['pips_sl']} pips)")
            print(f"      Take Prof : {sig['take_profit']:.5f}  (+{sig['pips_tp']} pips)")
            print(f"      Ratio R/R : 1 : {sig['rr']}")
        print(f"\n  [3] Saisir des niveaux manuels")
        print(f"  [4] Annuler")
        print(_DIV2)

        choice = input("\n  Choisir le signal a envoyer [1/2/3/4] : ").strip()

        if choice == "4":
            print("  <- Annule.")
            return None

        if choice in ("1", "2"):
            idx = int(choice) - 1
            sig = signals[idx]
            lot_raw = input("  Lot size [vide=defaut EA]: ").strip()
            lot = float(lot_raw) if lot_raw else None
            conf_raw = input("  Confiance 0-100% [defaut=75]: ").strip()
            try:
                confidence = float(conf_raw) / 100.0 if conf_raw else 0.75
            except ValueError:
                confidence = 0.75
            return {
                "recommendation": rec,
                "confidence": max(0.0, min(1.0, confidence)),
                "entry_price":   sig["entry_price"],
                "stop_loss":     sig["stop_loss"],
                "take_profit":   sig["take_profit"],
                "execution_type": sig["exec_type"],
                "lot": lot,
            }

    # Pas de signaux calcules ou choix manuel (option 3)
    print("\n  Saisie manuelle des parametres.")
    d = input(f"  Direction [{rec}] (BUY/SELL/HOLD): ").strip().upper()
    rec = d if d in ("BUY", "SELL", "HOLD") else rec
    p = dict(params)
    p["entry_price"] = _ask_float("Entry price", p.get("entry_price"))
    p["stop_loss"]   = _ask_float("Stop Loss  ", p.get("stop_loss"))
    p["take_profit"] = _ask_float("Take Profit", p.get("take_profit"))

    lot_raw = input("  Lot size [vide=defaut EA]: ").strip()
    lot = float(lot_raw) if lot_raw else None

    exec_type = "market"
    if p.get("entry_price"):
        et = input("  Type [limit/stop/market, defaut=limit]: ").strip().lower()
        exec_type = et if et in ("limit", "stop", "market") else "limit"

    conf_raw = input("  Confiance 0-100% [defaut=75]: ").strip()
    try:
        confidence = float(conf_raw) / 100.0 if conf_raw else 0.75
    except ValueError:
        confidence = 0.75

    return {
        "recommendation": rec,
        "confidence": max(0.0, min(1.0, confidence)),
        "entry_price":   p.get("entry_price"),
        "stop_loss":     p.get("stop_loss"),
        "take_profit":   p.get("take_profit"),
        "execution_type": exec_type if p.get("entry_price") else "market",
        "lot": lot,
    }


# ---------------------------------------------------------------------------
# Extraction des niveaux d'alerte depuis l'analyse Claude
# ---------------------------------------------------------------------------

def extract_alert_levels(expert_analysis: str) -> Dict[str, Any]:
    """
    Parse le texte de l'analyse Claude pour extraire les niveaux BUY/SELL
    a surveiller (ex: 'Si prix descend a 4480-4485' ou 'sur rebond a 4555-4560').
    Retourne {buy_level, sell_level, buy_reason, sell_reason}.
    """
    result: Dict[str, Any] = {
        "buy_level": None, "sell_level": None,
        "buy_reason": None, "sell_reason": None,
    }
    if not expert_analysis:
        return result

    # Recherche patterns numeriques apres mots-cles
    buy_patterns  = [
        r'(?:BUY|achat|support|descend.*?)\D*([0-9]{3,6}(?:[.,][0-9]{1,5})?)',
        r'([0-9]{3,6}(?:[.,][0-9]{1,5})?)\s*(?:BB|boll|support|inferieure)',
    ]
    sell_patterns = [
        r'(?:SELL|vente|resistance|remonte.*?|rebond.*?)\D*([0-9]{3,6}(?:[.,][0-9]{1,5})?)',
        r'([0-9]{3,6}(?:[.,][0-9]{1,5})?)\s*(?:SMA|resistance|superieure|rejet)',
    ]

    # Chercher dans les lignes contenant BUY/SELL
    for line in expert_analysis.split("\n"):
        line_up = line.upper()
        for pat in buy_patterns:
            m = _re.search(pat, line, _re.IGNORECASE)
            if m and result["buy_level"] is None:
                try:
                    v = float(m.group(1).replace(",", "."))
                    if v > 10:
                        result["buy_level"]  = v
                        result["buy_reason"] = line.strip()[:100]
                except ValueError:
                    pass
        for pat in sell_patterns:
            m = _re.search(pat, line, _re.IGNORECASE)
            if m and result["sell_level"] is None:
                try:
                    v = float(m.group(1).replace(",", "."))
                    if v > 10:
                        result["sell_level"]  = v
                        result["sell_reason"] = line.strip()[:100]
                except ValueError:
                    pass

    return result


def push_alert_levels(symbol: str, expert_analysis: str,
                      buy_level: Optional[float] = None,
                      sell_level: Optional[float] = None,
                      buy_reason: str = "",
                      sell_reason: str = "") -> bool:
    """Enregistre les niveaux d'alerte sur le serveur (EA les poll et envoie push MT5)."""
    # Si pas de niveaux fournis, les extraire de l'analyse Claude
    if buy_level is None or sell_level is None:
        extracted = extract_alert_levels(expert_analysis)
        buy_level  = buy_level  or extracted.get("buy_level")
        sell_level = sell_level or extracted.get("sell_level")
        buy_reason  = buy_reason  or extracted.get("buy_reason", "")
        sell_reason = sell_reason or extracted.get("sell_reason", "")

    if not buy_level and not sell_level:
        return False

    payload = {
        "symbol":      symbol,
        "buy_level":   buy_level,
        "sell_level":  sell_level,
        "buy_reason":  buy_reason  or "Niveau support bridge",
        "sell_reason": sell_reason or "Niveau resistance bridge",
    }
    try:
        r = requests.post(f"{_SERVER_URL}/alert-levels", json=payload, timeout=10)
        r.raise_for_status()
        print(f"  [OK] Alertes niveaux: BUY@{buy_level} SELL@{sell_level} -> MT5 push active")
        return True
    except Exception as e:
        print(f"  [!] /alert-levels : {e}")
        return False


# ---------------------------------------------------------------------------
# Push vers ai_server
# ---------------------------------------------------------------------------

def _clean_symbol_for_server(symbol: str) -> str:
    """
    Nettoie le nom d affichage du symbole pour l envoyer au serveur MT5.
    Ex: 'Or — XAUUSD (-> frxXAUUSD)' -> 'XAUUSD'
        'Crash 150 Index' -> 'CRASH 150 INDEX'
        'EURUSD=X' -> 'EURUSD'
    """
    s = symbol.strip().upper()
    # Supprimer les parties entre parentheses et apres les tirets longs
    import re as _re2
    s = _re2.sub(r'\(.*?\)', '', s).strip()
    s = _re2.sub(r'[—–→].*', '', s).strip()
    s = _re2.sub(r'\s+', ' ', s).strip()
    # Supprimer suffixes yfinance
    s = s.replace('=X', '').replace('=F', '').strip()
    # Si vide apres nettoyage, garder l original en majuscules
    return s if s else symbol.strip().upper()


def push_manual_report(symbol: str, data: Dict[str, Any], reasoning: str) -> bool:
    payload = {
        "symbol": _clean_symbol_for_server(symbol),
        "recommendation": data["recommendation"],
        "confidence": data["confidence"],
        "reasoning": reasoning[:4000],
        "execution_type": data.get("execution_type"),
        "entry_price": data.get("entry_price"),
        "stop_loss":   data.get("stop_loss"),
        "take_profit": data.get("take_profit"),
    }
    try:
        r = requests.post(f"{_SERVER_URL}/tradingagents/manual-report",
                          json=payload, timeout=10)
        r.raise_for_status()
        print(f"  [OK] Manuel report envoye : {r.json()}")
        return True
    except Exception as e:
        print(f"  [ERR] manual-report : {e}")
        return False


def push_session_bias(symbol: str, direction: str, confidence: float,
                      expires_hours: int = 8) -> bool:
    """Stocke le biais de session TradingAgents dans ai_server (/session-bias).
    L'EA MT5 lit ce biais toutes les heures et filtre les entrées en conséquence.
    expires_hours=8 : couvre une session London+NY complète.
    """
    try:
        payload = {
            "symbol":       _clean_symbol_for_server(symbol),
            "direction":    direction.upper(),   # BUY | SELL | NEUTRAL
            "confidence":   round(confidence, 4),
            "expires_hours": expires_hours,
        }
        r = requests.post(f"{_SERVER_URL}/session-bias", json=payload, timeout=5)
        if r.status_code == 200:
            print(f"[bridge] /session-bias stocké : {direction} conf={confidence:.0%} "
                  f"(expire dans {expires_hours}h)")
            return True
        print(f"  [WARN] /session-bias HTTP {r.status_code}")
        return False
    except Exception as e:
        print(f"  [WARN] /session-bias inaccessible : {e}")
        return False


def push_pending_order(symbol: str, data: Dict[str, Any], status: str = "ready") -> bool:
    sym_clean = _clean_symbol_for_server(symbol)
    action    = data["recommendation"].upper()
    cat       = _get_symbol_category(symbol)

    # ── Règle 0 : BOOM = BUY uniquement, CRASH = SELL uniquement ──
    if cat == "BOOM" and action == "SELL":
        print(f"  [!] Règle 0 BOOM: SELL rejeté sur {sym_clean} → forcé en BUY")
        action = "BUY"
    elif cat == "CRASH" and action == "BUY":
        print(f"  [!] Règle 0 CRASH: BUY rejeté sur {sym_clean} → forcé en SELL")
        action = "SELL"

    # ── Calcul lot minimum si non fourni ──
    lot = data.get("lot")
    if lot is None:
        # Lot minimum par catégorie (Deriv)
        # Boom/Crash synthétiques : 0.2
        # XAUUSD (Gold), Forex, Volatility, Index : 0.01
        if cat in ("BOOM", "CRASH"):
            lot_min = 0.2
        else:
            lot_min = 0.01  # XAUUSD, EURUSD, Volatility, indices
        # Calcul risque 2% sur capital cible $20
        capital  = 20.0
        risk_amt = capital * 0.02  # $0.40
        entry    = data.get("entry_price")
        sl       = data.get("stop_loss")
        if entry and sl and abs(float(entry) - float(sl)) > 0:
            sl_dist     = abs(float(entry) - float(sl))
            pip_val_est = 0.01  # valeur approximative pip/lot pour synthétiques
            calc_lot    = risk_amt / (sl_dist * pip_val_est)
            lot = max(lot_min, round(calc_lot - (calc_lot % lot_min), 2))
        else:
            lot = lot_min
        print(f"  [bridge] Lot calculé: {lot} (min={lot_min}, risque ~${lot*0.01:.3f})")

    # ── TP obligatoire pour Boom/Crash (spike = sortie rapide) ──
    tp = data.get("take_profit")
    entry = data.get("entry_price")
    sl    = data.get("stop_loss")
    if tp is None and entry and sl:
        sl_dist = abs(float(entry) - float(sl))
        if action == "BUY":
            tp = round(float(entry) + sl_dist * 2.0, 5)  # RR 1:2
        else:
            tp = round(float(entry) - sl_dist * 2.0, 5)
        print(f"  [bridge] TP calcule automatiquement: {tp} (RR 1:2)")

    payload = {
        "symbol":         sym_clean,
        "action":         action.lower(),
        "execution_type": data.get("execution_type", "market"),
        "entry_price":    entry,
        "stop_loss":      sl,
        "take_profit":    tp,
        "lot":            lot,
        "confidence":     data.get("confidence", 0.75),
        "source":         "tradingagents_cli",
        "comment":        "TA_BRIDGE",
        "status":         status,
    }
    try:
        r = requests.post(f"{_SERVER_URL}/pending-order", json=payload, timeout=10)
        r.raise_for_status()
        resp = r.json()
        print(f"  [OK] Ordre en queue : id={resp.get('order_id')} | "
              f"{payload['action'].upper()} @ {payload.get('entry_price')} "
              f"SL={payload.get('stop_loss')} TP={payload.get('take_profit')} "
              f"lot={lot}")
        return True
    except Exception as e:
        print(f"  [ERR] /pending-order : {e}")
        return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="TradBOT Bridge - TradingAgents -> MT5",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples:
  .\bridge.bat                        # wizard complet (symbol, date, analystes, LLM...)
  .\bridge.bat --symbol EURUSD        # rapide : symbol fixe, reste depuis .env
  .\bridge.bat --symbol XAUUSD --auto # rapide + pas de confirmation
  .\bridge.bat --symbol EURUSD --no-pending  # rapport seul, pas d ordre MT5
        """,
    )
    parser.add_argument("--symbol", "-s", default=None,
                        help="Symbole MT5. Si absent: wizard complet lance")
    parser.add_argument("--date", "-d", default=str(date.today()),
                        help="Date YYYY-MM-DD (mode rapide seulement)")
    parser.add_argument("--analysts", "-a", default="market,social,news,fundamentals",
                        help="Analystes separes par virgule (defaut: tous). "
                             "Choix: market, social, news, fundamentals")
    parser.add_argument("--no-pending", action="store_true",
                        help="N'envoie pas d'ordre pending, rapport manuel seulement")
    parser.add_argument("--auto", action="store_true",
                        help="Pas de confirmation interactive avant envoi")
    parser.add_argument("--no-tv", action="store_true",
                        help="Desactive analyse TradingView MCP (TA seul)")
    parser.add_argument("--no-whatsapp", action="store_true",
                        help="Pas d'envoi WhatsApp unifie")
    parser.add_argument("--followup", action="store_true",
                        help="Lance le monitor de suivi 10 min en arriere-plan apres envoi")
    args = parser.parse_args()

    # Valider les analystes (noms exacts TradingAgents)
    _valid = {"market", "social", "news", "fundamentals"}
    analysts = [a.strip().lower() for a in args.analysts.split(",") if a.strip()]
    # Alias : sentiment -> social
    analysts = ["social" if a == "sentiment" else a for a in analysts]
    analysts = [a for a in analysts if a in _valid] or ["market", "social"]

    # --- Choisir le mode ---
    if args.symbol is None:
        # Mode 1 : selecteur interactif par categorie de broker
        sym_label, ticker_id, vendor = select_symbol_interactive()
        result = run_quick(sym_label, args.date,
                           analysts=analysts,
                           data_ticker=ticker_id,
                           vendor=vendor)
    else:
        # Mode 2 : symbol fourni en argument (auto-detection vendor)
        sym = args.symbol.strip()
        ticker_id = _mt5_to_yfinance(sym)
        vendor = "deriv" if any(ticker_id.upper().startswith(p)
                                for p in ("BOOM","CRASH","1HZ","R_","FRX")) else "yfinance"
        result = run_quick(sym, args.date,
                           analysts=analysts,
                           data_ticker=ticker_id,
                           vendor=vendor)

    symbol          = result["symbol"]
    signal_rating   = result["signal_rating"]
    final_state     = result["final_state"]
    indicators      = result.get("indicators")
    expert_analysis = result.get("expert_analysis", "")
    rec             = _normalize_rating(signal_rating)
    params          = _extract_order_params(final_state)

    # Calculer les 2 signaux depuis ai_server
    computed_signals: List[Dict] = []
    current_price_main = None
    if indicators and rec in ("BUY", "SELL"):
        cp = indicators.get("current_price")
        at = indicators.get("atr")
        if cp and at:
            current_price_main = float(cp)
            computed_signals = compute_signals(symbol, rec, float(cp), float(at))
    if not computed_signals and rec in ("BUY", "SELL"):
        lvl = compute_entry_levels(symbol, rec)
        computed_signals = lvl.get("signals", [])
        if not current_price_main:
            current_price_main = lvl.get("current_price")

    # Afficher le rapport terminal
    print_report(symbol, signal_rating, final_state, params)

    # Afficher les signaux calcules dans le terminal
    if computed_signals:
        print(f"\n  Prix actuel : {current_price_main}")
        for sig in computed_signals[:2]:
            lots = compute_lot_sizes(sig.get("entry_price"), sig.get("stop_loss"))
            lot_str = "  ".join([f"{k}:lot={v['lot']}" for k, v in lots.items()])
            print(f"  [{sig['label']}] {sig['action']} {sig['exec_type'].upper()}"
                  f" @ {sig['entry_price']}  SL:{sig['stop_loss']} (-{sig['pips_sl']} pips)"
                  f"  TP:{sig['take_profit']} (+{sig['pips_tp']} pips)  RR 1:{sig['rr']}"
                  f"  | {lot_str}")

    # Afficher l'analyse expert Claude
    if expert_analysis:
        print(f"\n{'='*60}")
        print("  ANALYSE EXPERT SCALPING (Claude)")
        print('='*60)
        print(expert_analysis)

    reasoning = (str(final_state.get("final_trade_decision") or "") + "\n\n"
                 + str(final_state.get("trader_investment_plan") or ""))

    if args.auto:
        sig0 = computed_signals[0] if computed_signals else {}
        confirmed = {
            "recommendation": rec,
            "confidence": 0.75,
            "entry_price":   sig0.get("entry_price") or params.get("entry_price"),
            "stop_loss":     sig0.get("stop_loss")   or params.get("stop_loss"),
            "take_profit":   sig0.get("take_profit") or params.get("take_profit"),
            "execution_type": sig0.get("exec_type", "market"),
            "lot": None,
        }
    else:
        confirmed = interactive_confirm(rec, params, signals=computed_signals)

    # Sauvegarder le rapport Word (meme si signal annule)
    trade_date = args.date if args.symbol else str(date.today())
    # Injecter l'analyse expert dans final_state pour l'inclure dans le Word
    if expert_analysis:
        final_state = dict(final_state)
        final_state["expert_scalp_analysis"] = expert_analysis

    # ── TradingView MCP Kola ─────────────────────────────────────
    tv_raw: Dict[str, Any] = {}
    tv_summary: Dict[str, Any] = {"success": False, "error": "skipped"}
    tv_comparison: Dict[str, Any] = {"verdict": "SKIPPED", "allow_pending": True}

    if not args.no_tv and _UNIFIED_BRIDGE_OK and confirmed is not None:
        print("\n[bridge] Analyse TradingView (MCP Kola + CDP)...")
        tv_raw, tv_summary = run_tv_analysis_for_bridge(symbol)
        if tv_summary.get("success"):
            print(f"  TV: {tv_summary.get('direction')} | "
                  f"confluence={tv_summary.get('confluence_score')} | "
                  f"prix={tv_summary.get('current_price')}")
        else:
            print(f"  [!] TV: {tv_summary.get('error', tv_raw.get('error', 'echec'))}")
        tv_comparison = compare_ta_and_tv(rec, tv_summary)
        print(f"  Convergence: {tv_comparison.get('message')}")
        if tv_comparison.get("aligned"):
            confirmed = merge_confirmed_with_tv(confirmed, tv_summary, tv_comparison)
    elif not args.no_tv and not _UNIFIED_BRIDGE_OK:
        print(f"  [!] unified_bridge: {_UNIFIED_BRIDGE_ERR}")

    report_path = save_report_word(symbol, trade_date, signal_rating, final_state, params,
                                   confirmed=confirmed if confirmed else None,
                                   indicators=indicators,
                                   tv_summary=tv_summary,
                                   tv_comparison=tv_comparison)

    # Envoi automatique du rapport Word en pièce jointe WhatsApp
    if report_path and not args.no_whatsapp:
        try:
            sys.path.insert(0, str(_HERE))
            from send_tradingagents_report import send_whatsapp_file  # type: ignore
            rec_now = (confirmed or {}).get("recommendation", signal_rating) or signal_rating
            caption = (
                f"📊 *RAPPORT TRADINGAGENTS — {symbol}*\n"
                f"Direction : *{rec_now}*\n"
                f"Date : {trade_date}\n\n"
                f"Rapport complet en pièce jointe 👇"
            )
            print(f"\n[bridge] Envoi rapport Word par WhatsApp → {report_path.name} ...")
            ok = send_whatsapp_file(str(report_path), caption)
            print(f"  {'✅ Rapport envoyé' if ok else '❌ Échec envoi rapport'}")
        except Exception as _wa_err:
            print(f"  [!] Envoi WhatsApp rapport : {_wa_err}")

    if confirmed is None:
        print("\n[bridge] Signal annule. Rapport Word sauvegarde.")
        return

    print(f"\n[bridge] Envoi vers {_SERVER_URL} ...")
    push_manual_report(symbol, confirmed, reasoning)

    bias_conf = confirmed.get("confidence") or 0.70
    push_session_bias(symbol, rec, float(bias_conf))

    if expert_analysis:
        push_alert_levels(symbol, expert_analysis)

    allow_pending = tv_comparison.get("allow_pending", True)
    if not args.no_tv and _UNIFIED_BRIDGE_OK:
        allow_pending = bool(tv_comparison.get("aligned"))

    if not args.no_pending and confirmed["recommendation"] in ("BUY", "SELL"):
        if allow_pending:
            push_pending_order(symbol, confirmed)
            push_unified_state(symbol, confirmed, tv_summary, tv_comparison)
        elif (not args.no_tv and _UNIFIED_BRIDGE_OK
              and tv_comparison.get("verdict") == "CONFLICT"):
            # Conflit TA vs TV — pousser en CONFLICT_PENDING et boucle de résolution
            conflict_conf = round(float(confirmed.get("confidence") or 0.75) * 0.7, 4)
            conflict_data = dict(confirmed)
            conflict_data["confidence"] = conflict_conf
            push_pending_order(symbol, conflict_data, status="CONFLICT_PENDING")
            push_unified_state(symbol, conflict_data, tv_summary, tv_comparison)
            ta_dir = tv_comparison.get("ta_direction", rec)
            tv_dir = tv_comparison.get("tv_direction", "?")
            print(f"  [CONFLIT] TA={ta_dir} vs TV={tv_dir} — ordre CONFLICT_PENDING cree")
            print(f"  [CONFLIT] Boucle resolution: 3 re-scans TV x 5min (15min max)")
            wa_conflict = (
                f"*CONFLIT SIGNAL* [{datetime.utcnow().strftime('%H:%M UTC')}]\n"
                f"Symbole: *{symbol}*\n"
                f"TradingAgents: *{ta_dir}* ({conflict_conf:.0%})\n"
                f"TradingView MCP: *{tv_dir}* (confluence={tv_summary.get('confluence_score', '?')})\n"
                f"---\n"
                f"Situation: rebond TV court-terme vs tendance TA HTF\n"
                f"Action: attente realignement TV - refresh auto toutes les 5min\n"
                f"Delai max: 15min (3 tentatives)"
            )
            send_unified_whatsapp(wa_conflict)
            resolve_conflict_loop(
                symbol, ta_dir,
                interval_sec=300,
                max_retries=3,
            )
        else:
            print("  [!] Pending NON envoye — pas de convergence TA/TV")

    # WhatsApp unifie (un seul message)
    if not args.no_whatsapp and _UNIFIED_BRIDGE_OK:
        wa_msg = format_unified_whatsapp(
            symbol, rec, float(bias_conf), confirmed,
            tv_summary, tv_comparison,
            expert_snippet=expert_analysis[:500] if expert_analysis else "",
        )
        send_unified_whatsapp(wa_msg)

    if args.followup:
        import subprocess
        follow_script = _HERE / "bridge_followup_monitor.py"
        if follow_script.exists():
            subprocess.Popen(
                [sys.executable, str(follow_script),
                 "--symbol", symbol, "--interval", "600"],
                cwd=str(_TRADBOT_ROOT),
                creationflags=subprocess.CREATE_NEW_CONSOLE if sys.platform == "win32" else 0,
            )
            print("  [OK] Monitor suivi 10 min demarre en arriere-plan")

    print("\n[bridge] OK. L'EA MT5 recevra le signal au prochain appel /decision.")
    print(f"  Status  : GET {_SERVER_URL}/tradingagents/realtime/status")
    print(f"  Ordres  : GET {_SERVER_URL}/pending-order")
    print(f"  Rapports: {_REPORTS_DIR}\n")


if __name__ == "__main__":
    main()
