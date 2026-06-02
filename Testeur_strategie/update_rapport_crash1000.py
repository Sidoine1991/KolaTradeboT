#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Met a jour le rapport Backtest DerivEAPro avec les resultats Crash 1000.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r"D:\Dev\TradBOT\Testeur_strategie\Rapport_Backtest_DerivEAPro_v6.02_FINAL.docx"

doc = Document(DOC_PATH)

# ── 1. Marquer la tache Crash 1000 comme completee dans section 7.4 ─────────
for para in doc.paragraphs:
    if "Backtest sur Crash 1000 avec memes parametres" in para.text:
        for run in para.runs:
            if "Backtest sur Crash 1000" in run.text:
                run.text = run.text.replace(
                    "Backtest sur Crash 1000 avec memes parametres",
                    "[COMPLETE - voir section 8] Backtest sur Crash 1000"
                )
                run.bold = True
        break

# ── 2. Ajouter section 8 apres la derniere page ──────────────────────────────
doc.add_page_break()

doc.add_heading("8. Backtest Crash 1000 Index", level=1)

doc.add_paragraph(
    "Ce backtest valide les parametres de production (Test 3) sur Crash 1000 Index, "
    "instrument SELL-only (spikes baissiers). Periode : 2025.01.01 - 2026.01.01, "
    "qualite historique 100% ticks reels."
)

# ── 8.1 Parametres utilises ─────────────────────────────────────────────────
doc.add_heading("8.1 Parametres utilises (identiques Test 3 Boom 1000)", level=2)

t_params = doc.add_table(rows=1, cols=3)
t_params.style = "Light List Accent 2"
hdr = t_params.rows[0].cells
hdr[0].text = "Parametre"
hdr[1].text = "Valeur"
hdr[2].text = "Notes"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

params = [
    ("InpSL_ATR",             "1.0",   "SL = 1x ATR"),
    ("InpTP_ATR",             "1.5",   "TP = 1.5x ATR"),
    ("InpTimeStopMinutes",    "8",     "Sortie forcee a 8 min"),
    ("InpQuickExitMinProfit", "0.30",  "Sortie rapide >= $0.30"),
    ("InpMinSignalQuality",   "70%",   "Filtre qualite signal"),
    ("InpUseGOMFilter",       "true",  "GOM aligne obligatoire"),
    ("InpEnableSpikeCrash",   "true",  "Detection spike Crash active"),
    ("InpSynthZScoreMin",     "0.9",   "Seuil ZScore (a optimiser)"),
    ("InpSynthBodyAtrMult",   "0.32",  "Seuil corps (a optimiser)"),
    ("InpImminenceThresh",    "60.0",  "Score imminence minimum"),
]
for p, v, n in params:
    row = t_params.add_row()
    row.cells[0].text = p
    row.cells[1].text = v
    row.cells[2].text = n

doc.add_paragraph()

# ── 8.2 Resultats comparatifs ───────────────────────────────────────────────
doc.add_heading("8.2 Resultats du Backtest Crash 1000 vs Boom 1000", level=2)

t_res = doc.add_table(rows=1, cols=3)
t_res.style = "Light List Accent 2"
hdr2 = t_res.rows[0].cells
hdr2[0].text = "Metrique"
hdr2[1].text = "Crash 1000 (M1, 2025)"
hdr2[2].text = "Boom 1000 (Ref. Test 3)"
for cell in hdr2:
    for run in cell.paragraphs[0].runs:
        run.bold = True

results = [
    ("Symbole",           "Crash 1000 Index",   "Boom 1000 Index"),
    ("Periode",           "2025 M1 (1 an)",      "2019-2026 M1 (7 ans)"),
    ("Depot initial",     "$1 000",              "$1 000"),
    ("Qualite historique","100% ticks reels",    "100% ticks reels"),
    ("Barres / Tiques",   "525 528 / 30 372 824","N/A"),
    ("Profit Total Net",  "-$16.54",             "+$196.90"),
    ("Profit Brut",       "+$37.87",             "+$607.84"),
    ("Perte Brute",       "-$54.41",             "-$410.94"),
    ("Retention Profit",  "-44%",                "32%"),
    ("Total Trades",      "138",                 "1 438"),
    ("Positions Courtes", "138 (30.43%)",        "0"),
    ("Positions Longues", "0",                   "1 438 (74.20%)"),
    ("Win Rate",          "30.43%",              "74.20%"),
    ("Profit Factor",     "0.696",               "1.48"),
    ("Max Drawdown",      "1.97% ($19.75)",      "1.11%"),
    ("Sharpe Ratio",      "-5.00",               "45.82"),
    ("Recovery Factor",   "-0.837",              "17.64"),
    ("Avg Win",           "+$0.90",              "+$0.57"),
    ("Avg Loss",          "-$0.57",              "-$1.11"),
    ("Ratio Win/Loss",    "1.59",                "0.51"),
    ("Duree moy. trade",  "7 min 04s",           "3 min 30s"),
    ("Max pertes consec.", "10 (-$5.59)",         "N/A"),
]
for label, v_crash, v_boom in results:
    row = t_res.add_row()
    row.cells[0].text = label
    row.cells[1].text = v_crash
    row.cells[2].text = v_boom

doc.add_paragraph()

# ── 8.3 Analyse ─────────────────────────────────────────────────────────────
doc.add_heading("8.3 Analyse des Resultats", level=2)

doc.add_heading("Facteurs explicatifs des performances inferieures", level=3)

analyses = [
    (
        "1. Win Rate 30% vs 74% :",
        "Le seuil de detection spike (ZScore=0.9, BodyAtrMult=0.32) a ete calibre sur Boom 1000. "
        "Les spikes Crash ont une morphologie differente : mouvement brutal puis rebond rapide. "
        "Beaucoup de faux signaux sont detectes, d ou le faible win rate de 30.43%."
    ),
    (
        "2. Ratio Avg Win / Avg Loss favorable mais insuffisant :",
        "Sur Crash 1000 : Avg Win = $0.90, Avg Loss = $0.57 (ratio 1.59). "
        "Ce bon ratio n est pas suffisant avec seulement 30% de wins : "
        "42 trades gagnants x $0.90 = $37.87 vs 96 perdants x $0.57 = $54.41. "
        "Il faudrait un win rate >= 36% pour etre profitable avec ce ratio."
    ),
    (
        "3. Periode limitee (1 an, 138 trades) :",
        "Le backtest Boom 1000 couvre 7 ans avec 1 438 trades (significatif statistiquement). "
        "Le Crash 1000 ne couvre que 2025 avec 138 trades. "
        "Ce volume est insuffisant pour une conclusion definitive."
    ),
    (
        "4. Trades uniquement SELL :",
        "Crash 1000 = 100% positions courtes. L EA fonctionne en BUY sur Boom. "
        "La logique de detection / gestion est identique, mais la psychologie du marche "
        "differe (rebonds brutaux post-spike Crash plus violents)."
    ),
]
for title, text in analyses:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(" " + text)

doc.add_paragraph()

# ── 8.4 Verdict ─────────────────────────────────────────────────────────────
doc.add_heading("8.4 Verdict et Plan d Optimisation", level=2)

p_verdict = doc.add_paragraph()
p_verdict.add_run("VERDICT : ").bold = True
p_verdict.add_run(
    "Les parametres Test 3 (Boom 1000) ne sont PAS directement applicables au Crash 1000. "
    "Reoptimisation des parametres de detection requise avant tout deploiement."
)

doc.add_paragraph()
doc.add_paragraph("Plan d optimisation Crash 1000 :")

plan = [
    "1. Etendre la periode : backtest 2022-2026 (min 3 ans, ~400-500 trades)",
    "2. Optimiser InpSynthZScoreMin : tester plage 0.7 - 1.2 (par pas de 0.1)",
    "3. Optimiser InpSynthBodyAtrMult : tester plage 0.20 - 0.50 (par pas de 0.05)",
    "4. Tester InpTP_ATR=2.0 et 2.5 (ratio Win/Loss 1.59 suggere un TP plus large)",
    "5. Evaluer InpImminenceThresh : tester 50, 60, 70 sur Crash",
    "6. Objectif : win rate >= 45% + profit factor >= 1.30",
    "7. Ne PAS deployer en production Crash 1000 avant ces ajustements",
]
for item in plan:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph()

# ── 8.5 Synthese multi-symbol ────────────────────────────────────────────────
doc.add_heading("8.5 Synthese Multi-Symbol (etat actuel)", level=2)

t_multi = doc.add_table(rows=1, cols=5)
t_multi.style = "Light List Accent 2"
hdr3 = t_multi.rows[0].cells
for i, h in enumerate(["Symbole", "Periode", "Trades", "Win Rate", "Statut"]):
    hdr3[i].text = h
    for run in hdr3[i].paragraphs[0].runs:
        run.bold = True

symbols = [
    ("Boom 1000", "2019-2026 M1", "1 438",  "74.20%", "PRODUCTION READY"),
    ("Crash 1000","2025 M1",       "138",    "30.43%", "REOPTIMISATION REQUISE"),
    ("Boom 500",  "A tester",      "-",      "-",      "EN ATTENTE"),
    ("Crash 500", "A tester",      "-",      "-",      "EN ATTENTE"),
    ("V75 (Vol)", "A tester",      "-",      "-",      "EN ATTENTE"),
]
for s_data in symbols:
    row = t_multi.add_row()
    for i, val in enumerate(s_data):
        row.cells[i].text = val

doc.add_paragraph()

# ── Note de mise a jour ──────────────────────────────────────────────────────
doc.add_heading("Note de Mise a Jour", level=2)
doc.add_paragraph(
    "Section 8 ajoutee le 01/06/2026. "
    "Backtest Crash 1000 realise avec parametres Test 3 (Boom 1000). "
    "Source : ReportTester-5775742_c1000.xlsx | Deriv-Demo Build 5836. "
    "Prochaine etape : reoptimisation parametres detection spike Crash 1000 "
    "puis nouveau backtest sur periode 2022-2026."
)

# ── Sauvegarder ─────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"Rapport mis a jour : {DOC_PATH}")
print("Section 8 (Crash 1000) ajoutee avec succes.")
